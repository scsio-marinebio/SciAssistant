# Copyright (c) 2025 Huawei Technologies Co., Ltd. All rights reserved.
import os
import json
import random
import subprocess
import requests
import re
import shutil
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import logging
from dataclasses import dataclass
from urllib.parse import urlparse
import tempfile
import time
import asyncio
import aiohttp
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from datetime import datetime, timedelta
import dateutil.parser
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlparse
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET
from urllib.parse import quote
from collections import Counter
import inspect
import sys
from functools import wraps
from typing import Optional

import feedparser
from .paper import Paper
from config.logging_config import get_logger
logger = get_logger()
# from markdown_pdf import MarkdownPdf, Section  # 改用 ReportLab

# ReportLab imports for PDF generation
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Preformatted, \
        Image as RLImage, HRFlowable, Flowable
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT
    
    # 尝试导入 Bookmark 和 TableOfContents（用于创建 PDF 书签）
    try:
        from reportlab.platypus.tableofcontents import TableOfContents
        from reportlab.platypus.paragraph import Paragraph as BaseParagraph
        BOOKMARK_AVAILABLE = True
    except ImportError:
        BOOKMARK_AVAILABLE = False
        logger.info("提示: TableOfContents 不可用，PDF 将不包含书签导航")

    REPORTLAB_AVAILABLE = True
except ImportError as e:
    REPORTLAB_AVAILABLE = False
    BOOKMARK_AVAILABLE = False
    logger.warning(f"警告: ReportLab 未安装或导入失败: {e}，请运行: pip install reportlab")
except Exception as e:
    REPORTLAB_AVAILABLE = False
    BOOKMARK_AVAILABLE = False
    logger.error(f"错误: ReportLab 导入时发生异常: {e}")

# 尝试导入matplotlib用于渲染数学公式
try:
    import matplotlib

    matplotlib.use('Agg')  # 使用非GUI后端
    import matplotlib.pyplot as plt
    from matplotlib import mathtext

    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    logger.warning("警告: matplotlib 未安装，数学公式将以文本形式显示。安装: pip install matplotlib")

try:
    from config.config import get_config
except ImportError:
    import sys

    sys.path.append(str(Path(__file__).parent.parent.parent))
    from config.config import get_config

# Import the optimized Faiss-based manager (fallback to JSON if Faiss not available)
try:
    from knowledge.vector_store import auto_index_task_completion_optimized, get_optimized_knowledge_manager

    FAISS_AVAILABLE = True
except ImportError:
    try:
        from ..knowledge.vector_store import auto_index_task_completion_optimized, get_optimized_knowledge_manager

        FAISS_AVAILABLE = True
    except ImportError:
        # Knowledge module not available, provide stub implementations
        FAISS_AVAILABLE = False


        def auto_index_task_completion_optimized(config, task_summary):
            """Stub implementation - knowledge module not available"""
            logging.getLogger(__name__).debug("Knowledge indexing skipped - module not available")
            return True


        def auto_index_task_completion(config, task_summary):
            """Stub implementation - knowledge module not available"""
            logging.getLogger(__name__).debug("Knowledge indexing skipped - module not available")
            return True


        get_optimized_knowledge_manager = None

logger = logging.getLogger(__name__)

proxy = {}


@dataclass
class MCPToolResult:
    """Standard result format for MCP tools"""
    success: bool
    data: Any = None
    error: str = None
    metadata: Dict[str, Any] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "success": self.success,
            "data": self.data,
            "error": self.error,
            "metadata": self.metadata or {}
        }


def _render_latex_to_image(latex_text: str, output_path: Path = None, fontsize: int = 12) -> Path:
    """
    使用matplotlib将LaTeX公式渲染为图片

    Args:
        latex_text: LaTeX格式的数学公式
        output_path: 输出图片路径,如果为None则自动生成
        fontsize: 字体大小

    Returns:
        图片文件路径
    """
    if not MATPLOTLIB_AVAILABLE:
        return None

    try:
        # 创建临时文件路径
        if output_path is None:
            import hashlib
            hash_name = hashlib.md5(latex_text.encode()).hexdigest()
            output_path = Path(tempfile.gettempdir()) / f"latex_{hash_name}.png"

        # 如果图片已存在,直接返回
        if output_path.exists():
            return output_path

        # 确保LaTeX文本被$包裹
        if not latex_text.startswith('$'):
            latex_text = f'${latex_text}$'

        # 创建图形
        fig = plt.figure(figsize=(8, 1))
        fig.patch.set_facecolor('white')

        # 渲染LaTeX
        text = fig.text(0.5, 0.5, latex_text,
                        fontsize=fontsize,
                        ha='center',
                        va='center',
                        usetex=False)  # 使用matplotlib的mathtext而不是真正的LaTeX

        # 调整边界
        fig.tight_layout(pad=0.1)

        # 保存为PNG
        plt.savefig(output_path,
                    dpi=150,
                    bbox_inches='tight',
                    pad_inches=0.05,
                    facecolor='white',
                    edgecolor='none')
        plt.close(fig)

        return output_path

    except Exception as e:
        logger.error(f"警告: LaTeX渲染失败: {e}")
        return None


def _wrap_special_symbol(symbol: str, fallback: str = None) -> str:
    """为特殊符号添加字体回退支持"""
    if fallback:
        return f'<font name="Arial">{symbol}</font>'
    return symbol


def _strip_escaped_font_tags(text: str) -> str:
    text = re.sub(r'&lt;\s*/?\s*font\b[^&]*?&gt;', '', text, flags=re.IGNORECASE)
    text = re.sub(r'&amp;lt;\s*/?\s*font\b[^&]*?&amp;gt;', '', text, flags=re.IGNORECASE)
    text = re.sub(r'&#60;\s*/?\s*font\b[^#]*?&#62;', '', text, flags=re.IGNORECASE)
    text = re.sub(r'&#x3c;\s*/?\s*font\b[^#]*?&#x3e;', '', text, flags=re.IGNORECASE)
    return text


def _strip_all_font_tags(text: str) -> str:
    text = re.sub(r'</?\s*font\b[^>]*>', '', text, flags=re.IGNORECASE)
    text = _strip_escaped_font_tags(text)
    return text


def _simplify_latex(latex_text: str) -> str:
    """
    简化LaTeX数学公式为可读文本
    将常见的LaTeX命令转换为Unicode数学符号

    Args:
        latex_text: LaTeX格式的数学公式

    Returns:
        简化后的文本
    """
 
    latex_text = _strip_all_font_tags(latex_text)
 
    # 常见LaTeX命令映射 - 使用有序字典确保处理顺序
    # 重要：必须先处理长命令，再处理短命令，避免部分匹配
    replacements = [
        # 希腊字母（按字母顺序）
        (r'\\alpha', 'α'),
        (r'\\beta', 'β'),
        (r'\\gamma', 'γ'),
        (r'\\Gamma', 'Γ'),
        (r'\\delta', 'δ'),
        (r'\\Delta', 'Δ'),
        (r'\\epsilon', 'ε'),
        (r'\\varepsilon', 'ε'),
        (r'\\zeta', 'ζ'),
        (r'\\eta', 'η'),
        (r'\\theta', 'θ'),
        (r'\\Theta', 'Θ'),
        (r'\\iota', 'ι'),
        (r'\\kappa', 'κ'),
        (r'\\lambda', 'λ'),
        (r'\\Lambda', 'Λ'),
        (r'\\mu', 'μ'),
        (r'\\nu', 'ν'),
        (r'\\xi', 'ξ'),
        (r'\\Xi', 'Ξ'),
        (r'\\pi', 'π'),
        (r'\\Pi', 'Π'),
        (r'\\rho', 'ρ'),
        (r'\\sigma', 'σ'),
        (r'\\Sigma', 'Σ'),
        (r'\\tau', 'τ'),
        (r'\\upsilon', 'υ'),
        (r'\\Upsilon', 'Υ'),
        (r'\\phi', 'φ'),
        (r'\\Phi', 'Φ'),
        (r'\\varphi', 'φ'),
        (r'\\chi', 'χ'),
        (r'\\psi', 'ψ'),
        (r'\\Psi', 'Ψ'),
        (r'\\omega', 'ω'),
        (r'\\Omega', 'Ω'),

        # 数学运算符
        # 注意：某些符号在宋体中可能无法显示，使用备选方案
        # nabla: 使用倒三角形 ▽ (U+25BD) 代替 ∇ (U+2207)，因为后者在某些字体中缺失
        (r'\\nabla', '▽'),
        # partial: 宋体不支持 ∂ (U+2202)
        # 使用 ð (U+00F0 eth) 代替，这是拉丁字母，看起来很像 ∂，大多数字体都支持
        (r'\\partial', 'ð'),  # 使用 eth 代替偏导数符号
        (r'\\infty', '∞'),

        # 关系符号
        (r'\\leq', '≤'),
        (r'\\le', '≤'),
        (r'\\geq', '≥'),
        (r'\\ge', '≥'),
        (r'\\neq', '≠'),
        (r'\\ne', '≠'),
        (r'\\approx', '≈'),
        (r'\\equiv', '≡'),
        (r'\\sim', '∼'),
        (r'\\propto', '∝'),

        # 二元运算符
        (r'\\times', '×'),
        (r'\\cdot', '·'),
        (r'\\div', '÷'),
        (r'\\pm', '±'),
        (r'\\mp', '∓'),
        (r'\\oplus', '⊕'),
        (r'\\otimes', '⊗'),

        # 集合符号
        (r'\\in', '∈'),
        (r'\\notin', '∉'),
        (r'\\subset', '⊂'),
        (r'\\supset', '⊃'),
        (r'\\subseteq', '⊆'),
        (r'\\supseteq', '⊇'),
        (r'\\cup', '∪'),
        (r'\\cap', '∩'),
        (r'\\emptyset', '∅'),

        # 逻辑符号
        (r'\\forall', '∀'),
        (r'\\exists', '∃'),
        (r'\\neg', '¬'),
        (r'\\land', '∧'),
        (r'\\lor', '∨'),

        # 大型运算符
        (r'\\sum', '∑'),
        (r'\\prod', '∏'),
        (r'\\int', '∫'),
        (r'\\oint', '∮'),
        (r'\\iint', '∬'),
        (r'\\iiint', '∭'),

        # 箭头
        (r'\\rightarrow', '→'),
        (r'\\to', '→'),
        (r'\\leftarrow', '←'),
        (r'\\gets', '←'),
        (r'\\leftrightarrow', '↔'),
        (r'\\Rightarrow', '⇒'),
        (r'\\Leftarrow', '⇐'),
        (r'\\Leftrightarrow', '⇔'),

        # 其他常用符号
        (r'\\sqrt', '√'),
        (r'\\angle', '∠'),
        (r'\\perp', '⊥'),
        (r'\\parallel', '∥'),

        # 上下标和特殊格式（具体数字）
        # 注意：不要在这里转换为Unicode上下标字符（₀₁₂等），因为宋体可能不支持
        # 应该在后面统一转换为HTML <sub>/<sup> 标签
        # 只保留常用的平方、立方等特殊符号
        (r'\^2', '²'),
        (r'\^3', '³'),

        # 文本命令
        (r'\\text\{([^}]+)\}', r'\1'),
        (r'\\mathrm\{([^}]+)\}', r'\1'),
        (r'\\mathbf\{([^}]+)\}', r'\1'),
        (r'\\mathit\{([^}]+)\}', r'\1'),

        # 分数(简化显示)
        (r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)'),

        # 帽子和修饰符号
        # 注意：组合字符在 ReportLab 中可能无法正确显示，使用简化表示
        (r'\\hat\{([^}]+)\}', r'\1^'),  # 使用上标^表示帽子
        (r'\\hat ([a-zA-Z])', r'\1^'),
        (r'\\bar\{([^}]+)\}', r'\1‾'),  # 使用上划线
        (r'\\tilde\{([^}]+)\}', r'\1~'),  # 使用波浪号
        (r'\\vec\{([^}]+)\}', r'<b>\1</b>'),  # 向量使用粗体表示（标准数学记号）

        # 括号
        (r'\\left\(', '('),
        (r'\\right\)', ')'),
        (r'\\left\[', '['),
        (r'\\right\]', ']'),
        (r'\\left\{', '{'),
        (r'\\right\}', '}'),
        (r'\\left', ''),
        (r'\\right', ''),
        (r'\\big', ''),
        (r'\\Big', ''),
        (r'\\bigg', ''),
        (r'\\Bigg', ''),

        # 空格
        (r'\\\\', ' '),  # 换行
        (r'\\,', ' '),  # 小空格
        (r'\\;', ' '),  # 中空格
        (r'\\quad', '  '),  # 大空格
        (r'\\qquad', '    '),  # 更大空格
    ]

    result = latex_text

    # 按顺序应用所有替换
    for pattern, replacement in replacements:
        result = re.sub(pattern, replacement, result)

    # 处理通用的上标 ^{...} 转换为HTML
    def convert_superscript(match):
        content = match.group(1)
        return f'<sup>{content}</sup>'

    result = re.sub(r'\^\{([^}]+)\}', convert_superscript, result)
    result = re.sub(r'\^([0-9a-zA-Z])', r'<sup>\1</sup>', result)

    # 处理通用的下标 _{...} 转换为HTML
    def convert_subscript(match):
        content = match.group(1)
        return f'<sub>{content}</sub>'

    result = re.sub(r'\_\{([^}]+)\}', convert_subscript, result)
    result = re.sub(r'\_([0-9a-zA-Z])', r'<sub>\1</sub>', result)

    # 清理剩余的反斜杠和花括号
    # 注意：这里要小心，不要清理掉已经转换好的Unicode符号
    result = result.replace('\\', '')
    result = re.sub(r'\{([^}]*)\}', r'\1', result)

    # 后处理：为可能无法显示的符号添加字体标记或替换
    # 这些符号在宋体中可能缺失，需要特殊处理
    problematic_symbols = {
        # 偏导数符号：已在前面替换为 ð (eth)
        # nabla: 已在前面替换为 ▽
        # 向量箭头：组合字符，需要移除
        '⃗': ('', 'vec'),  # 向量箭头（组合字符）- 移除
    }

    # 检查并标记特殊符号
    for symbol, (replacement, name) in problematic_symbols.items():
        if symbol in result and replacement is not None:
            # 如果有替代符号，直接替换
            result = result.replace(symbol, replacement)

    return result


def _escape_url_for_html_attr(url: str) -> str:
    """转义 URL 中的特殊字符，使其可以安全地用于 HTML href 属性"""
    if not url:
        return url
    # 转义 & 为 &amp;（必须首先处理，避免重复转义）
    url = url.replace('&', '&amp;')
    # 转义引号
    url = url.replace('"', '&quot;')
    return url


def _process_inline_formatting(text: str) -> str:
    """
    将 Markdown 行内格式转换为 ReportLab 可解析的安全 HTML，
    保证标签平衡，避免 PDF 生成时的解析错误。
    """
    
    # 先将 id 转换为 name
    text = re.sub(r'<a\s+id="([^"]+)"', r'<a name="\1"', text)

    # 处理带 style 属性的链接，转义 URL 中的特殊字符
    def fix_styled_link(m):
        url = m.group(1)
        safe_url = _escape_url_for_html_attr(url)
        return f'<a href="{safe_url}" color="#04B5BB">'
    
    text = re.sub(
        r'<a\s+href="([^"]+)"\s+style="[^"]*">',
        fix_styled_link,
        text
    )

    # 3. 为没有颜色属性的 href 链接添加颜色，同时转义 URL
    # 匹配: <a href="..."> (但不匹配已有color属性的)
    def fix_uncolored_link(m):
        url = m.group(1)
        safe_url = _escape_url_for_html_attr(url)
        return f'<a href="{safe_url}" color="#04B5BB">'
    
    text = re.sub(
        r'<a\s+href="([^"]+)"(?!\s+color)>',
        fix_uncolored_link,
        text
    )
    
    # 恢复并增强Unicode上下标处理
    # 这一步非常关键，因为用户经常直接复制粘贴包含Unicode上标（如 ⁻¹⁶）的文本
    # 而这些字符在标准中文字体（如宋体）中通常不支持，导致显示为空白

    # Unicode上标映射
    superscript_map = {
        '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4',
        '⁵': '5', '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9',
        '⁺': '+', '⁻': '-', '⁼': '=', '⁽': '(', '⁾': ')', 'ⁿ': 'n'
    }

    # Unicode下标映射
    subscript_map = {
        '₀': '0', '₁': '1', '₂': '2', '₃': '3', '₄': '4',
        '₅': '5', '₆': '6', '₇': '7', '₈': '8', '₉': '9',
        '₊': '+', '₋': '-', '₌': '=', '₍': '(', '₎': ')',
        'ₐ': 'a', 'ₑ': 'e', 'ₒ': 'o', 'ₓ': 'x', 'ₕ': 'h',
        'ₖ': 'k', 'ₗ': 'l', 'ₘ': 'm', 'ₙ': 'n', 'ₚ': 'p',
        'ₛ': 's', 'ₜ': 't'
    }

    # 转换Unicode上标为HTML sup标签
    for unicode_char, normal_char in superscript_map.items():
        if unicode_char in text:
            text = text.replace(unicode_char, f'<sup>{normal_char}</sup>')

    # 转换Unicode下标为HTML sub标签
    for unicode_char, normal_char in subscript_map.items():
        if unicode_char in text:
            text = text.replace(unicode_char, f'<sub>{normal_char}</sub>')

    # 先保护数学公式 $...$ (必须在数学符号处理之前，避免公式内的希腊字母被错误地包裹font标签)
    math_formulas = []

    def protect_math(match):
        formula = match.group(1)
        placeholder = f"__MATH_FORMULA_{len(math_formulas)}__"
        math_formulas.append(formula)
        return placeholder

    # 【关键】在匹配$...$公式之前，先保护所有<a>标签（含href/name属性）
    # 防止$...$公式匹配跨越HTML锚点标签，导致_simplify_latex将属性中的_转为<sub>
    # 例如: "公式$x$的值...<a href="#ref-msg_123">...$y$" 中两个$之间包含href属性
    anchor_tag_placeholders = []
    def protect_anchor_tag(match):
        placeholder = f"__ANCHOR_TAG_{len(anchor_tag_placeholders)}__"
        anchor_tag_placeholders.append(match.group(0))
        return placeholder
    # 保护完整的<a>标签（包括开标签和闭标签之间的所有内容，以及自闭合的<a name="..."></a>）
    text = re.sub(r'<a\s[^>]*>.*?</a>', protect_anchor_tag, text, flags=re.IGNORECASE | re.DOTALL)
    # 保护独立的<a name="..."></a>锚点定义
    text = re.sub(r'<a\s[^>]*></a>', protect_anchor_tag, text, flags=re.IGNORECASE)

    # 保护行内数学公式 $...$
    # 使用非贪婪匹配,确保只匹配成对的$
    text = re.sub(r'\$([^\$]+?)\$', protect_math, text)

    # 恢复锚点标签占位符
    for i, original in enumerate(anchor_tag_placeholders):
        text = text.replace(f"__ANCHOR_TAG_{i}__", original)

    # 清理孤立的$符号(没有配对的)
    # 统计剩余的$数量,如果是奇数,说明有孤立的$
    dollar_count = text.count('$')
    if dollar_count > 0:
        # 移除所有剩余的孤立$符号
        text = text.replace('$', '')

    # 特殊字符处理（必须在公式保护之后，避免污染公式内容）
    text = text.replace('μ', 'µ').replace('µ', '<font name="Arial">µ</font>')
    text = text.replace('ŷ', '<font name="Arial">ŷ</font>')

    # 处理宋体不支持的数学符号，使用 Arial 字体显示（Windows 系统自带）
    # 注意：这里只处理公式外的数学符号，公式内的符号由 _simplify_latex 处理
    math_symbols = [
        # 希腊字母
        'α', 'β', 'γ', 'δ', 'ε', 'ζ', 'η', 'θ', 'ι', 'κ', 'λ', 'ν', 'ξ', 'ο', 'π',
        'ρ', 'σ', 'τ', 'υ', 'φ', 'χ', 'ψ', 'ω',
        'Α', 'Β', 'Γ', 'Δ', 'Ε', 'Ζ', 'Η', 'Θ', 'Ι', 'Κ', 'Λ', 'Μ', 'Ν', 'Ξ', 'Ο',
        'Π', 'Ρ', 'Σ', 'Τ', 'Υ', 'Φ', 'Χ', 'Ψ', 'Ω',
        # 数学运算符
        '∇', '∂', '∞', '∑', '∏', '∫', '∬', '∭', '∮', '√',
        '≤', '≥', '≠', '≈', '≡', '∼', '∝', '±', '∓', '×', '÷',
        '∈', '∉', '⊂', '⊃', '⊆', '⊇', '∪', '∩', '∅',
        '∀', '∃', '¬', '∧', '∨',
        '→', '←', '↔', '⇒', '⇐', '⇔', '↑', '↓',
        '∠', '⊥', '∥', '▽',
    ]
    for sym in math_symbols:
        if sym in text:
            text = text.replace(sym, f'<font name="Arial">{sym}</font>')

    # 处理文献引用格式（必须在普通Markdown链接之前处理）
    # 格式1: [数字] 标题/文件名，URL.pdf，日期 - PDF文件引用（根据URL是否以.pdf结尾判断）
    def replace_pdf_reference(match):
        num = match.group(1)
        title = match.group(2)  # 标题或文件名（可以不含.pdf）
        url = match.group(3)  # PDF URL
        date = match.group(4) if len(match.groups()) >= 4 and match.group(4) else ''
        
        # 清理和转义 URL
        url = url.strip()
        url = re.sub(r'["\'/]+$', '', url)
        url_escaped = url.replace('&', '&amp;')
        
        # 使用回形针图标📎 (U+1F4CE) 表示可下载的PDF文档
        # 处理无法确定月份的情况，只显示年份
        if date and date.strip():
            date_str = date.strip()
            if '无法确定月份' in date_str:
                # 提取年份（匹配4位数字+年）
                year_match = re.search(r'(\d{4})年', date_str)
                if year_match:
                    date_str = f'{year_match.group(1)}年'
            date_part = f', {date_str}'
        else:
            date_part = ''
        
        # 使用 📎 图标表示PDF文件（已下载完整内容）
        # 注意：添加style属性确保图标不受父元素斜体样式影响
        return f'[{num}] <font name="EmojiFont" style="font-style: normal;">📎</font> {title}, <a href="{url_escaped}" color="#04B5BB">{url_escaped}</a>{date_part}'

    # 改进的正则表达式0：优先匹配标题包含.pdf的引用（不管URL是什么）
    # 这样可以正确识别标题含.pdf但URL不是PDF的情况，如：[42] 2401.10359v1.pdf, https://arxiv.org/abs/2401.10359
    # 使用负向前瞻，跳过已经有图标的引用（<font name="EmojiFont">）
    text = re.sub(r'\[(\d+)\]\s(?!<font name="EmojiFont">)(.+?\.pdf.+?)，(https?://[^\s，]+)(?:，(.+?))?(?=\s*\n|\s*$)',
                  replace_pdf_reference, text, flags=re.IGNORECASE | re.MULTILINE)
    
    # 改进的正则表达式1：根据URL是否以.pdf结尾来判断PDF引用（不管标题是什么）
    # 这样可以正确识别标题不含.pdf但URL是PDF的情况，如：[1] BD CD Marker Handbook, https://example.com/file.pdf
    # 使用负向前瞻，跳过已经有图标的引用（<font name="EmojiFont">）
    text = re.sub(r'\[(\d+)\]\s(?!<font name="EmojiFont">)(.+?)，(https?://[^\s，]+?\.pdf)(?:，(.+?))?(?=\s*\n|\s*$)',
                  replace_pdf_reference, text, flags=re.IGNORECASE | re.MULTILINE)

    # 改进的正则表达式1.1：匹配arxiv.org/pdf/格式的PDF链接（不以.pdf结尾）
    # 使用负向前瞻，跳过已经有图标的引用
    text = re.sub(r'\[(\d+)\]\s(?!<font name="EmojiFont">)(.+?)，(https?://arxiv\.org/pdf/[^\s，]+)(?:，(.+?))?(?=\s*\n|\s*$)',
                  replace_pdf_reference, text, flags=re.IGNORECASE | re.MULTILINE)

    # 格式1.2: [数字] 标题，学术论文URL，日期 - 学术论文引用（使用 📄 图标）
    def replace_academic_reference(match):
        num = match.group(1)
        title = match.group(2)
        url = match.group(3)
        date_str = match.group(4) if match.group(4) else None
        
        # 清理和转义 URL
        url = url.strip()
        url = re.sub(r'["\'/]+$', '', url)
        url_escaped = url.replace('&', '&amp;')
        
        if date_str:
            # 清理日期字符串
            if '无法确定月份' in date_str:
                year_match = re.search(r'(\d{4})年', date_str)
                if year_match:
                    date_str = f'{year_match.group(1)}年'
            date_part = f', {date_str}'
        else:
            date_part = ''
        
        # 使用 📄 图标表示学术论文（已引用完整内容）
        # 注意：添加style属性确保图标不受父元素斜体样式影响
        return f'[{num}] <font name="EmojiFont" style="font-style: normal;">📄</font> {title}, <a href="{url_escaped}" color="#04B5BB">{url_escaped}</a>{date_part}'
    
    # 匹配学术论文引用：arXiv/PubMed/学术期刊网站（非 PDF 链接）
    # 使用负向前瞻，跳过已经有图标的引用
    # 扩展范围：包含 Nature、Lancet、MDPI、Frontiers、Springer、Wiley 等学术期刊
    # 添加 (?:www\.)? 支持，因为很多学术网站URL包含www前缀（如www.mdpi.com）
    academic_pattern = r'\[(\d+)\]\s(?!<font name="EmojiFont">)(.+?)，(https?://(?:(?:www\.)?arxiv\.org/(?:abs|html)/|pubmed\.ncbi\.nlm\.nih\.gov/|pmc\.ncbi\.nlm\.nih\.gov/|ncbi\.nlm\.nih\.gov/pubmed/|(?:www\.)?medrxiv\.org/content/|(?:www\.)?biorxiv\.org/content/|(?:www\.)?nature\.com/articles/(?!d41586-)|(?:www\.)?thelancet\.com/(?:journals/|article/)|(?:www\.)?science\.org/doi/|(?:www\.)?cell\.com/|(?:www\.)?nejm\.org/doi/|(?:www\.)?mdpi\.com/|(?:www\.)?frontiersin\.org/(?:journals/|articles/)|(?:www\.)?plos\.org/|link\.springer\.com/article/|onlinelibrary\.wiley\.com/doi/|(?:www\.)?bmj\.com/content/|jamanetwork\.com/journals/)[^\s，]+)(?:，(.+?))?(?=\s*\n|\s*$)'
    text = re.sub(academic_pattern, replace_academic_reference, text, flags=re.IGNORECASE | re.MULTILINE)

    # 格式2: [数字] 标题，URL，日期 - 网页引用
    def replace_web_reference(match):
        num = match.group(1)
        title = match.group(2)
        url = match.group(3)
        date_str = match.group(4) if match.group(4) else None
        
        # 清理 URL 中可能的特殊字符和 HTML 标签残留
        url = url.strip()
        # 移除 URL 末尾可能的引号、斜杠等
        url = re.sub(r'["\'/]+$', '', url)
        # 转义 URL 中的特殊字符（如 &）
        url_escaped = url.replace('&', '&amp;')
        
        if date_str:
            # 清理日期字符串中的"无法确定月份"等文本
            if '无法确定月份' in date_str:
                # 提取年份（匹配4位数字+年）
                year_match = re.search(r'(\d{4})年', date_str)
                if year_match:
                    date_str = f'{year_match.group(1)}年'
            date_part = f', {date_str}'
        else:
            date_part = ''
        # 使用font标签指定emoji字体来显示图标，ReportLab会自动回退到支持该字符的字体
        # 注意：添加style属性确保图标不受父元素斜体样式影响
        return f'[{num}] <font name="EmojiFont" style="font-style: normal;">🌐</font> {title}, <a href="{url_escaped}" color="#04B5BB">{url_escaped}</a>{date_part}'

    # 改进的正则表达式2：匹配非PDF的URL引用（网页引用）
    # 使用负向前瞻，跳过已经有图标的引用（<font name="EmojiFont">）
    text = re.sub(r'\[(\d+)\]\s(?!<font name="EmojiFont">)(.+?)，(https?://[^\s，]+?)(?:，(.+?))?(?=\s*\n|\s*$)',
                  replace_web_reference, text, flags=re.IGNORECASE | re.MULTILINE)

    # 处理 Markdown 链接 [text](url) (必须在其他格式之前处理)
    def replace_markdown_link(match):
        link_text = match.group(1)
        link_url = match.group(2)
        # ReportLab 的 <a> 标签格式
        return f'<a href="{link_url}" color="#04B5BB">{link_text}</a>'

    # 处理Markdown链接，支持URL中包含括号（如wiki链接、论文DOI等）
    # 匹配模式：URL可包含一层嵌套括号，如 https://example.com/page(1).html
    text = re.sub(r'\[([^\]]+)\]\(([^()\s]*(?:\([^)]*\)[^()\s]*)*)\)', replace_markdown_link, text)

    # 处理行内代码 `code` (必须在粗体和斜体之前处理,避免冲突)
    # 使用占位符保护代码块，防止后续的粗体/斜体正则误匹配代码块内的HTML标签或内容
    inline_codes = []

    def protect_code(match):
        code_content = match.group(1)
        # HTML 转义
        code_content = code_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        # 检查是否包含中文字符
        has_chinese = bool(re.search(r'[\u4e00-\u9fff]', code_content))

        if has_chinese:
            # 包含中文,使用宋体,稍小字号,添加灰色背景
            html = f'<font name="SimSun" size="9" color="#333333" backColor="#f5f5f5">{code_content}</font>'
        else:
            # 纯英文/数字,使用 Arial 字体
            html = f'<font name="Arial" size="9">{code_content}</font>'

        placeholder = f"__INLINE_CODE_{len(inline_codes)}__"
        inline_codes.append(html)
        return placeholder

    text = re.sub(r'`([^`]+)`', protect_code, text)

    # 粗斜体：优先处理 ***text***
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<font name="SimHei"><b><i>\1</i></b></font>', text)
    # 粗体：非贪婪匹配，允许中间包含其他字符（如斜体的*）
    text = re.sub(r'\*\*(.+?)\*\*', r'<font name="SimHei"><b>\1</b></font>', text)
    # 斜体：非贪婪匹配，排除**的情况，且要求内容两侧无空白（避免匹配数学公式中的*）
    text = re.sub(r'(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)', r'<i>\1</i>', text)

    for i, html in enumerate(inline_codes):
        text = text.replace(f"__INLINE_CODE_{i}__", html)
    for i, formula in enumerate(math_formulas):
        display = _simplify_latex(formula)
        text = text.replace(f"__MATH_FORMULA_{i}__", f'<font size="9.5"><i>{display}</i></font>')

    text = _strip_escaped_font_tags(text)

    # 移除不支持或无意义的标签
    text = re.sub(r'</?\s*nobr\b[^>]*>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<hr\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<div\b[^>]*>(.*?)</div>', r'\1', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'</?div\b[^>]*>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<span\b[^>]*>(.*?)</span>', r'\1', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'</?span\b[^>]*>', '', text, flags=re.IGNORECASE)
    for tag in ['section', 'article', 'header', 'footer', 'nav', 'aside', 'main']:
        text = re.sub(rf'<{tag}\b[^>]*>(.*?)</{tag}>', r'\1', text, flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(rf'</?\s*{tag}\b[^>]*>', '', text, flags=re.IGNORECASE)

    text = re.sub(r'<br\s*>', '<br/>', text, flags=re.IGNORECASE)
    text = re.sub(r'<br\s*>([^<]*)</br>', r'<br/> \1', text, flags=re.IGNORECASE)
    text = re.sub(r'</br>', '', text, flags=re.IGNORECASE)
    text = text.replace('<br/>', ' <br/> ')

    # 修复 HTML 标签中没有引号的属性值（仅处理标签级属性，不处理 href 值内的 URL 参数）
    # 使用负向前瞻确保不在引号内匹配
    # 只匹配：空格 + 属性名 + = + 非引号值（不在已有引号属性值内）
    def fix_unquoted_attr_safe(text_val):
        # 先保护所有已有的引号属性值
        protected = []
        def protect(m):
            protected.append(m.group(0))
            return f'__PROTECTED_ATTR_{len(protected)-1}__'
        
        # 保护 attr="value" 和 attr='value' 格式
        text_val = re.sub(r'(\w+)="[^"]*"', protect, text_val)
        text_val = re.sub(r"(\w+)='[^']*'", protect, text_val)
        
        # 现在修复未引号的属性（只在标签内，且前面是空格）
        for attr in ['color', 'size', 'name', 'href', 'face', 'backColor']:
            text_val = re.sub(rf'(\s){attr}=([^\s>"\']+)', rf'\1{attr}="\2"', text_val, flags=re.IGNORECASE)
        
        # 恢复保护的属性
        for i, val in enumerate(protected):
            text_val = text_val.replace(f'__PROTECTED_ATTR_{i}__', val)
        
        return text_val
    
    # 只对 HTML 标签应用修复
    text = re.sub(r'<[^>]+>', lambda m: fix_unquoted_attr_safe(m.group(0)), text)

    def _sanitize_reportlab_links(value: str) -> str:
        def repl(m: re.Match) -> str:
            href = (m.group(1) or '').strip()
            body = m.group(2) or ''
            full_tag = m.group(0)
            if re.match(r'^(https?://|mailto:|#)', href, flags=re.IGNORECASE):
                # 对于有效的 URL，确保转义特殊字符
                safe_href = _escape_url_for_html_attr(href)
                # 替换原始 href 为转义后的版本
                return full_tag.replace(f'href="{href}"', f'href="{safe_href}"')
            if re.match(r'^www\.', href, flags=re.IGNORECASE) or re.match(r'^[\w.-]+\.[a-z]{2,}([/?#]|$)', href, flags=re.IGNORECASE):
                safe_href = _escape_url_for_html_attr(f'https://{href}')
                return f'<a href="{safe_href}" color="#04B5BB">{body}</a>'
            return body

        return re.sub(r'<a\s+href="([^"]+)"[^>]*>(.*?)</a>', repl, value, flags=re.IGNORECASE | re.DOTALL)

    text = _sanitize_reportlab_links(text)

    text = re.sub(r'<(font|b|i|sub|sup)\b[^>]*>\s*</\1>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<a(?![^>]*\bname=)[^>]*>\s*</a>', '', text, flags=re.IGNORECASE)

    # 在 _sanitize_reportlab_links 之后保护 HTML 标签中的属性值
    # 避免 _apply_english_font_markup 等后续处理影响锚点名称
    protected_attrs = []
    def protect_attr_value(m):
        attr_name = m.group(1)
        attr_value = m.group(2)
        # 使用不包含下划线的占位符，避免被下标转换影响
        placeholder = f"PROTECTEDATTRVAL{len(protected_attrs)}ENDPROTECTED"
        protected_attrs.append((attr_name, attr_value))
        return f'{attr_name}="{placeholder}"'
    
    # 保护 name, href 属性值
    text = re.sub(r'(name|href)="([^"]*)"', protect_attr_value, text, flags=re.IGNORECASE)

    text = _apply_english_font_markup(text)

    # 处理特殊符号，确保使用支持这些符号的字体，并设置颜色
    # 勾号设为绿色
    green_symbols = ['✓', '✔']
    for sym in green_symbols:
        text = text.replace(sym, f'<font name="SymbolFont" color="green">{sym}</font>')

    # 叉号设为红色
    red_symbols = ['✕', '✖', '✗', '✘']
    for sym in red_symbols:
        text = text.replace(sym, f'<font name="SymbolFont" color="red">{sym}</font>')

    # 转义非HTML标签的裸尖括号（如 <S, A, P, R, γ>），避免ReportLab误解析
    # 策略：保护合法HTML标签 → 转义剩余的 < > → 恢复合法标签
    _valid_html_tags = re.compile(r'</?(?:a|font|b|i|sub|sup|br|para)\b[^>]*/?>', re.IGNORECASE)
    _html_placeholders = []
    def _protect_valid_tag(m):
        placeholder = f"__VALID_HTML_{len(_html_placeholders)}__"
        _html_placeholders.append(m.group(0))
        return placeholder
    text = _valid_html_tags.sub(_protect_valid_tag, text)
    # 转义剩余的裸尖括号
    text = text.replace('<', '&lt;').replace('>', '&gt;')
    # 恢复合法HTML标签
    for i, original in enumerate(_html_placeholders):
        text = text.replace(f"__VALID_HTML_{i}__", original)

    # 平衡内联标签，修正缺失或多余的闭合
    tag_regex = re.compile(r'</?\s*(a|font|b|i|sub|sup)\b[^>]*>', re.IGNORECASE)

    def balance_inline_tags(value: str) -> str:
        parts: List[str] = []
        stack: List[str] = []
        last = 0
        for m in tag_regex.finditer(value):
            parts.append(value[last:m.start()])
            token = m.group(0)
            name = m.group(1).lower()
            closing = token.startswith('</')
            if closing:
                if name in stack:
                    while stack and stack[-1] != name:
                        parts.append(f'</{stack.pop()}>')
                    if stack and stack[-1] == name:
                        stack.pop()
                        parts.append(token)
                # 未匹配的孤立闭合直接丢弃
            else:
                stack.append(name)
                parts.append(token)
            last = m.end()
        parts.append(value[last:])
        while stack:
            parts.append(f'</{stack.pop()}>')
        return ''.join(parts)

    # 恢复被保护的属性值
    for i, (attr_name, attr_value) in enumerate(protected_attrs):
        text = text.replace(f'PROTECTEDATTRVAL{i}ENDPROTECTED', attr_value)
    
    return balance_inline_tags(text)


_EN_TOKEN_RE = re.compile(r"[A-Za-z][A-Za-z0-9._@/+\-]*")


def _apply_english_font_markup(text: str, font_name: str = "Arial") -> str:
    parts = re.split(r'(<[^>]+>)', text)
    font_stack: List[bool] = []
    in_anchor: int = 0  # 跟踪是否在 <a> 标签内（支持嵌套计数）

    def in_locked_font() -> bool:
        return any(font_stack)

    def wrap_tokens(s: str) -> str:
        return _EN_TOKEN_RE.sub(lambda m: f'<font name="{font_name}">{m.group(0)}</font>', s)

    out: List[str] = []
    for part in parts:
        if not part:
            continue
        if part.startswith('<') and part.endswith('>'):
            # 跟踪 <a> 标签的开始和结束
            if re.match(r'<\s*a\b', part, flags=re.IGNORECASE) and not part.rstrip().endswith('/>'):
                in_anchor += 1
            elif re.match(r'</\s*a\s*>', part, flags=re.IGNORECASE):
                in_anchor = max(0, in_anchor - 1)
            
            if re.match(r'<\s*font\b', part, flags=re.IGNORECASE) and not part.rstrip().endswith('/>'):
                locked = bool(re.search(r'\b(name|face)\s*=', part, flags=re.IGNORECASE))
                font_stack.append(locked)
            elif re.match(r'</\s*font\s*>', part, flags=re.IGNORECASE):
                if font_stack:
                    font_stack.pop()
            out.append(part)
            continue

        # 在 <a> 标签内或已锁定字体时，不进行英文字体包裹
        out.append(part if (in_locked_font() or in_anchor > 0) else wrap_tokens(part))

    return ''.join(out)


class PDFBookmark(Flowable):
    """自定义 PDF 书签类，用于在 PDF 中创建可点击的目录导航"""
    def __init__(self, title: str, level: int, key: str = None):
        Flowable.__init__(self)
        self.title = title
        self.level = level
        self.key = key or f'bookmark_{id(self)}'
    
    def wrap(self, availWidth, availHeight):
        """计算书签占用的空间（不占用空间）"""
        return 0, 0

    def draw(self):
        """在 PDF 中绘制（添加）书签"""
        # 使用 canvas.bookmarkPage 创建书签锚点
        self.canv.bookmarkPage(self.key)
        # 使用 canvas.addOutlineEntry 添加到 PDF 大纲（目录）
        self.canv.addOutlineEntry(self.title, self.key, level=self.level)


def _find_font_with_priority(font_name: str, system_paths: List[str]) -> Optional[str]:
    """
    按优先级查找字体文件：先查找 ./Font/ 目录，再回退到系统路径
    
    Args:
        font_name: 字体名称（如 'simhei', 'simsun', 'arial'）
        system_paths: 系统字体路径列表（按优先级排序）
    
    Returns:
        找到的字体文件路径，如果都找不到则返回 None
    """
    # 1. 优先从项目根目录的 Font/ 目录中查找
    # 获取当前文件所在目录，然后向上查找项目根目录
    current_file = Path(__file__).resolve()
    # 从 deepdiver_v2/src/tools/mcp_tools.py 向上3级到项目根目录
    project_root = current_file.parent.parent.parent.parent
    font_dir = project_root / "Font"

    if font_dir.exists() and font_dir.is_dir():
        # 支持的字体文件扩展名
        extensions = ['.ttf', '.ttc', '.TTF', '.TTC']
        
        for ext in extensions:
            # 尝试不同的文件名格式
            font_file = font_dir / f"{font_name}{ext}"
            if font_file.exists():
                logger.info(f"找到字体文件: {font_file}")
                return str(font_file.absolute())
            
            # 尝试大写文件名
            font_file_upper = font_dir / f"{font_name.upper()}{ext}"
            if font_file_upper.exists():
                logger.info(f"找到字体文件: {font_file_upper}")
                return str(font_file_upper.absolute())
    else:
        logger.warning(f"Font目录不存在: {font_dir}")
    # 2. 回退到系统路径查找
    for system_path in system_paths:
        if system_path and os.path.exists(system_path):
            logger.info(f"使用系统字体: {system_path}")
            return system_path
    
    # 3. 都找不到
    logger.warning(f"字体 '{font_name}' 在 {font_dir} 和系统路径中均未找到")
    return None


def generate_pdf_with_reportlab(markdown_content: str, output_path: Path) -> bool:
    """
    使用 ReportLab 将 Markdown 内容转换为 PDF，支持中文字体（黑体标题，宋体正文）

    Args:
        markdown_content: Markdown 格式的内容
        output_path: 输出 PDF 文件路径

    Returns:
        bool: 是否成功生成 PDF
    """
    if not REPORTLAB_AVAILABLE:
        logger.error("错误: ReportLab 未安装")
        return False

    try:
        # 注册中文字体（优先从 ./Font/ 目录加载，回退到系统字体）
        import platform
        system = platform.system()

        # 根据操作系统选择系统字体路径（作为回退选项）
        if system == "Windows":
            # Windows 系统字体路径
            simsun_system_paths = ["C:/Windows/Fonts/simsun.ttf"]
            simhei_system_paths = ["C:/Windows/Fonts/simhei.ttf"]
            arial_system_paths = ["C:/Windows/Fonts/arial.ttf"]
            symbol_system_paths = ["C:/Windows/Fonts/seguisym.ttf"]
            emoji_system_paths = ["C:/Windows/Fonts/seguiemj.ttf", "C:/Windows/Fonts/seguisym.ttf"]
        elif system == "Linux":
            # Linux 系统字体路径
            simsun_system_paths = ["/usr/share/fonts/dejavu/SIMSUN.TTC"]
            simhei_system_paths = ["/usr/share/fonts/dejavu/SIMHEI.TTF"]
            arial_system_paths = ["/usr/share/fonts/dejavu/ARIAL.TTF"]
            symbol_system_paths = ["/usr/share/fonts/dejavu/DejaVuSans.ttf"]
            emoji_system_paths = [
                "/usr/share/fonts/truetype/noto/NotoEmoji-Regular.ttf",
                "/usr/share/fonts/noto/NotoEmoji-Regular.ttf",
                "/usr/share/fonts/google-noto-emoji/NotoEmoji-Regular.ttf",
                "/usr/share/fonts/dejavu/DejaVuSans.ttf"
            ]
        else:  # macOS
            simsun_system_paths = ["/System/Library/Fonts/STHeiti Light.ttc"]
            simhei_system_paths = ["/System/Library/Fonts/STHeiti Medium.ttc"]
            arial_system_paths = ["/Library/Fonts/Arial.ttf"]
            symbol_system_paths = ["/System/Library/Fonts/AppleSymbols.ttf"]
            emoji_system_paths = ["/System/Library/Fonts/Apple Color Emoji.ttc"]

        # 使用优先级查找加载字体
        simsun_path = _find_font_with_priority('simsun', simsun_system_paths)
        simhei_path = _find_font_with_priority('simhei', simhei_system_paths)
        arial_path = _find_font_with_priority('arial', arial_system_paths)
        symbol_path = _find_font_with_priority('symbol', symbol_system_paths)

        # 注册字体
        try:
            from reportlab.pdfbase.pdfmetrics import registerFontFamily
            
            # 注册宋体
            if simsun_path:
                pdfmetrics.registerFont(TTFont('SimSun', simsun_path))
                # 注册宋体字体族（让 bold 和 italic 都指向同一个字体）
                registerFontFamily('SimSun', normal='SimSun', bold='SimSun', italic='SimSun', boldItalic='SimSun')
            else:
                raise FileNotFoundError("宋体字体文件未找到")
            
            # 注册黑体
            if simhei_path:
                pdfmetrics.registerFont(TTFont('SimHei', simhei_path))
                # 注册黑体字体族（让 bold 和 italic 都指向同一个字体）
                registerFontFamily('SimHei', normal='SimHei', bold='SimHei', italic='SimHei', boldItalic='SimHei')
            else:
                raise FileNotFoundError("黑体字体文件未找到")
            
            # 注册Arial字体（可选）
            if arial_path:
                pdfmetrics.registerFont(TTFont('Arial', arial_path))
                registerFontFamily('Arial', normal='Arial', bold='Arial', italic='Arial', boldItalic='Arial')
            else:
                logger.warning("Arial字体文件未找到，将使用默认字体")

            # 注册符号字体（可选）
            if symbol_path:
                pdfmetrics.registerFont(TTFont('SymbolFont', symbol_path))
                registerFontFamily('SymbolFont', normal='SymbolFont', bold='SymbolFont', italic='SymbolFont', boldItalic='SymbolFont')
            elif arial_path:
                pdfmetrics.registerFont(TTFont('SymbolFont', arial_path))
                registerFontFamily('SymbolFont', normal='SymbolFont', bold='SymbolFont', italic='SymbolFont', boldItalic='SymbolFont')
                logger.info(f"使用Arial作为符号字体备选")
            else:
                logger.warning("符号字体文件未找到")

            # 注册emoji字体（用于显示📎 🌐等图标）
            emoji_registered = False
            emoji_path = None

            # 尝试从多个路径加载emoji字体
            for sys_path in emoji_system_paths:
                emoji_path = _find_font_with_priority('emoji', [sys_path])
                if emoji_path:
                    try:
                        pdfmetrics.registerFont(TTFont('EmojiFont', emoji_path))
                        emoji_registered = True
                        break
                    except Exception as e:
                        logger.warning(f"注册Emoji字体失败 {emoji_path}: {e}")
            
            # 如果emoji字体注册失败，使用备选方案
            if not emoji_registered:
                if symbol_path:
                    pdfmetrics.registerFont(TTFont('EmojiFont', symbol_path))
                    logger.info(f"使用符号字体作为Emoji备选: {symbol_path}")
                    emoji_registered = True
                elif arial_path:
                    pdfmetrics.registerFont(TTFont('EmojiFont', arial_path))
                    logger.info(f"使用Arial字体作为Emoji备选: {arial_path}")
                    emoji_registered = True
                else:
                    logger.warning("Emoji字体文件未找到，emoji图标可能无法显示")

            # 注册emoji字体族（让bold和italic都指向同一个字体，避免"Can't map determine family/bold/italic"错误）
            if emoji_registered:
                from reportlab.pdfbase.pdfmetrics import registerFontFamily
                registerFontFamily('EmojiFont', normal='EmojiFont', bold='EmojiFont', italic='EmojiFont',
                                   boldItalic='EmojiFont')

        except Exception as e:
            logger.warning(f"警告: 无法注册字体: {e}，将使用默认字体")
            # 使用 ReportLab 内置的中文字体作为后备
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
            pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))

        # 从 Markdown 内容中提取第一个标题作为 PDF 标题元数据
        pdf_title = ''
        for md_line in markdown_content.split('\n'):
            md_line_stripped = md_line.strip()
            if md_line_stripped.startswith('# '):
                pdf_title = md_line_stripped[2:].strip()
                break

        # 如果没有找到 H1 标题，使用文件名（去掉扩展名）作为标题
        if not pdf_title:
            pdf_title = output_path.stem if hasattr(output_path, 'stem') else Path(str(output_path)).stem

        # 创建 PDF 文档
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2.5 * cm,  # 增加底部边距以容纳页码
            title=pdf_title,
            author='PanguAI'
        )

        # 定义页码回调函数
        def add_page_number(canvas, doc):
            """
            在每页底部中央添加页码
            """
            page_num = canvas.getPageNumber()
            text = f"第 {page_num} 页"
            canvas.saveState()
            canvas.setFont('SimSun', 9)
            canvas.drawCentredString(A4[0] / 2, 1.2 * cm, text)
            canvas.restoreState()

        # 定义样式
        styles = getSampleStyleSheet()

        # 正文样式（宋体）
        style_normal = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName='SimSun',
            fontSize=10.5,
            leading=16,  # 行间距 1.5倍
            alignment=TA_LEFT,  # 改为左对齐,避免两端对齐导致的格式问题
            spaceBefore=3,  # 减少段落间距
            spaceAfter=3,
            # 设置加粗文本使用黑体
            bulletFontName='SimHei',
            wordWrap='CJK'  # 使用CJK换行规则
        )

        # 一级标题样式（文章标题：黑体加粗，居中）
        style_h1 = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontName='SimHei',
            fontSize=18,
            leading=27,
            alignment=TA_CENTER,  # 居中
            spaceBefore=12,
            spaceAfter=10,
            textColor=colors.black,
            fontWeight='bold'
        )

        # 二级标题样式（章节标题：黑体加粗，左对齐）
        style_h2 = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontName='SimHei',
            fontSize=14,
            leading=21,
            alignment=TA_LEFT,  # 左对齐
            spaceBefore=10,
            spaceAfter=6,
            textColor=colors.black,
            fontWeight='bold'
        )

        # 三级标题样式（章节内二级标题：黑体加粗，左对齐）
        style_h3 = ParagraphStyle(
            'CustomHeading3',
            parent=styles['Heading3'],
            fontName='SimHei',
            fontSize=12,
            leading=18,
            alignment=TA_LEFT,  # 左对齐
            spaceBefore=8,
            spaceAfter=5,
            textColor=colors.black,
            fontWeight='bold'
        )

        # 四级标题样式（小节标题：黑体加粗，左对齐）
        style_h4 = ParagraphStyle(
            'CustomHeading4',
            parent=styles['Heading3'],
            fontName='SimHei',
            fontSize=11,
            leading=16,
            alignment=TA_LEFT,
            spaceBefore=6,
            spaceAfter=4,
            textColor=colors.black,
            fontWeight='bold'
        )

        # 五级标题样式（小节标题：黑体加粗，左对齐）
        style_h5 = ParagraphStyle(
            'CustomHeading5',
            parent=styles['Heading3'],
            fontName='SimHei',
            fontSize=10.5,
            leading=15,
            alignment=TA_LEFT,
            spaceBefore=5,
            spaceAfter=3,
            textColor=colors.black,
            fontWeight='bold'
        )

        # 六级标题样式（小节标题：黑体加粗，左对齐）
        style_h6 = ParagraphStyle(
            'CustomHeading6',
            parent=styles['Heading3'],
            fontName='SimHei',
            fontSize=10,
            leading=14,
            alignment=TA_LEFT,
            spaceBefore=4,
            spaceAfter=3,
            textColor=colors.black,
            fontWeight='bold'
        )

        # 右对齐样式（用于署名等）
        style_right = ParagraphStyle(
            'CustomRight',
            parent=styles['Normal'],
            fontName='SimSun',
            fontSize=10.5,
            leading=16,
            alignment=TA_RIGHT,
            spaceBefore=3,
            spaceAfter=3
        )

        # 代码块样式
        style_code = ParagraphStyle(
            'CustomCode',
            parent=styles['Code'],
            fontName='SimSun',  # 改用宋体支持中文,虽然不是等宽但能正常显示
            fontSize=9,
            leading=12,
            alignment=TA_LEFT,
            spaceBefore=3,
            spaceAfter=3,
            leftIndent=15,
            rightIndent=15,
            backColor=colors.Color(0.95, 0.95, 0.95),  # 浅灰色背景
            wordWrap='CJK'  # 允许自动换行
        )

        # 报告统计信息样式（浅蓝色背景，带边框）
        style_stats = ParagraphStyle(
            'StatsInfo',
            parent=styles['Normal'],
            fontName='SimSun',
            fontSize=10,
            leading=15,
            alignment=TA_LEFT,
            spaceBefore=8,
            spaceAfter=8,
            leftIndent=12,
            rightIndent=12,
            backColor=colors.Color(0.94, 0.97, 1.0),  # 浅蓝色背景 #F0F8FF
            borderColor=colors.Color(0.7, 0.85, 0.95),  # 蓝色边框
            borderWidth=1,
            borderPadding=8
        )

        # 参考来源样式（改善长URL换行）
        style_reference = ParagraphStyle(
            'Reference',
            parent=styles['Normal'],
            fontName='SimSun',
            fontSize=9.5,
            leading=14,
            alignment=TA_LEFT,
            spaceBefore=4,
            spaceAfter=4,
            leftIndent=0,
            firstLineIndent=0,
            wordWrap='CJK',
            allowWidows=1,
            allowOrphans=1
        )

        # 解析 Markdown 并转换为 ReportLab 元素
        story = []
        lines = markdown_content.split('\n')

        i = 0
        in_code_block = False
        in_math_block = False
        in_stats_section = False  # 标记是否在报告统计信息区块
        in_reference_section = False  # 标记是否在参考来源区块
        code_block_lines = []
        math_block_lines = []
        
        # 跟踪上一个标题的大纲层级，防止层级跳跃（如从1直接跳到3）导致的错误
        last_outline_level = -1

        while i < len(lines):
            line = lines[i].strip()

            # 处理数学公式块 $$
            if line == '$$':
                if not in_math_block:
                    # 开始数学公式块
                    in_math_block = True
                    math_block_lines = []
                else:
                    # 结束数学公式块
                    in_math_block = False
                    if math_block_lines:
                        # 数学公式使用居中、斜体显示
                        math_text = '\n'.join(math_block_lines)

                        # 尝试渲染为图片
                        if MATPLOTLIB_AVAILABLE:
                            img_path = _render_latex_to_image(math_text, fontsize=8)
                            if img_path and img_path.exists():
                                # 使用图片
                                img = RLImage(str(img_path))
                                # 调整图片大小,保持宽高比(缩小尺寸)
                                img._restrictSize(8 * cm, 2 * cm)
                                # 居中显示
                                story.append(Spacer(1, 0.05 * cm))
                                story.append(img)
                                story.append(Spacer(1, 0.05 * cm))
                            else:
                                # 图片渲染失败,使用文本
                                math_text_display = _simplify_latex(math_text)
                                math_style = ParagraphStyle(
                                    'MathFormula',
                                    parent=style_normal,
                                    alignment=TA_CENTER,
                                    fontName='SimSun',
                                    fontSize=10,
                                    leading=14,
                                    spaceBefore=5,
                                    spaceAfter=5
                                )
                                story.append(Paragraph(f'<i>{math_text_display}</i>', math_style))
                        else:
                            # matplotlib不可用,简化LaTeX为可读文本
                            math_text_display = _simplify_latex(math_text)
                            math_style = ParagraphStyle(
                                'MathFormula',
                                parent=style_normal,
                                alignment=TA_CENTER,
                                fontName='SimSun',
                                fontSize=10,
                                leading=14,
                                spaceBefore=5,
                                spaceAfter=5
                            )
                            story.append(Paragraph(f'<i>{math_text_display}</i>', math_style))
                i += 1
                continue

            # 如果在数学公式块中，收集公式行
            if in_math_block:
                math_block_lines.append(lines[i].strip())
                i += 1
                continue

            # 处理代码块
            if line.startswith('```'):
                if not in_code_block:
                    # 开始代码块
                    in_code_block = True
                    code_block_lines = []
                else:
                    # 结束代码块
                    in_code_block = False
                    if code_block_lines:
                        # 使用Preformatted保持代码格式
                        code_text = '\n'.join(code_block_lines)

                        # 检测是否包含中文
                        has_chinese = bool(re.search(r'[\u4e00-\u9fff]', code_text))

                        # 根据是否包含中文选择字体
                        if has_chinese:
                            code_style_dynamic = ParagraphStyle(
                                'CodeWithChinese',
                                parent=style_code,
                                fontName='SimSun'
                            )
                        else:
                            code_style_dynamic = ParagraphStyle(
                                'CodeEnglishOnly',
                                parent=style_code,
                                fontName='Arial'
                            )

                        # 转义特殊字符,避免显示为黑色方框
                        code_text_escaped = code_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                        # 替换换行符、空格和制表符以支持Paragraph并保留格式，同时允许自动换行
                        code_content_html = code_text_escaped.replace('\n', '<br/>').replace(' ', '&nbsp;').replace(
                            '\t', '&nbsp;&nbsp;&nbsp;&nbsp;')

                        font_name = 'SimSun' if has_chinese else 'Arial'
                        try:
                            # 使用Paragraph代替Preformatted以支持自动换行
                            story.append(Paragraph(f'<font name="{font_name}" size="9">{code_content_html}</font>',
                                                   code_style_dynamic))
                        except Exception as e:
                            logger.error(f"渲染代码块失败: {e}")
                            # 后备方案：使用Preformatted（虽然会溢出，但至少能显示）
                            story.append(Preformatted(code_text_escaped, code_style_dynamic))
                i += 1
                continue

            # 如果在代码块中，收集代码行
            if in_code_block:
                code_block_lines.append(lines[i])  # 使用原始行，保留缩进
                i += 1
                continue

            # 跳过空行(不添加额外间距,由段落样式控制)
            if not line:
                i += 1
                continue

            # 处理表格
            # 改进表格检测:必须满足真实表格的特征
            # 1. 包含多个 | 符号(至少3个,形成至少2列)
            # 2. 下一行也有 | 符号
            # 3. 检查是否是分隔行(包含---)
            is_likely_table = (line.count('|') >= 3 and  # 至少3个|,形成2列
                               i + 1 < len(lines) and
                               lines[i + 1].count('|') >= 3)

            # 进一步检查:如果是数学公式(不是表格),需要同时满足:
            # 1. 包含 = 和括号的组合
            # 2. 下一行不是表格分隔行(不包含---)
            next_line = lines[i + 1] if i + 1 < len(lines) else ''
            is_separator_line = bool(re.match(r'^\|[\s\-:|]+\|$', next_line.strip()))

            # 如果下一行是分隔行,说明当前行是表头,即使包含$也是表格
            has_math_expr = ('=' in line and '(' in line and ')' in line and not is_separator_line)

            is_table = is_likely_table and not has_math_expr

            if is_table:
                table_lines = [line]
                i += 1
                # 收集表格的所有行
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1

                # 解析表格
                if len(table_lines) >= 2:
                    # 移除分隔行（通常是第二行，包含 --- ）
                    header_line = table_lines[0]
                    data_lines = [l for l in table_lines[1:] if not re.match(r'^\|[\s\-:|]+\|$', l)]

                    # 解析表头
                    # 注意：不能使用 if cell.strip() 过滤，否则会丢失空单元格
                    # 正确做法是先去除行首尾的 |，然后分割
                    headers = [cell.strip() for cell in header_line.strip().strip('|').split('|')]

                    # 解析数据行
                    table_data = [headers]
                    for data_line in data_lines:
                        # 同样保留空单元格
                        cells = [cell.strip() for cell in data_line.strip().strip('|').split('|')]

                        # 补齐或截断单元格以匹配表头列数
                        if len(cells) < len(headers):
                            cells.extend([''] * (len(headers) - len(cells)))
                        elif len(cells) > len(headers):
                            cells = cells[:len(headers)]

                        table_data.append(cells)

                    # 创建表格
                    if table_data:
                        # 将表格单元格内容包装为Paragraph以支持自动换行
                        wrapped_data = []
                        for row_idx, row in enumerate(table_data):
                            wrapped_row = []
                            for cell in row:
                                # 使用_process_inline_formatting处理单元格内容
                                # 这会处理加粗、斜体、下标、上标等所有格式
                                cell_text = _process_inline_formatting(str(cell))

                                if row_idx == 0:
                                    # 表头使用黑体
                                    p = Paragraph(f'<font name="SimHei"><b>{cell_text}</b></font>', style_normal)
                                else:
                                    # 数据行使用宋体
                                    p = Paragraph(f'<font name="SimSun">{cell_text}</font>', style_normal)
                                wrapped_row.append(p)
                            wrapped_data.append(wrapped_row)

                        # 计算可用宽度（页面宽度 - 左右边距）
                        available_width = A4[0] - 4 * cm  # A4宽度减去左右边距各2cm

                        # 根据列数平均分配列宽
                        num_cols = len(table_data[0]) if table_data else 0
                        if num_cols > 0:
                            col_width = available_width / num_cols
                            col_widths = [col_width] * num_cols
                        else:
                            col_widths = None

                        # 创建表格，设置列宽
                        t = Table(wrapped_data, colWidths=col_widths)
                        t.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.95, 0.95, 0.95)),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white]),
                            ('LEFTPADDING', (0, 0), (-1, -1), 6),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                            ('TOPPADDING', (0, 0), (-1, -1), 6),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ]))
                        story.append(t)
                continue

            # 处理换页符 <div style="page-break-before: always;">
            if 'page-break-before' in line.lower():
                story.append(PageBreak())
                i += 1
                continue

            # 处理横线 <hr>
            if line.startswith('<hr') or line.strip() == '<hr>' or line.strip() == '<hr/>':
                # 添加横线
                story.append(
                    HRFlowable(width="100%", thickness=1, color=colors.Color(0.898, 0.906, 0.922), spaceBefore=5,
                               spaceAfter=5))
                i += 1
                continue

            # 处理包含 text-align: right 的 div 标签
            if '<div style="text-align: right;">' in line or "<div style='text-align: right;'>" in line:
                match = re.search(r'<div[^>]*>\s*(.*?)\s*</div>', line)
                if match:
                    text = _process_inline_formatting(match.group(1).strip())
                    story.append(Paragraph(text, style_right))
                    i += 1
                    continue

            # 处理标题（从最长的开始匹配，避免误匹配）
            target_level = -1
            style = None
            prefix_len = 0
            
            if line.startswith('###### '):
                target_level = 5
                style = style_h6
                prefix_len = 7
            elif line.startswith('##### '):
                target_level = 4
                style = style_h5
                prefix_len = 6
            elif line.startswith('#### '):
                target_level = 3
                style = style_h4
                prefix_len = 5
            elif line.startswith('### '):
                target_level = 2
                style = style_h3
                prefix_len = 4
            elif line.startswith('## '):
                target_level = 1
                style = style_h2
                prefix_len = 3
            elif line.startswith('# '):
                target_level = 0
                style = style_h1
                prefix_len = 2
            
            if target_level != -1:
                raw_text = line[prefix_len:].strip()
                text = _process_inline_formatting(raw_text)
                
                # 计算安全的大纲层级：
                # 1. 如果目标层级 <= 上一个层级（回退或同级），直接使用目标层级
                # 2. 如果目标层级 > 上一个层级（深入），最多只能比上一级大1（防止跳级）
                if target_level <= last_outline_level:
                    # 回退或同级：直接使用目标层级，确保同级标题保持同一书签层级
                    safe_level = target_level
                else:
                    # 深入：最多只能比上一个层级大1
                    safe_level = min(target_level, last_outline_level + 1)
                last_outline_level = safe_level
                
                # 创建书签（目录项），使用纯文本作为书签标题
                clean_title = re.sub(r'<[^>]+>', '', raw_text)  # 移除 HTML 标签
                bookmark_key = f'heading_{len(story)}'
                
                story.append(PDFBookmark(clean_title, safe_level, bookmark_key))
                # 显式指定字体名称，避免 ps2tt 映射错误
                story.append(Paragraph(f'<font name="SimHei"><b>{text}</b></font>', style))
                
                # 检测是否进入"报告统计信息"区块
                if '报告统计信息' in clean_title or 'Report Statistics' in clean_title:
                    in_stats_section = True
                    in_reference_section = False
                # 检测是否进入"参考来源"区块
                elif '参考来源' in clean_title or 'References' in clean_title:
                    in_reference_section = True
                    in_stats_section = False
                # 检测是否离开这些特殊区块（遇到其他标题）
                elif (in_stats_section or in_reference_section) and clean_title not in ['报告统计信息', 'Report Statistics', '参考来源', 'References']:
                    in_stats_section = False
                    in_reference_section = False
                
            elif line.startswith('* ') or line.startswith('- '):
                # 无序列表
                text = line[2:].strip()
                # 处理列表项中的格式
                text = _process_inline_formatting(text)
                
                # 如果在统计信息区块内，使用特殊样式
                if in_stats_section:
                    story.append(Paragraph(f'<font name="SimSun">• {text}</font>', style_stats))
                else:
                    story.append(Paragraph(f'\u2022 {text}', style_normal))
            else:
                # 检测纯文本格式的子标题（如 "2.1 标题" "2.1.1 标题" "3.2 Title"）
                # 匹配模式：数字.数字[.数字...] 空格 标题文字（非空，且不以标点结尾）
                plain_heading_match = re.match(r'^(\d+(?:\.\d+)+)\s+(.+)$', line)
                if plain_heading_match:
                    heading_number = plain_heading_match.group(1)  # 如 "2.1" 或 "2.1.1"
                    heading_text_raw = plain_heading_match.group(2).strip()
                    # 排除误判：如果文字很长（超过80字符）或以句号等结尾，可能不是标题
                    is_likely_heading = (
                        len(heading_text_raw) < 80 and
                        not re.search(r'[。？！.?!,，;；]$', heading_text_raw) and
                        len(heading_text_raw) > 0
                    )
                    if is_likely_heading:
                        # 根据编号层级确定标题级别：2.1 → h3(level 2), 2.1.1 → h4(level 3)
                        dot_count = heading_number.count('.')
                        if dot_count == 1:
                            sub_style = style_h3
                            target_level = 2  # h3 → outline level 2
                        elif dot_count == 2:
                            sub_style = style_h4
                            target_level = 3  # h4 → outline level 3
                        elif dot_count == 3:
                            sub_style = style_h5
                            target_level = 4
                        else:
                            sub_style = style_h6
                            target_level = 5

                        full_title = f"{heading_number} {heading_text_raw}"
                        text = _process_inline_formatting(full_title)

                        # 计算安全的大纲层级（与markdown标题逻辑一致）
                        if target_level <= last_outline_level:
                            safe_level = target_level
                        else:
                            safe_level = min(target_level, last_outline_level + 1)
                        last_outline_level = safe_level

                        # 创建书签
                        clean_title = re.sub(r'<[^>]+>', '', full_title)
                        bookmark_key = f'heading_{len(story)}'
                        story.append(PDFBookmark(clean_title, safe_level, bookmark_key))
                        story.append(Paragraph(f'<font name="SimHei"><b>{text}</b></font>', sub_style))

                        i += 1
                        continue

                # 处理行内格式
                line = _process_inline_formatting(line)
                
                # 检测是否是参考文献条目（以 [数字] 开头）
                if in_reference_section and re.match(r'^\[\d+\]', line):
                    # 参考文献条目使用特殊样式，改善长URL换行
                    story.append(Paragraph(line, style_reference))
                else:
                    # 普通段落
                    story.append(Paragraph(line, style_normal))

            i += 1

        # 生成 PDF（应用页码回调函数）
        try:
            doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
        except ValueError as ve:
            if 'undefined destination target' in str(ve) or 'format not resolved' in str(ve):
                logger.warning(f"PDF内部链接解析失败，将移除内部锚点链接后重试: {ve}")
                # 重建story，移除所有内部href（#开头的链接）保留文本
                cleaned_story = []
                for item in story:
                    if isinstance(item, Paragraph):
                        raw = item.text
                        # 移除指向内部锚点的<a href="#...">标签，保留链接文本
                        cleaned = re.sub(r'<a\s+href="#[^"]*"[^>]*>(.*?)</a>', r'\1', raw, flags=re.IGNORECASE | re.DOTALL)
                        cleaned_story.append(Paragraph(cleaned, item.style))
                    else:
                        cleaned_story.append(item)
                doc2 = SimpleDocTemplate(
                    str(output_path), pagesize=A4,
                    rightMargin=2 * cm, leftMargin=2 * cm,
                    topMargin=2 * cm, bottomMargin=2.5 * cm,
                    title=pdf_title, author='PanguAI'
                )
                doc2.build(cleaned_story, onFirstPage=add_page_number, onLaterPages=add_page_number)
                logger.info("已移除内部链接后成功生成PDF")
            else:
                raise
        return True

    except Exception as e:
        logger.error(f"生成 PDF 失败: {e}")
        import traceback
        traceback.print_exc()
        return False


class MCPTools:
    """Multi Agent System MCP Tools Implementation"""

    def __init__(self, workspace_path: str = None):
        self.config = get_config()
        self.workspace_path = Path(workspace_path) if workspace_path else Path.cwd()
        self.workspace_path.mkdir(exist_ok=True, parents=True)

        # Session context for workspace-aware operations
        self.session_id = None
        self.session_workspace_path = None
        self.full_workspace_path = os.path.realpath(self.workspace_path)
        if not self.full_workspace_path.endswith(os.sep):
            self.full_workspace_path += os.sep

        # 初始化 username，尝试从 workspace 配置文件读取，否则使用默认值
        self.username = self._get_username_from_workspace()
        
        # medRxiv API configuration
        self.BASE_URL = "https://api.biorxiv.org/details/medrxiv"
        self.session = requests.Session()
        self.timeout = 30
        self.max_retries = 3
        
        # Academic sites configuration for targeted search
        self.academic_sites_enabled = True  # Enable by default for research tasks

    def _get_username_from_workspace(self) -> str:
        """
        从 workspace 配置文件中读取用户名
        
        Returns:
            用户名，如果读取失败则返回默认值 "用户"
        """
        try:
            username_file = self.workspace_path / '.username'
            if username_file.exists():
                with open(username_file, 'r', encoding='utf-8') as f:
                    username = f.read().strip()
                    if username:
                        return username
        except Exception as e:
            logger.debug(f"读取用户名失败: {e}")

        return "用户"  # 默认用户名

    def set_session_context(self, session_id: str, session_workspace_path: str):
        """Set session context for workspace-aware operations"""
        self.session_id = session_id
        self.session_workspace_path = Path(session_workspace_path)
        # Update workspace path to session-specific path
        self.workspace_path = self.session_workspace_path
        self.full_workspace_path = os.path.realpath(self.workspace_path)
        if not self.full_workspace_path.endswith(os.sep):
            self.full_workspace_path += os.sep
        # 更新 username
        self.username = self._get_username_from_workspace()
        logger.info(
            f"Set session context - ID: {session_id}, Workspace: {session_workspace_path}, Username: {self.username}")

    def get_session_context(self) -> Dict[str, Any]:
        """Get current session context"""
        return {
            "session_id": self.session_id,
            "session_workspace_path": str(self.session_workspace_path) if self.session_workspace_path else None,
            "workspace_path": str(self.workspace_path)
        }
    
    def _get_academic_sites_list(self, query: str = "") -> List[str]:
        """
        Get list of academic websites for targeted searching.
        Ordered by open access priority (higher success rate for content crawling).
        
        Args:
            query: Optional search query for dynamic site selection
        
        Returns:
            List of academic website domains, prioritized by accessibility and relevance
        """
        return [
            # ===== TIER 1: Preprint Servers & Full Open Access (Crawl Success Rate: 90%+) =====
            "arxiv.org",            # arXiv - Physics, Math, CS preprints (fully open)
            "biorxiv.org",          # bioRxiv - Biology preprints (fully open)
            "plos.org",             # PLOS - Public Library of Science (fully open)
            "frontiersin.org",      # Frontiers journals (fully open access)
            "mdpi.com",             # MDPI journals (fully open access)
            
            # ===== TIER 2: AI/ML Conference Proceedings & Open Repositories (Success Rate: 90%+) =====
            "openreview.net",       # OpenReview - AI/ML conference papers (fully open)
            "papers.nips.cc",       # NeurIPS papers (fully open)
            "semanticscholar.org",  # Semantic Scholar (AI-powered search)
            
            # ===== TIER 3: Top Journals with Partial Open Access (Success Rate: 50-70%) =====
            "nature.com",           # Nature Publishing Group (some open articles)
            "science.org",          # Science/AAAS (some open articles)
            "pnas.org",             # PNAS (some open articles)
            
            # ===== TIER 4: Major Publishers & Discipline-Specific (Success Rate: 30-50%) =====
            "link.springer.com",    # SpringerLink (Springer articles)
            "onlinelibrary.wiley.com",  # Wiley Online Library (AGU, Chemistry, etc.)
            "ieeexplore.ieee.org",  # IEEE Xplore (engineering/CS) - 精确化
            "agu.org",              # American Geophysical Union (earth science)
            
            # ===== Below are NOT used in top-15 site targeting =====
            "pubscholar.cn",        # PubScholar China (Chinese academic search)
            "cell.com",             # Cell Press (some open articles)
            "springer.com",         # Springer Nature (mixed access)
            "academic.oup.com",     # Oxford University Press - 精确化
            "cambridge.org",        # Cambridge University Press
            "tandfonline.com",      # Taylor & Francis Online
            "annualreviews.org",    # Annual Reviews
            "acm.org",              # ACM Digital Library (CS)
            "rsc.org",              # Royal Society of Chemistry
            "asm.org",              # American Society for Microbiology
            "geoscienceworld.org", # GeoScienceWorld
            
            # ===== TIER 6: Medical & Life Sciences (Success Rate: 40-60%) =====
            "thelancet.com",        # The Lancet
            "bmj.com",              # BMJ journals
            "nejm.org",             # New England Journal of Medicine
            "jamanetwork.com",      # JAMA Network
            
            # ===== TIER 7: Scientific Societies (Success Rate: 40-60%) =====
            "aaas.org",             # AAAS (American Association for the Advancement of Science)
            "royalsocietypublishing.org",  # Royal Society Publishing
            "aps.org",              # American Physical Society
            "acs.org",              # American Chemical Society
            
            # ===== TIER 8: Databases & Indexes (Success Rate: 20-40%) =====
            "scopus.com",           # Scopus (Elsevier database)
            "proquest.com",         # ProQuest databases
            "sciencedirect.com"     # ScienceDirect (Elsevier)
        ]
    
    def _get_dynamic_academic_sites(self, query: str, base_count: int = 15) -> List[str]:
        """
        Dynamically adjust academic site priority based on query keywords.
        
        Args:
            query: Search query string
            base_count: Number of sites to return (default 15)
        
        Returns:
            List of academic sites with discipline-specific sites boosted to top
        """
        query_lower = query.lower()
        
        # Get base list
        all_sites = self._get_academic_sites_list()
        
        # Define discipline-specific sites and their keywords (支持中英文)
        # 全面覆盖所有主要学科领域，确保任何查询都能匹配到相关专业网站
        discipline_boosts = {
            # ===== Engineering & Technology (工程技术) =====
            "ieee.org": [
                # 电气电子工程
                "engineering", "electrical", "electronic", "circuit", "signal processing", "robotics", "automation",
                "semiconductor", "microelectronics", "power systems", "control systems", "embedded systems",
                "telecommunications", "wireless", "antenna", "radar", "sensor", "actuator", "mechatronics",
                "工程", "电气", "电子", "电路", "信号处理", "机器人", "自动化", "控制", "半导体", "微电子",
                "电力系统", "控制系统", "嵌入式", "通信", "无线", "天线", "雷达", "传感器", "执行器", "机电"
            ],
            "acm.org": [
                # 计算机科学
                "algorithm", "software", "programming", "computing", "database", "graphics", "computer science",
                "artificial intelligence", "machine learning", "deep learning", "neural network", "data mining",
                "computer vision", "natural language processing", "nlp", "distributed systems", "cloud computing",
                "cybersecurity", "cryptography", "blockchain", "human-computer interaction", "hci",
                "算法", "软件", "编程", "程序", "计算", "数据库", "图形", "计算机", "人工智能", "机器学习",
                "深度学习", "神经网络", "数据挖掘", "计算机视觉", "自然语言处理", "分布式", "云计算",
                "网络安全", "密码学", "区块链", "人机交互"
            ],
            "openreview.net": [
                # AI/ML 会议论文（ICLR, NeurIPS, ICML 等）
                "machine learning", "deep learning", "neural network", "artificial intelligence", "reinforcement learning",
                "transformer", "attention mechanism", "generative model", "gan", "vae", "diffusion model",
                "computer vision", "natural language processing", "nlp", "representation learning", "meta-learning",
                "few-shot learning", "transfer learning", "self-supervised learning", "contrastive learning",
                "graph neural network", "gnn", "optimization", "gradient descent", "backpropagation",
                "机器学习", "深度学习", "神经网络", "人工智能", "强化学习", "生成模型", "对比学习"
            ],
            "papers.nips.cc": [
                # NeurIPS 会议论文（神经信息处理系统）
                "neural", "learning", "optimization", "bayesian", "probabilistic", "inference", "statistical learning",
                "deep learning", "machine learning", "reinforcement learning", "supervised learning", "unsupervised learning",
                "semi-supervised", "active learning", "online learning", "bandit", "kernel method", "svm",
                "neural network", "cnn", "rnn", "lstm", "transformer", "attention", "autoencoder",
                "神经", "学习", "优化", "贝叶斯", "概率", "推理", "统计学习"
            ],
            
            # ===== Earth & Environmental Science (地球与环境科学) =====
            "agu.org": [
                # 地球科学
                "climate", "earth", "geology", "geophysics", "atmosphere", "ocean", "environmental",
                "meteorology", "hydrology", "glaciology", "seismology", "volcanology", "tectonics",
                "climate change", "global warming", "carbon cycle", "water cycle", "ecosystem",
                "remote sensing", "gis", "paleoclimate", "oceanography", "marine science",
                "气候", "地球", "地质", "地球物理", "大气", "海洋", "环境", "生态", "气象", "水文",
                "冰川", "地震", "火山", "构造", "气候变化", "全球变暖", "碳循环", "水循环", "生态系统",
                "遥感", "海洋学", "海洋科学"
            ],
            "geoscienceworld.org": [
                "mineral", "petroleum", "sediment", "paleontology", "stratigraphy", "geochemistry",
                "矿物", "石油", "沉积", "古生物", "地层", "地球化学"
            ],
            
            # ===== Chemistry & Materials Science (化学与材料科学) =====
            "rsc.org": [
                # 化学与材料
                "chemistry", "chemical", "molecule", "synthesis", "catalyst", "polymer", "material", "ceramic", "composite",
                "thermal", "heat resistant", "insulation", "nanomaterial", "nanoparticle", "graphene", "carbon nanotube",
                "electrochemistry", "photochemistry", "spectroscopy", "chromatography", "crystallography",
                "superconductor", "semiconductor material", "battery", "fuel cell", "solar cell", "photovoltaic",
                "coating", "corrosion", "metallurgy", "alloy", "steel", "aluminum", "titanium",
                "化学", "分子", "合成", "催化", "聚合物", "材料", "陶瓷", "复合材料", "高温", "耐热", "隔热", "绝缘",
                "纳米材料", "纳米颗粒", "石墨烯", "碳纳米管", "电化学", "光化学", "光谱", "色谱", "晶体",
                "超导", "半导体材料", "电池", "燃料电池", "太阳能电池", "光伏", "涂层", "腐蚀", "冶金", "合金", "钢", "铝", "钛"
            ],
            "acs.org": [
                "organic chemistry", "inorganic", "analytical chemistry", "biochemistry", "material science",
                "pharmaceutical chemistry", "medicinal chemistry", "polymer chemistry", "surface chemistry",
                "有机化学", "无机", "分析化学", "生物化学", "材料科学", "药物化学", "医药化学", "高分子化学", "表面化学"
            ],
            
            # ===== Medical & Life Sciences (医学与生命科学) =====
            "thelancet.com": [
                # 临床医学
                "clinical", "patient", "disease", "treatment", "diagnosis", "medical", "surgery", "therapy",
                "cancer", "oncology", "cardiology", "neurology", "psychiatry", "pediatrics", "radiology",
                "pathology", "immunology", "infectious disease", "vaccine", "antibody", "inflammation",
                "临床", "患者", "疾病", "治疗", "诊断", "医学", "医疗", "手术", "疗法", "癌症", "肿瘤",
                "心脏病", "神经", "精神", "儿科", "放射", "病理", "免疫", "传染病", "疫苗", "抗体", "炎症"
            ],
            "bmj.com": [
                "health", "medicine", "epidemiology", "public health", "healthcare", "prevention",
                "健康", "医学", "流行病", "公共卫生", "卫生", "医疗保健", "预防"
            ],
            "nejm.org": [
                "therapy", "drug", "pharmaceutical", "clinical trial", "randomized controlled trial",
                "疗法", "药物", "制药", "临床试验", "试验", "随机对照"
            ],
            "cell.com": [
                # 生命科学
                "cell biology", "molecular biology", "genetics", "protein", "gene", "genome", "dna", "rna",
                "stem cell", "crispr", "gene editing", "transcription", "translation", "enzyme", "metabolism",
                "signaling pathway", "apoptosis", "autophagy", "epigenetics", "microbiome", "proteomics",
                "细胞", "分子生物", "遗传", "基因", "蛋白质", "蛋白", "基因组", "突变", "表达", "干细胞",
                "基因编辑", "转录", "翻译", "酶", "代谢", "信号通路", "凋亡", "自噬", "表观遗传", "微生物组", "蛋白质组"
            ],
            
            # ===== Physics & Astronomy (物理与天文) =====
            "aps.org": [
                # 物理
                "physics", "quantum", "particle", "condensed matter", "optics", "photonics", "laser",
                "plasma", "nuclear", "atomic", "molecular", "thermodynamics", "statistical mechanics",
                "relativity", "cosmology", "astrophysics", "gravitational wave", "dark matter", "dark energy",
                "物理", "量子", "粒子", "凝聚态", "光学", "力学", "光子", "激光", "等离子体", "核", "原子",
                "分子", "热力学", "统计力学", "相对论", "宇宙学", "天体物理", "引力波", "暗物质", "暗能量"
            ],
        }
        
        # Calculate boost scores for each site
        boost_scores = {}
        for site, keywords in discipline_boosts.items():
            score = sum(1 for kw in keywords if kw in query_lower)
            if score > 0:
                boost_scores[site] = score
        
        # If no specific discipline detected, return base list
        if not boost_scores:
            return all_sites[:base_count]
        
        # Separate boosted sites from base list
        boosted_sites = sorted(boost_scores.keys(), key=lambda s: boost_scores[s], reverse=True)
        base_sites = [s for s in all_sites if s not in boosted_sites]
        
        # Merge: keep top 5 base sites, insert boosted sites, then fill remaining
        top_base = base_sites[:5]  # Always keep top 5 (arxiv, biorxiv, etc.)
        remaining_base = base_sites[5:]
        
        # Construct final list
        final_sites = top_base + boosted_sites + remaining_base
        
        return final_sites[:base_count]
    
    def _deduplicate_search_results(self, results: List[Dict]) -> List[Dict]:
        """
        Deduplicate search results using multiple strategies (方案三实现)
        
        Deduplication strategies:
        1. URL normalization and deduplication (primary)
        2. Title normalization and hash-based deduplication
        3. DOI-based deduplication (if available)
        
        Args:
            results: List of search result dictionaries
            
        Returns:
            Deduplicated list of results
        """
        from hashlib import md5
        import re
        
        seen_urls = set()
        seen_title_hashes = set()
        seen_dois = set()
        deduplicated = []
        duplicate_count = 0
        
        for result in results:
            # Extract key information
            url = result.get('link', '').strip()
            title = result.get('title', '').strip()
            
            # Try to extract DOI from URL or snippet
            doi = None
            snippet = result.get('snippet', '')
            doi_pattern = r'10\.\d{4,}/[^\s]+'
            doi_match = re.search(doi_pattern, url + ' ' + snippet)
            if doi_match:
                doi = doi_match.group(0).lower()
            
            # Strategy 1: DOI deduplication (highest priority)
            if doi and doi in seen_dois:
                logger.debug(f"[DEDUP] Skipping duplicate DOI: {doi}")
                duplicate_count += 1
                continue
            
            # Strategy 2: URL deduplication
            if url:
                # Normalize URL: remove protocol, www, trailing slash, query parameters
                normalized_url = url.lower()
                normalized_url = re.sub(r'^https?://', '', normalized_url)
                normalized_url = re.sub(r'^www\.', '', normalized_url)
                normalized_url = re.sub(r'[?#].*$', '', normalized_url)  # Remove query and fragment
                normalized_url = normalized_url.rstrip('/')
                
                if normalized_url in seen_urls:
                    logger.debug(f"[DEDUP] Skipping duplicate URL: {url}")
                    duplicate_count += 1
                    continue
                
                seen_urls.add(normalized_url)
            
            # Strategy 3: Title hash deduplication
            if title:
                # Normalize title: lowercase, remove punctuation, remove extra spaces
                normalized_title = re.sub(r'[^\w\s]', '', title.lower())
                normalized_title = ' '.join(normalized_title.split())
                
                if len(normalized_title) > 10:  # Only hash meaningful titles
                    title_hash = md5(normalized_title.encode()).hexdigest()
                    
                    if title_hash in seen_title_hashes:
                        logger.debug(f"[DEDUP] Skipping duplicate title: {title[:50]}...")
                        duplicate_count += 1
                        continue
                    
                    seen_title_hashes.add(title_hash)
            
            # Record DOI
            if doi:
                seen_dois.add(doi)
            
            deduplicated.append(result)
        
        if duplicate_count > 0:
            logger.info(f"[DEDUP] Removed {duplicate_count} duplicates from {len(results)} results, kept {len(deduplicated)} unique items")
        
        return deduplicated
    
    def _validate_workspace_path(self, path: str) -> Path:
        """Validate that a path is within the workspace directory"""
        if os.path.isabs(path):
            raise Exception(f"Path '{path}' is absolute. Only relative paths are allowed.")
        joined_path = os.path.join(self.workspace_path, path)
        full_joined_path = os.path.realpath(joined_path)
        if not full_joined_path.startswith(self.full_workspace_path):
            raise Exception(f"Path '{path}' is outside workspace directory.")
        return Path(full_joined_path)

    def _safe_join(self, path: str) -> Path:
        """Alias for _validate_workspace_path for backward compatibility"""
        return self._validate_workspace_path(path)

    # ================ WEB SEARCH TOOLS ================

    def batch_web_search(
            self,
            queries: List[str],
            max_results_per_query: int = 15,
            max_workers: int = 5,
            academic_sites: bool = True,
            fallback_to_general: bool = True,
            min_results_threshold: int = 5
    ) -> MCPToolResult:
        """
        Batch web search using configurable search provider with concurrent processing.
        
        Supports academic site targeting to prioritize results from professional academic websites
        such as Nature, Science, IEEE, ACM, Springer, etc.
        
        Users need to implement their own search provider. Below is an example available:
        [
            {
                "query": "search query",
                "search_results": [
                    {
                        "title": "Page title",
                        "link": "https://example.com",
                        "snippet": "Description snippet",
                        "date": "Feb 8, 2022",
                    },
                    ...
                ]
            },
            ...
        ]
        
        Args:
            queries: List of search queries
            max_results_per_query: Maximum search results per query
            max_workers: Maximum number of concurrent search requests
            academic_sites: If True, prioritize academic websites (Nature, Science, IEEE, etc.)
            fallback_to_general: If True, automatically fallback to general search when academic results are insufficient
            min_results_threshold: Minimum number of results required before triggering fallback (default: 5)
        """
        try:
            from config.config import get_search_engine_config
            search_config = get_search_engine_config()

            if not search_config:
                return MCPToolResult(
                    success=False,
                    error="Search engine not configured"
                )

            # Ensure we never return more than 15 results per query
            actual_max_results = min(max_results_per_query, 15)

            def search_single_query(query: str) -> Dict[str, Any]:
                """Search a single query"""
                try:
                    search_results = self._generic_search(query, actual_max_results, search_config, academic_sites)

                    if not search_results.success:
                        return {
                            'query': query,
                            'success': False,
                            'error': search_results.error,
                            'results': []
                        }

                    # Process search results
                    search_data = search_results.data
                    search_data["organic"] = search_data["organic"][:actual_max_results]

                    return {
                        'query': query,
                        'success': True,
                        'results': search_data,
                        'timestamp': time.time()
                    }

                except Exception as e:
                    logger.error(f"Error searching query '{query}': {e}")
                    return {
                        'query': query,
                        'success': False,
                        'error': str(e),
                        'results': []
                    }

            # Execute searches concurrently
            all_results = []
            with ThreadPoolExecutor(max_workers=min(max_workers, len(queries))) as executor:
                # Submit all search tasks
                future_to_query = {executor.submit(search_single_query, query): query for query in queries}

                # Collect results as they complete
                for future in as_completed(future_to_query):
                    try:
                        result = future.result()
                        all_results.append(result)
                    except Exception as e:
                        query = future_to_query[future]
                        logger.error(f"Error processing search for '{query}': {e}")
                        all_results.append({
                            'query': query,
                            'success': False,
                            'error': str(e),
                            'results': []
                        })

            # Sort results to maintain original query order
            query_order = {query: i for i, query in enumerate(queries)}
            all_results.sort(key=lambda x: query_order.get(x['query'], float('inf')))

            # Fallback strategy: supplement with general search if results are insufficient
            fallback_queries = []
            fallback_triggered = False
            
            if academic_sites and fallback_to_general:
                for result in all_results:
                    if result.get('success', False):
                        organic_count = len(result.get('results', {}).get('organic', []))
                        if organic_count < min_results_threshold:
                            fallback_queries.append(result['query'])
                            logger.warning(
                                f"[SEARCH_FALLBACK] Query '{result['query'][:50]}...' has only {organic_count} results, "
                                f"triggering general search fallback"
                            )
                
                if fallback_queries:
                    fallback_triggered = True
                    logger.warning(
                        f"[SEARCH_FALLBACK] {len(fallback_queries)} queries need supplemental general search"
                    )
                    
                    # Execute general search for insufficient queries
                    def search_general_query(query: str) -> Dict[str, Any]:
                        try:
                            # Use general search (academic_sites=False)
                            search_results = self._generic_search(query, actual_max_results, search_config, academic_sites=False)
                            if not search_results.success:
                                return {'query': query, 'success': False, 'error': search_results.error, 'results': []}
                            
                            search_data = search_results.data
                            search_data["organic"] = search_data["organic"][:actual_max_results]
                            return {'query': query, 'success': True, 'results': search_data, 'timestamp': time.time()}
                        except Exception as e:
                            logger.error(f"Error in general search for '{query}': {e}")
                            return {'query': query, 'success': False, 'error': str(e), 'results': []}
                    
                    # Execute fallback searches
                    fallback_results = []
                    with ThreadPoolExecutor(max_workers=min(max_workers, len(fallback_queries))) as executor:
                        future_to_query = {executor.submit(search_general_query, query): query for query in fallback_queries}
                        for future in as_completed(future_to_query):
                            try:
                                result = future.result()
                                fallback_results.append(result)
                            except Exception as e:
                                query = future_to_query[future]
                                logger.error(f"Error processing fallback search for '{query}': {e}")
                    
                    # Merge fallback results with original results
                    fallback_map = {r['query']: r for r in fallback_results}
                    for i, result in enumerate(all_results):
                        query = result['query']
                        if query in fallback_map and fallback_map[query].get('success', False):
                            # Merge results: academic results first, then general results
                            academic_organic = result.get('results', {}).get('organic', [])
                            general_organic = fallback_map[query].get('results', {}).get('organic', [])
                            
                            # Deduplicate by URL
                            seen_urls = {item['link'] for item in academic_organic}
                            unique_general = [item for item in general_organic if item['link'] not in seen_urls]
                            
                            # Combine results
                            merged_organic = academic_organic + unique_general[:actual_max_results - len(academic_organic)]
                            all_results[i]['results']['organic'] = merged_organic
                            
                            logger.info(
                                f"[SEARCH_FALLBACK] Merged results for '{query[:50]}...': "
                                f"{len(academic_organic)} academic + {len(unique_general)} general = {len(merged_organic)} total"
                            )
            
            # Apply global deduplication across all queries (改进的方案三实现)
            # 步骤1: 收集所有查询的结果
            all_organic_results = []
            query_indices = []  # 记录每个结果属于哪个查询
            
            for idx, result in enumerate(all_results):
                if result.get('success', False) and 'results' in result:
                    organic = result['results'].get('organic', [])
                    for item in organic:
                        all_organic_results.append(item)
                        query_indices.append(idx)
            
            total_results_before = len(all_organic_results)
            
            # 步骤2: 全局去重（跨查询）
            if all_organic_results:
                deduplicated_results = self._deduplicate_search_results(all_organic_results)
                total_results_after = len(deduplicated_results)
                
                # 步骤3: 将去重后的结果分配回各个查询
                # 创建 URL 到去重结果的映射
                deduplicated_url_map = {item.get('link', ''): item for item in deduplicated_results}
                
                # 按原查询分配去重后的结果
                query_result_map = {i: [] for i in range(len(all_results))}
                seen_urls = set()  # 跟踪已分配的 URL，避免重复分配
                
                for item, query_idx in zip(all_organic_results, query_indices):
                    url = item.get('link', '')
                    # 如果这个 URL 在去重结果中，且还没被分配过
                    if url in deduplicated_url_map and url not in seen_urls:
                        query_result_map[query_idx].append(deduplicated_url_map[url])
                        seen_urls.add(url)
                
                # 更新各查询的结果
                for idx, result in enumerate(all_results):
                    if result.get('success', False) and 'results' in result:
                        result['results']['organic'] = query_result_map[idx]
                
                if total_results_before > total_results_after:
                    logger.info(
                        f"[DEDUP_SUMMARY] Global deduplication across {len(all_results)} queries: "
                        f"{total_results_before} → {total_results_after} results "
                        f"({total_results_before - total_results_after} duplicates removed, "
                        f"{(total_results_before - total_results_after) / total_results_before * 100:.1f}% reduction)"
                    )
            else:
                total_results_after = 0

            return MCPToolResult(
                success=True,
                data=all_results,
                metadata={
                    'total_queries': len(queries),
                    'successful_queries': len([r for r in all_results if r.get('success', False)]),
                    'concurrent_workers': min(max_workers, len(queries)),
                    'fallback_triggered': fallback_triggered,
                    'fallback_queries_count': len(fallback_queries) if fallback_triggered else 0,
                    'deduplication': {
                        'results_before': total_results_before,
                        'results_after': total_results_after,
                        'duplicates_removed': total_results_before - total_results_after
                    }
                }
            )

        except Exception as e:
            logger.error(f"Batch web search failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def _generic_search(self, query: str, max_results: int, config: Dict[str, Any], 
                       academic_sites: bool = True) -> MCPToolResult:
        """
        Generic search function with academic site targeting support.
        
        This function returns results in the standard format wrapped in MCPToolResult:
        
        search_res = {
            "organic": [
                {
                    "title": "Page title",
                    "link": "https://example.com",
                    "snippet": "Description snippet",
                    "date": "Feb 8, 2022"
                }
            ]
        }

        return MCPToolResult(success=True, data=search_res)
        
        Args:
            query: Search query string
            max_results: Maximum number of results to return
            config: Search engine configuration
            academic_sites: If True, prioritize academic websites in search
        
        Notes:
        1. It is recommended to use search engine APIs that comply with relevant safety and regulatory requirements.
        The user assumes full responsibility for any safety issues, legal consequences, or policy violations
        arising from the use of the search engine results.
        
        2. User requests may be indirectly transmitted to the search engine API in the form of search queries.
        It is the user's responsibility to implement appropriate measures to protect personal privacy and
        sensitive information. We assume no liability for privacy-related issues arising from such transmission.
        """
        try:
            # Debug: Log function entry
            logger.warning(f"[SEARCH_DEBUG] _generic_search called with academic_sites={academic_sites}, query={query[:50]}...")
            
            # Get academic sites list (dynamic or static)
            academic_sites_list = self._get_dynamic_academic_sites(query, base_count=15)
            logger.warning(f"[SEARCH_DEBUG] Academic sites list length: {len(academic_sites_list) if academic_sites_list else 0}")
            logger.warning(f"[SEARCH_DEBUG] Top 5 sites for this query: {', '.join(academic_sites_list[:5])}")
            
            # Enhance query with academic site targeting if enabled
            enhanced_query = query
            if academic_sites and academic_sites_list:
                # Add site: operators to prioritize academic sources
                # Use OR logic to search across multiple academic sites
                # Increased from 10 to 15 sites for better coverage (still well under 2048 char limit)
                site_filters = " OR ".join([f"site:{site}" for site in academic_sites_list[:15]])
                enhanced_query = f"({query}) ({site_filters})"
                
                # Enhanced debug logging for verification
                logger.warning(f"[ACADEMIC_SEARCH_DEBUG] ========== Academic Search Enabled ==========")
                logger.warning(f"[ACADEMIC_SEARCH_DEBUG] Original query: {query}")
                logger.warning(f"[ACADEMIC_SEARCH_DEBUG] Enhanced query: {enhanced_query}")
                logger.warning(f"[ACADEMIC_SEARCH_DEBUG] Using {len(academic_sites_list[:15])} academic sites (from {len(academic_sites_list)} total)")
                logger.warning(f"[ACADEMIC_SEARCH_DEBUG] Top 15 sites: {', '.join(academic_sites_list[:15])}")
                logger.warning(f"[ACADEMIC_SEARCH_DEBUG] Query length: {len(enhanced_query)} chars (limit: 2048)")
                logger.warning(f"[ACADEMIC_SEARCH_DEBUG] ================================================")
                logger.info(f"Academic search enabled. Enhanced query with {len(academic_sites_list[:15])} academic site filters.")
            
            url = config['base_url']

            payload = json.dumps({
                "q": enhanced_query,
                "num": max_results
            })

            headers = {
                'X-API-KEY': config["api_keys"][0],
                'Content-Type': 'application/json'
            }

            response = requests.request("POST", url, headers=headers, data=payload, timeout=config["timeout"],
                                        proxies=proxy)
            response.raise_for_status()

            return MCPToolResult(success=True, data=response.json())
        except Exception as e:
            return MCPToolResult(success=False, error=f"Generic search failed: {e}")

    def _extract_google_search_date(self, search_item: Dict[str, Any]) -> Optional[str]:
        """Extract publication date from Google Search result"""
        try:
            # Check pagemap metatags for various date formats
            pagemap = search_item.get('pagemap', {})
            metatags = pagemap.get('metatags', [{}])

            if metatags:
                meta = metatags[0]
                # Common date meta tags
                date_fields = [
                    'article:published_time',
                    'article:modified_time',
                    'date',
                    'pubdate',
                    'published',
                    'datePublished',
                    'dateModified',
                    'dc.date',
                    'dc.date.created',
                    'creation_date'
                ]

                for field in date_fields:
                    if field in meta and meta[field]:
                        return meta[field]

            # Check for news articles with publish date
            newsarticle = pagemap.get('newsarticle', [{}])
            if newsarticle and newsarticle[0].get('datepublished'):
                return newsarticle[0]['datepublished']

            # Check article schema
            article = pagemap.get('article', [{}])
            if article and article[0].get('datepublished'):
                return article[0]['datepublished']

            return None

        except Exception as e:
            logger.info(f"Error extracting date from search result: {e}")
            return None

    @staticmethod
    def _extract_publication_date_from_html(url: str) -> Optional[str]:
        """Extract publication date directly from webpage HTML"""
        try:
            # Fetch HTML content
            response = requests.get(url, timeout=10, headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            })
            response.raise_for_status()

            soup = BeautifulSoup(response.content, 'html.parser')

            # Common meta tags for publication date
            meta_selectors = [
                'meta[property="article:published_time"]',
                'meta[property="article:modified_time"]',
                'meta[name="date"]',
                'meta[name="pubdate"]',
                'meta[name="published"]',
                'meta[name="datePublished"]',
                'meta[name="publication-date"]',
                'meta[property="og:published_time"]',
                'meta[name="DC.date"]',
                'meta[name="DC.date.created"]',
                'meta[itemprop="datePublished"]',
                'meta[itemprop="dateModified"]'
            ]

            for selector in meta_selectors:
                meta_tag = soup.select_one(selector)
                if meta_tag:
                    content = meta_tag.get('content') or meta_tag.get('datetime')
                    if content:
                        try:
                            # Parse and standardize the date
                            parsed_date = dateutil.parser.parse(content)
                            return parsed_date.isoformat()
                        except ValueError:
                            continue

            # Check for time tags with datetime attribute
            time_tags = soup.find_all('time', {'datetime': True})
            for time_tag in time_tags:
                try:
                    parsed_date = dateutil.parser.parse(time_tag['datetime'])
                    return parsed_date.isoformat()
                except ValueError:
                    continue

            # JSON-LD structured data
            json_ld_scripts = soup.find_all('script', {'type': 'application/ld+json'})
            for script in json_ld_scripts:
                try:
                    data = json.loads(script.string)
                    if isinstance(data, dict):
                        date_fields = ['datePublished', 'dateCreated', 'dateModified']
                        for field in date_fields:
                            if field in data:
                                parsed_date = dateutil.parser.parse(data[field])
                                return parsed_date.isoformat()
                    elif isinstance(data, list):
                        for item in data:
                            if isinstance(item, dict):
                                for field in date_fields:
                                    if field in item:
                                        parsed_date = dateutil.parser.parse(item[field])
                                        return parsed_date.isoformat()
                except ValueError:
                    continue

            return None

        except Exception as e:
            logger.debug(f"Error extracting publication date from {url}: {e}")
            return None

    def _content_extractor(self, url: str, max_tokens: int, config: Dict[str, Any]) -> MCPToolResult:
        """Get content using URL content extractor"""
        max_retry_num = 5
        sleep_time = 5
        retry_num = 0
        while True:
            retry_num += 1
            try:
                api_key = random.choice(config.get('api_keys', ['default_key']))
                headers = {
                    'Authorization': f'Bearer {api_key}',
                    'Content-Type': 'application/json'
                }

                # Users need to implement this - placeholder for custom URL crawler
                # raise NotImplementedError(
                #     "URL crawler not implemented. Please implement your own URL crawling logic. "
                #     "The function should extract text content from URLs and return it in a structured format "
                #     "with metadata like title, publication date, and word count."
                # )

                # Example implementation for content extractor (commented out):
                crawler_url = f"{config.get('base_url', 'https://api.content-extractor.com')}/{url}"
                response = requests.get(crawler_url, headers=headers, timeout=config.get('timeout', 30))
                response.raise_for_status()

                content = response.text

                # Truncate if needed
                if max_tokens and len(content.split()) > max_tokens:
                    words = content.split()[:max_tokens]
                    content = ' '.join(words) + '...'

                return MCPToolResult(success=True, data=content)

            except Exception as e:
                if retry_num == max_retry_num:
                    return MCPToolResult(success=False, error=f"Content extractor failed: {e}")
                else:
                    time.sleep(sleep_time)

    def url_crawler(
            self,
            documents: List[Dict],
            max_tokens_per_url: int = 100000,
            include_metadata: bool = True,
            max_workers: int = 10
    ) -> MCPToolResult:
        """
        Extract LLM-friendly content from URLs using configurable crawler service.
        Content is saved to specified file paths.

        Users need to implement their own URL crawler. The return format should include:
        - Extracted text content from the URL
        - Metadata like title, publication date, word count
        - Success/error status for each URL

        Args:
            documents: List of document dictionaries containing:
                - url: Web page URL to extract
                - file_path: Local path to save extracted content
                - title: (Optional) Web page title
                - time: (Optional) Web page publication time
            max_tokens_per_url: Maximum tokens per URL result
            include_metadata: Whether to include metadata about extraction
            max_workers: Maximum number of concurrent extraction requests
        """
        try:
            from config.config import get_url_crawler_config
            crawler_config = get_url_crawler_config()

            if not crawler_config:
                return MCPToolResult(
                    success=False,
                    error="URL crawler not configured"
                )

            def process_single_document(doc: Dict) -> Dict[str, Any]:
                """Process a single document: extract content and save to file"""
                url = doc['url']
                file_path = doc['file_path']
                title = doc.get('title')
                doc_time = doc.get('time')

                result_base = {
                    'url': url,
                    'file_path': file_path,
                    'title': title,
                    'time': doc_time,
                    'success': False,
                    'error': None,
                    'content_length': 0,
                    'word_count': 0,
                    'publication_date': None,
                    'extraction_timestamp': time.time(),
                    'write_success': False
                }

                try:
                    # Extract publication date from the webpage
                    publication_date = self._extract_publication_date_from_html(url)
                    result_base["publication_date"] = publication_date

                    # Extract content using content extractor
                    content_result = self._content_extractor(url, max_tokens_per_url, crawler_config)

                    if not content_result.success:
                        result_base['error'] = content_result.error
                        return result_base

                    content = content_result.data
                    if not content:
                        result_base['error'] = "Extracted content is empty"
                        return result_base

                    # Save content to file
                    write_result = self.file_write(
                        file_path=file_path,
                        content=content,
                        create_dirs=True
                    )

                    if not write_result.success:
                        result_base['error'] = f"File write failed: {write_result.error}"
                        return result_base

                    # Build success result
                    result = {
                        **result_base,
                        'success': True,
                        'content_length': len(content),
                        'word_count': len(content.split()),
                        'publication_date': publication_date,
                        'write_success': True
                    }

                    if include_metadata:
                        result['metadata'] = {
                            'truncated': len(content.split()) >= max_tokens_per_url,
                            'has_publication_date': publication_date is not None,
                            'date_extraction_method': 'html_parsing' if publication_date else None,
                            'file_size': len(content.encode('utf-8'))
                        }

                    return result

                except Exception as e:
                    logger.error(f"Error processing document {url}: {e}")
                    # Try to extract publication date even if processing failed
                    try:
                        publication_date = self._extract_publication_date_from_html(url)
                    except:
                        publication_date = None

                    return {
                        **result_base,
                        'error': str(e),
                        'publication_date': publication_date
                    }

            # Execute processing concurrently
            results = []
            with ThreadPoolExecutor(max_workers=min(max_workers, len(documents))) as executor:
                # Submit all processing tasks
                future_to_doc = {executor.submit(process_single_document, doc): doc for doc in documents}

                # Collect results as they complete
                for future in as_completed(future_to_doc):
                    doc = future_to_doc[future]
                    try:
                        result = future.result()
                        results.append(result)
                    except Exception as e:
                        url = doc['url']
                        logger.error(f"Error processing extraction for '{url}': {e}")

                        # Try to extract publication date even if processing failed
                        try:
                            publication_date = self._extract_publication_date_from_html(url)
                        except:
                            publication_date = None

                        results.append({
                            'url': url,
                            'file_path': doc['file_path'],
                            'title': doc.get('title'),
                            'time': doc.get('time'),
                            'success': False,
                            'error': str(e),
                            'publication_date': publication_date,
                            'extraction_timestamp': time.time(),
                            'write_success': False
                        })

            # Sort results to maintain original document order
            url_order = {doc['url']: i for i, doc in enumerate(documents)}
            results.sort(key=lambda x: url_order.get(x['url'], float('inf')))

            successful_extractions = len([r for r in results if r.get('success', False)])
            successful_writes = len([r for r in results if r.get('write_success', False)])

            return MCPToolResult(
                success=True,
                data=results,
                metadata={
                    'total_documents': len(documents),
                    'successful_extractions': successful_extractions,
                    'successful_writes': successful_writes,
                    'failed_processing': len(documents) - successful_extractions,
                    'concurrent_workers': min(max_workers, len(documents))
                }
            )

        except Exception as e:
            logger.error(f"URL crawler batch processing failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def _extract_original_filename(self, filename: str) -> str:
        """
        从文件名中提取原始文件名（去掉file_id前缀和缓存文件的.txt后缀）

        Args:
            filename: 可能包含file_id前缀的文件名，格式如 'file_id_filename.ext' 或 'filename.ext'

        Returns:
            原始文件名（去掉file_id前缀和缓存文件的.txt后缀）
        """
        result = filename
        if '_' in filename:
            parts = filename.split('_', 1)
            # 如果第一部分是file_id（8位以上十六进制），则使用第二部分
            if len(parts) > 1 and len(parts[0]) >= 8 and re.match(r'^[a-f0-9]{8,}', parts[0].lower()):
                result = parts[1]
        
        # 去掉缓存文件的.txt后缀（如 .doc.txt, .docx.txt, .pdf.txt）
        # 但保留原生.txt文件的扩展名
        if (result.endswith('.doc.txt') or
                result.endswith('.docx.txt') or
                result.endswith('.pdf.txt')):
            result = result[:-4]
        
        return result

    def _extract_title_from_filename(self, filename: str) -> str:
        """
        从文件名中提取标题（去掉file_id前缀，保留文件扩展名如.pdf/.txt/.doc）

        Args:
            filename: 可能包含file_id前缀的文件名

        Returns:
            标题（保留原始文件扩展名，只去掉缓存文件的.txt后缀）
        """
        # 先提取原始文件名（去掉file_id前缀）
        original_filename = self._extract_original_filename(filename)
        # 只去掉缓存文件的.txt后缀（如 .doc.txt, .docx.txt, .pdf.txt）
        # 但保留原生.txt文件的扩展名
        if (original_filename.endswith('.doc.txt') or
                original_filename.endswith('.docx.txt') or
                original_filename.endswith('.pdf.txt')):
            original_filename = original_filename[:-4]
        return original_filename

    def _extract_title_from_file_content(self, file_path: Path) -> tuple:
        """
        从文件内容中提取标题和URL
        
        对于 arXiv 论文，会自动从文件名构建 URL

        Args:
            file_path: 文件路径

        Returns:
            (title, url_source) 元组
        """
        title = "Unknown Title"
        url_source = "Unknown URL"
        
        # 检查是否是 arXiv 文件（通过文件名格式判断）
        filename = file_path.name
        # arXiv paper_id 格式：YYMM.NNNNN[vN].txt (例如：2603.02208v1.txt)
        arxiv_match = re.match(r'^(\d{4}\.\d{5}(?:v\d+)?)\.txt$', filename)
        
        if arxiv_match:
            # 这是 arXiv 文件，从文件名构建 URL
            paper_id = arxiv_match.group(1)
            url_source = f"https://arxiv.org/abs/{paper_id}"
            logger.info(f"识别为 arXiv 文件: {filename}, URL: {url_source}")

        # 检查是否是 PubMed 文件（通过文件名和父目录判断）
        # PubMed 文件格式：{PMID}.txt 或 {PMID}_abstract.txt（纯数字PMID）
        pubmed_match = re.match(r'^(\d{5,10})(?:_abstract)?\.txt$', filename)
        if pubmed_match and ('pubmed' in str(file_path.parent).lower()):
            pmid = pubmed_match.group(1)
            url_source = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
            logger.info(f"识别为 PubMed 文件: {filename}, PMID: {pmid}, URL: {url_source}")

        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                lines = f.readlines()

                # 改进标题提取逻辑
                for i, line in enumerate(lines[:30]):  # 检查更多行
                    line = line.strip()
                    
                    # 跳过 arXiv 元数据行（如：arXiv:1206.3218v1 [math.FA] 14 Jun 2012）
                    if line.startswith('arXiv:') or line.startswith('arxiv:'):
                        continue
                    
                    # 处理"Title: xxx URL Source: yyy"格式（标题和URL在同一行）
                    if line.startswith('Title: ') and 'URL Source:' in line:
                        # 提取Title和URL Source之间的内容
                        title_part = line.split('URL Source:')[0]
                        title = title_part.replace('Title: ', '').strip()
                        # 清理HTML标签
                        title = re.sub(r'<[^>]+>', '', title).strip()
                        if title and len(title) >= 10:
                            logger.info(f"提取到标题 (行{i + 1}, 同行格式): {title[:50]}...")
                            break
                    
                    # 处理markdown标题格式
                    if line.startswith('#'):
                        title = line.strip('# ').strip()[:200]
                        # 清理HTML标签
                        title = re.sub(r'<[^>]+>', '', title).strip()
                        logger.info(f"提取到标题 (行{i + 1}): {title[:50]}...")
                        break
                    
                    # 处理普通标题（不包含http）
                    if line and 10 <= len(line) <= 200:
                        if 'http' not in line and not line.startswith('['):
                            title = line.strip()[:200]
                            # 去掉"Title: "前缀（如果存在）
                            if title.startswith('Title: '):
                                title = title[7:]
                            # 清理HTML标签
                            title = re.sub(r'<[^>]+>', '', title).strip()
                            logger.info(f"提取到标题 (行{i + 1}): {title[:50]}...")
                            break

                # 如果标题仍然未知，尝试从长文本行中提取第一句话作为标题
                # 适用于 PubMed abstract 等内容为单行长文本的文件
                if title == "Unknown Title" and lines:
                    first_content_line = ""
                    for line in lines[:10]:
                        stripped = line.strip()
                        if stripped and not stripped.startswith('URL Source:') and not stripped.startswith('http'):
                            first_content_line = stripped
                            break
                    if first_content_line and len(first_content_line) > 200:
                        # 提取第一句话（以句号、问号、感叹号结尾）
                        sentence_match = re.match(r'^(.{20,200}?[.!?])\s', first_content_line)
                        if sentence_match:
                            title = sentence_match.group(1).strip()
                        else:
                            # 没有明确的句子结束符，截取前200个字符
                            title = first_content_line[:200].strip()
                        title = re.sub(r'<[^>]+>', '', title).strip()
                        logger.info(f"从长文本提取标题: {title[:50]}...")

                # 如果不是 arXiv 文件，从内容中提取 URL
                if url_source == "Unknown URL":
                    # 改进URL提取逻辑：排除中文标点符号，确保URL不包含日期
                    for line in lines[:50]:  # 检查更多行以处理 HTML 文件
                        # 匹配URL，但排除中文标点符号（，。；：！？）、右方括号]、引号和HTML标签
                        url_match = re.search(r'https?://[^\s\]，。；：！？"\'<>]+', line)
                        if url_match:
                            url_source = url_match.group(0)
                            # 清理 URL 末尾可能的 HTML 标签残留
                            url_source = re.sub(r'["\'/]+$', '', url_source)  # 移除末尾的引号、斜杠
                            logger.info(f"提取到URL: {url_source[:50]}...")
                            break
                
                # 特殊处理：如果是 HTML 文件，尝试从 meta 标签或 canonical 链接提取更准确的信息
                if '<html' in ''.join(lines[:10]).lower() or '<!doctype html>' in ''.join(lines[:5]).lower():
                    logger.info(f"检测到 HTML 文件: {file_path.name}")
                    
                    # 尝试从 canonical 链接提取 URL
                    for line in lines[:100]:
                        canonical_match = re.search(r'<link\s+rel="canonical"\s+href="([^"]+)"', line)
                        if canonical_match:
                            url_source = canonical_match.group(1)
                            logger.info(f"从 canonical 链接提取 URL: {url_source}")
                            break
                    
                    # 尝试从 og:url 提取 URL（备用）
                    if url_source == "Unknown URL":
                        for line in lines[:100]:
                            og_url_match = re.search(r'<meta\s+property="og:url"\s+content="([^"]+)"', line)
                            if og_url_match:
                                url_source = og_url_match.group(1)
                                logger.info(f"从 og:url 提取 URL: {url_source}")
                                break
                    
                    # 尝试从 og:title 或 citation_title 提取标题
                    if title == "Unknown Title":
                        for line in lines[:100]:
                            # 优先使用 og:title
                            og_title_match = re.search(r'<meta\s+property="og:title"\s+content="([^"]+)"', line)
                            if og_title_match:
                                title = og_title_match.group(1)
                                title = re.sub(r'<[^>]+>', '', title).strip()  # 清理 HTML 标签
                                logger.info(f"从 og:title 提取标题: {title[:50]}...")
                                break
                            
                            # 备用：使用 citation_title
                            citation_title_match = re.search(r'<meta\s+name="citation_title"\s+content="([^"]+)"', line)
                            if citation_title_match:
                                title = citation_title_match.group(1)
                                title = re.sub(r'<[^>]+>', '', title).strip()  # 清理 HTML 标签
                                logger.info(f"从 citation_title 提取标题: {title[:50]}...")
                                break
        except Exception as e:
            logger.warning(f"警告: 无法读取研究文件 {file_path} - {str(e)}")

        return title, url_source

    def _extract_title_from_research_filename(self, file_path: str) -> str:
        """
        从research文件名中提取标题（备用方案）

        Args:
            file_path: 文件路径

        Returns:
            标题（如果文件名有效）
        """
        filename = os.path.basename(file_path)
        if filename and filename != file_path:
            title_candidate = os.path.splitext(filename)[0]
            # 如果文件名看起来有意义（不是随机字符串），使用它
            if len(title_candidate) > 3 and not re.match(r'^[a-f0-9]{32}', title_candidate):
                logger.info(f"从文件名提取标题: {title_candidate}")
                return title_candidate
        return "Unknown Title"

    # new
    def generate_abstract_and_keywords(self, article_content: str, user_query: str = "") -> Dict[str, str]:
        """
        使用LLM生成文章的标题、摘要和关键词

        Args:
            article_content: 完整的文章内容
            user_query: 用户的原始查询（可选，用于生成更贴切的摘要）

        Returns:
            包含 'title', 'abstract' 和 'keywords' 的字典
        """
        try:
            import requests
            config = get_config()
            model_config = config.get_custom_llm_config()
            # PANGU 模型配置
            PANGU_URL = model_config.get('url') or os.getenv('MODEL_REQUEST_URL', '')
            model_name = model_config.get('model') or os.getenv("MODEL_NAME", "")
            # 语言检测：优先根据user_query判断，其次根据文章内容判断
            # Priority: user_query language > article content language
            is_english_content = False
            
            # 首先检查user_query的语言
            if user_query:
                query_zh_count = len(re.findall(r'[\u4e00-\u9fff]', user_query))
                query_en_count = len(re.findall(r'[a-zA-Z]', user_query))
                # 如果query包含中文，则使用中文
                if query_zh_count > 0:
                    is_english_content = False
                    logger.info(f"Language detection: user_query contains Chinese ({query_zh_count} chars), using Chinese")
                # 如果query纯英文，则使用英文
                elif query_en_count > 10 and query_zh_count == 0:
                    is_english_content = True
                    logger.info(f"Language detection: user_query is English only, using English")
            
            # 如果user_query为空或无法判断，则根据文章内容判断
            if not user_query:
                sample_text = article_content[:5000]
                zh_count = len(re.findall(r'[\u4e00-\u9fff]', sample_text))
                en_count = len(re.findall(r'[a-zA-Z]', sample_text))
                if zh_count < 50 and en_count > 200:
                    is_english_content = True
                elif en_count > 0 and (zh_count / en_count) < 0.05:
                    is_english_content = True

            # 根据内容语言定制 Prompt
            if is_english_content:
                lang_instruction = "The article is in English. You MUST generate the Title, Abstract, and Keywords in English."
                format_instruction = """
                Please strictly follow this format:

                Title:
                [Write Title Here]

                Abstract:
                [Write Abstract Here]

                Keywords:
                [Keyword1; Keyword2; Keyword3]
                """
            else:
                lang_instruction = "文章内容包含中文，请务必用中文生成。"
                format_instruction = """
                请严格按照以下格式输出：

                标题：
                [在这里写标题]

                摘要：
                [在这里写摘要内容]

                关键词：
                [关键词1; 关键词2; 关键词3]
                """

            # 构建生成标题、摘要和关键词的prompt - 改进为更明确的格式
            prompt = f"""请仔细阅读以下文章内容，准确提取或生成标题、摘要和关键词。

                文章内容：
                {article_content}

                重要提示：
                - 你的任务是"提取"和"总结"，而不是"创作"。
                - 摘要必须完全基于文章内容，不得包含文章中未提及的信息。

                具体要求：
                1. 语言要求：{lang_instruction}
                2. 标题：准确反映文章主题，优先提取文章原本的标题。
                3. 摘要（重点）：
                      - 优先提取文章原有的摘要/Abstract部分（如果有）。
                      - 如果没有原有摘要，请仔细阅读全文，生成一个结构化的摘要。
                      - 摘要应隐含以下四个层次（按此顺序组织内容，但禁止出现层次标题）：
                         * 第一层：背景与目的
                         * 第二层：方法与过程  
                         * 第三层：主要结果与发现
                         * 第四层：结论与意义
                      - 明确禁止：禁止使用【】、（）等符号标注章节，禁止出现"背景与目的"、"方法与过程"、"主要结果与发现"、"结论与意义"等层次标题文字。
                      - 摘要长度控制在300字左右，必须具体、准确，拒绝空泛的套话。
                4. 关键词：提取5-8个最能代表文章核心主题的关键词，用分号（;）分隔。

                {format_instruction}
                """

            # 调用 PANGU 模型生成
            headers = {'Content-Type': 'application/json'}

            logger.info("正在调用 PANGU 模型生成标题、摘要和关键词...")
            response = requests.post(
                url=PANGU_URL,
                headers=headers,
                json={
                    "model": model_name,
                    "messages": [
                        {"role": "system", "content": "你是一位严谨的学术分析师。你的核心职责是基于给定的文本内容提取准确的信息。请务必客观、真实，严禁编造原文中不存在的内容。"},
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.3
                },
                verify=False,
                timeout=300
            )

            # 解析响应
            response.raise_for_status()
            result_data = response.json()
            result_text = result_data['choices'][0]['message']['content']

            # 打印原始返回结果用于调试
            logger.info("\n" + "=" * 60)
            logger.info(f"{model_name} 模型原始返回结果：")
            logger.info(result_text)
            logger.info("=" * 60 + "\n")

            # 解析结果 - 使用多种策略
            title = ""
            abstract = ""
            keywords = ""

            # 策略1：尝试标准格式提取（标题/Title：...摘要/Abstract：...关键词/Keywords：...）
            title_match = re.search(r'(?:标题|Title)[：:]\s*(.*?)(?=(?:摘要|Abstract)|(?:关键词|Keywords)|$)', result_text,
                                    re.DOTALL | re.IGNORECASE)
            if title_match:
                title = title_match.group(1).strip()
                logger.info(f"✓ 策略1成功提取标题（{len(title)}字）")

            abstract_match = re.search(r'(?:摘要|Abstract)[：:]\s*(.*?)(?=(?:关键词|Keywords)|$)', result_text,
                                       re.DOTALL | re.IGNORECASE)
            if abstract_match:
                abstract = abstract_match.group(1).strip()
                logger.info(f"✓ 策略1成功提取摘要（{len(abstract)}字）")

            keywords_match = re.search(r'(?:关键词|Keywords)[：:]\s*(.*?)$', result_text, re.DOTALL | re.IGNORECASE)
            if keywords_match:
                keywords = keywords_match.group(1).strip()
                logger.info(f"✓ 策略1成功提取关键词")

            # 策略2：如果策略1失败，尝试更宽松的匹配
            if not title:
                # 查找 "标题" 后面的内容
                title_match2 = re.search(r'(?:标题|title)[：:\s]*(.*?)(?=摘要|abstract|关键词|keywords|$)', result_text,
                                         re.IGNORECASE | re.DOTALL)
                if title_match2:
                    title = title_match2.group(1).strip()
                    logger.info(f"✓ 策略2成功提取标题（{len(title)}字）")

            if not abstract:
                # 查找 "摘要" 后面的内容，直到遇到 "关键词" 或文本结束
                abstract_match2 = re.search(r'(?:摘要|abstract)[：:\s]*(.*?)(?=关键词|keywords|$)', result_text,
                                            re.IGNORECASE | re.DOTALL)
                if abstract_match2:
                    abstract = abstract_match2.group(1).strip()
                    logger.info(f"✓ 策略2成功提取摘要（{len(abstract)}字）")

            if not keywords:
                # 查找 "关键词" 后面的内容
                keywords_match2 = re.search(r'(?:关键词|keywords)[：:\s]*(.*?)$', result_text,
                                            re.IGNORECASE | re.DOTALL)
                if keywords_match2:
                    keywords = keywords_match2.group(1).strip()
                    logger.info(f"✓ 策略2成功提取关键词")

            # 策略3：如果仍然失败，尝试按行分割
            if not title or not abstract or not keywords:
                lines = result_text.split('\n')
                in_title = False
                in_abstract = False
                in_keywords = False
                title_lines = []
                abstract_lines = []
                keywords_lines = []

                for line in lines:
                    line = line.strip()
                    if not line:
                        continue

                    # 检查是否是标题标题
                    if re.match(r'(?:标题|Title)[：:]?', line, re.IGNORECASE):
                        in_title = True
                        in_abstract = False
                        in_keywords = False
                        # 如果标题后面直接有内容，提取它
                        content = re.sub(r'^(?:标题|Title)[：:]?\s*', '', line, flags=re.IGNORECASE)
                        if content:
                            title_lines.append(content)
                        continue

                    # 检查是否是摘要标题
                    if re.match(r'(?:摘要|Abstract)[：:]?', line, re.IGNORECASE):
                        in_title = False
                        in_abstract = True
                        in_keywords = False
                        # 如果标题后面直接有内容，提取它
                        content = re.sub(r'^(?:摘要|Abstract)[：:]?\s*', '', line, flags=re.IGNORECASE)
                        if content:
                            abstract_lines.append(content)
                        continue

                    # 检查是否是关键词标题
                    if re.match(r'(?:关键词|Keywords)[：:]?', line, re.IGNORECASE):
                        in_title = False
                        in_keywords = True
                        in_abstract = False
                        # 如果标题后面直接有内容，提取它
                        content = re.sub(r'^(?:关键词|Keywords)[：:]?\s*', '', line, flags=re.IGNORECASE)
                        if content:
                            keywords_lines.append(content)
                        continue

                    # 收集内容
                    if in_title:
                        title_lines.append(line)
                    elif in_abstract:
                        abstract_lines.append(line)
                    elif in_keywords:
                        keywords_lines.append(line)

                if not title and title_lines:
                    title = ' '.join(title_lines)
                    logger.info(f"✓ 策略3成功提取标题（{len(title)}字）")

                if not abstract and abstract_lines:
                    abstract = ' '.join(abstract_lines)
                    logger.info(f"✓ 策略3成功提取摘要（{len(abstract)}字）")

                if not keywords and keywords_lines:
                    keywords = ' '.join(keywords_lines)
                    logger.info(f"✓ 策略3成功提取关键词")

            # 清理提取的内容
            if title:
                # 移除多余的空白和换行
                title = re.sub(r'\s+', ' ', title).strip()
                # 移除开头的标点符号
                title = re.sub(r'^[：:\-\s]+', '', title)

            if abstract:
                # 移除多余的空白和换行
                abstract = re.sub(r'\s+', ' ', abstract).strip()
                # 移除开头的标点符号
                abstract = re.sub(r'^[：:\-\s]+', '', abstract)

            if keywords:
                # 移除多余的空白和换行
                keywords = re.sub(r'\s+', ' ', keywords).strip()
                # 移除开头的标点符号
                keywords = re.sub(r'^[：:\-\s]+', '', keywords)

            # 验证结果
            if title and abstract and keywords:
                logger.info(f"\n✅ 成功生成标题、摘要（{len(abstract)}字）和关键词")
                logger.info(f"标题: {title}")
                logger.info(f"摘要预览: {abstract[:100]}...")
                logger.info(f"关键词: {keywords}")
            else:
                logger.info(f"\n⚠️ 提取不完整:")
                logger.info(f"  标题: {'✓' if title else '✗'} ({len(title) if title else 0}字)")
                logger.info(f"  摘要: {'✓' if abstract else '✗'} ({len(abstract) if abstract else 0}字)")
                logger.info(f"  关键词: {'✓' if keywords else '✗'}")

                # 如果提取失败，使用备用方案
                if not title:
                    # 尝试从文章内容中提取第一个一级标题
                    title_from_content = re.search(r'^#\s+(.+)$', article_content, re.MULTILINE)
                    if title_from_content:
                        title = title_from_content.group(1).strip()
                        logger.info(f"  使用备用方案：从内容中提取标题")
                    else:
                        title = "研究报告"
                        logger.info(f"  使用备用方案：设置默认标题")

                if not abstract and len(result_text) > 50:
                    abstract = result_text[:300].strip()
                    logger.info(f"  使用备用方案：提取前300字符作为摘要")

                if not keywords:
                    keywords = "未能提取关键词"
                    logger.info(f"  使用备用方案：设置默认关键词")

            return {
                "title": title if title else "研究报告",
                "abstract": abstract if abstract else "未能生成摘要",
                "keywords": keywords if keywords else "未能生成关键词"
            }

        except Exception as e:
            import traceback
            logger.error(f"❌ 警告: 生成标题、摘要和关键词失败 - {str(e)}")
            logger.error(f"错误详情:\n{traceback.format_exc()}")
            return {
                "title": "研究报告",
                "abstract": "摘要生成失败",
                "keywords": "关键词生成失败"
            }

    def file_read_dq(self, file_path: str, encoding: str = 'utf-8') -> MCPToolResult:
        """Read file content"""
        try:
            full_path = self._safe_join(file_path)

            if not full_path.exists():
                return MCPToolResult(
                    success=False,
                    error=f"File does not exist: {file_path}"
                )

            content = full_path.read_text(encoding=encoding)
            if len(content) > 40000:
                content = (
                    "Due to the content being too long, only the first 30,000 and last 10,000 characters are returned.\n"
                    "Below is the returned portion of the file content:\n\n"
                    f"First 30,000 characters:\n\n{content[:30000]}\n\n"
                    f"Last 10,000 characters:\n\n{content[-10000:]}"
                )

            return MCPToolResult(
                success=True,
                data=content,
                metadata={
                    'file_size': len(content),
                    'line_count': len(content.splitlines()),
                    'encoding': encoding
                }
            )

        except Exception as e:
            logger.error(f"File read failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def load_json(self, file_path: str, encoding: str = 'utf-8') -> MCPToolResult:
        """
        Read JSON format file 
        """
        try:
            full_path = self._safe_join(file_path)

            if not full_path.exists():
                return MCPToolResult(
                    success=False,
                    error=f"File does not exist: {file_path}"
                )

            res = []

            with open(full_path, "r", encoding=encoding, errors='ignore') as f:
                for idx, line in enumerate(f):
                    try:
                        ele = json.loads(line.strip())
                        res.append(ele)
                    except Exception as e:
                        logger.warning(f"Failed to process file: {e}")
                        continue

            return MCPToolResult(
                success=True,
                data=res,
                metadata={
                    'line_count': len(res),
                    'encoding': encoding
                }
            )

        except Exception as e:
            logger.error(f"File read failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def extract_author_and_title_for_reference(self, article_content: str, source_info: str = "") -> Dict[str, str]:
        """
        从文章内容中提取作者和标题信息，用于生成参考文献格式

        Args:
            article_content: 文章内容（从标题开始后的500个字符）
            source_info: 来源信息（文件路径、URL等）

        Returns:
            包含 'author', 'title', 'source' 的字典
        """
        try:
            import requests
            model_config = get_config().get_custom_llm_config()
            # PANGU 模型配置
            PANGU_URL = model_config.get('url') or os.getenv('MODEL_REQUEST_URL', '')
            model_name = model_config.get('model') or os.getenv("MODEL_NAME", "")
            # 提取文章开头的500个字符
            content_excerpt = article_content[:500] if len(article_content) > 500 else article_content

            # 构建提取作者和标题的prompt
            prompt = f"""请从以下文章内容中提取作者和标题信息。

重要要求：
1. 只提取作者姓名和文章标题，不要包含其他信息（如期刊名、卷期号、单位、资助信息等）
2. 作者：只提取人名，多个作者用"，"分隔。如果找不到作者，填写"[佚名]"
3. 标题：只提取文章的主标题。如果找不到标题，填写"无题"
4. 来源：只填写"{source_info if source_info else '未提供'}"，不要添加其他内容
5. 切勿虚构任何信息

文章内容片段（前500个字符）：
{content_excerpt}

请严格按照以下格式输出，每项只包含要求的内容：

作者：
[只填写作者姓名，如：张三 或 张三，李四 或 [佚名]]

标题：
[只填写文章标题，如：深度学习研究进展 或 无题]

来源：
{source_info if source_info else "未提供"}
"""

            # 调用 PANGU 模型生成
            headers = {'Content-Type': 'application/json'}

            logger.info("正在调用 PANGU 模型提取作者和标题信息...")
            response = requests.post(
                url=PANGU_URL,
                headers=headers,
                json={
                    "model": model_name,
                    "messages": [
                        {"role": "system",
                         "content": "你是一位专业的文献管理专家，擅长从文章内容中提取作者和标题信息，并整理为规范的参考文献格式。请严格按照用户要求的格式输出。"},
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.3
                },
                verify=False,
                timeout=300
            )

            # 解析响应
            response.raise_for_status()
            result_data = response.json()
            result_text = result_data['choices'][0]['message']['content']

            # 打印原始返回结果用于调试
            logger.info("\n" + "=" * 60)
            logger.info(f"{model_name} 模型原始返回结果（作者和标题提取）：")
            logger.info(result_text)
            logger.info("=" * 60 + "\n")

            # 解析结果
            author = ""
            title = ""
            source = ""

            # 策略1：尝试标准格式提取
            author_match = re.search(r'作者[：:]\s*(.*?)(?=标题|来源|$)', result_text, re.DOTALL)
            if author_match:
                author = author_match.group(1).strip()
                logger.info(f"✓ 成功提取作者: {author}")

            title_match = re.search(r'标题[：:]\s*(.*?)(?=来源|$)', result_text, re.DOTALL)
            if title_match:
                title = title_match.group(1).strip()
                logger.info(f"✓ 成功提取标题: {title}")

            source_match = re.search(r'来源[：:]\s*(.*?)$', result_text, re.DOTALL)
            if source_match:
                source = source_match.group(1).strip()
                logger.info(f"✓ 成功提取来源: {source}")

            # 策略2：如果策略1失败，尝试更宽松的匹配
            if not author:
                author_match2 = re.search(r'(?:作者|author)[：:\s]*(.*?)(?=标题|title|来源|source|$)', result_text,
                                          re.IGNORECASE | re.DOTALL)
                if author_match2:
                    author = author_match2.group(1).strip()
                    logger.info(f"✓ 策略2成功提取作者: {author}")

            if not title:
                title_match2 = re.search(r'(?:标题|title)[：:\s]*(.*?)(?=来源|source|$)', result_text,
                                         re.IGNORECASE | re.DOTALL)
                if title_match2:
                    title = title_match2.group(1).strip()
                    logger.info(f"✓ 策略2成功提取标题: {title}")

            if not source:
                source_match2 = re.search(r'(?:来源|source)[：:\s]*(.*?)$', result_text, re.IGNORECASE | re.DOTALL)
                if source_match2:
                    source = source_match2.group(1).strip()
                    logger.info(f"✓ 策略2成功提取来源: {source}")

            # 清理提取的内容
            if author:
                author = re.sub(r'\s+', ' ', author).strip()
                author = re.sub(r'^[：:\-\s]+', '', author)
                # 移除多余的换行
                author = author.replace('\n', ' ')

            if title:
                title = re.sub(r'\s+', ' ', title).strip()
                title = re.sub(r'^[：:\-\s]+', '', title)
                title = title.replace('\n', ' ')

            if source:
                source = re.sub(r'\s+', ' ', source).strip()
                source = re.sub(r'^[：:\-\s]+', '', source)
                source = source.replace('\n', ' ')

            # 使用默认值（如果提取失败）
            if not author or author == "":
                author = "[佚名]"
                logger.info("  使用默认作者: [佚名]")

            if not title or title == "":
                title = "无题"
                logger.info("  使用默认标题: 无题")

            if not source or source == "":
                source = source_info if source_info else "来源未知"
                logger.info(f"  使用默认来源: {source}")

            logger.info(f"\n✅ 成功提取参考文献信息")
            logger.info(f"作者: {author}")
            logger.info(f"标题: {title}")
            logger.info(f"来源: {source}")

            return {
                "author": author,
                "title": title,
                "source": source
            }

        except Exception as e:
            import traceback
            logger.error(f"❌ 警告: 提取作者和标题信息失败 - {str(e)}")
            logger.error(f"错误详情:\n{traceback.format_exc()}")
            return {
                "author": "[佚名]",
                "title": "无题",
                "source": source_info if source_info else "来源未知"
            }

    # new
    def insert_abstract_and_keywords_to_file(self, file_path: str, title: str = "", abstract: str = "",
                                             keywords: str = "", username: str = "用户"):
        """
        将标题、摘要和关键词插入到文件的最开头（第一章之前）

        Args:
            file_path: 文件路径
            title: 报告标题
            abstract: 摘要内容
            keywords: 关键词内容
            username: 用户名，用于生成信息文本
        """
        try:
            # 读取原始文件内容
            with open(file_path, 'r', encoding='utf-8') as f:
                original_content = f.read()

            # 【新增】检查并移除重复的标题（解决偶发性双标题Bug）
            if title:
                from difflib import SequenceMatcher
                lines = original_content.split('\n')
                first_content_idx = -1

                # 找到第一个非空行
                for i, line in enumerate(lines):
                    if line.strip():
                        first_content_idx = i
                        break

                if first_content_idx != -1:
                    first_line = lines[first_content_idx].strip()
                    # 检查是否是标题格式（# 开头或 **粗体**）
                    heading_match = re.match(r'^(#+\s*|\*\*)(.+)', first_line)
                    if heading_match:
                        # 提取标题文本（去除Markdown标记）
                        existing_title = heading_match.group(2)
                        # 如果是粗体结尾，也要去掉
                        if existing_title.endswith('**'):
                            existing_title = existing_title[:-2]
                        existing_title = existing_title.strip()

                        # 计算相似度
                        similarity = SequenceMatcher(None, existing_title, title).ratio()

                        # 如果相似度高，或者包含关系，则认为是重复标题
                        # 降低包含关系的误判风险：只有当现有标题长度接近新标题时才考虑包含关系
                        is_contained = (title in existing_title or existing_title in title)
                        len_ratio = min(len(title), len(existing_title)) / max(len(title), len(existing_title))

                        if similarity > 0.7 or (is_contained and len_ratio > 0.6):
                            print(f"检测到重复标题，已移除原文件开头的标题: {first_line}")
                            # 移除该行
                            lines.pop(first_content_idx)
                            # 移除紧随其后的空行
                            while first_content_idx < len(lines) and not lines[first_content_idx].strip():
                                lines.pop(first_content_idx)
                            # 更新original_content
                            original_content = '\n'.join(lines)

            # 构建标题、摘要和关键词部分
            header_section = ""

            # 判断语言是否为英文（基于标题和摘要中的中文字符数量）
            # 如果中文字符少于5个，认为是英文内容
            chinese_chars = re.findall(r'[\u4e00-\u9fff]', (title or "") + (abstract or ""))
            is_english = len(chinese_chars) < 5

            abstract_label = "Abstract" if is_english else "摘要"
            keywords_label = "Keywords" if is_english else "关键词"

            # 添加标题（作为一级标题）
            if title:
                header_section += f"# {title}\n\n"

            # 添加摘要（作为二级标题）
            if abstract:
                header_section += f"## {abstract_label}\n\n{abstract}\n\n"

            # 添加关键词（作为二级标题）
            if keywords:
                header_section += f"## {keywords_label}\n\n{keywords}\n\n"

            # 添加生成信息文本
            if is_english:
                footer_text = f'Generated by {username} and SciAssistant'
            else:
                footer_text = f'本文由用户{username}和SciAssistant共同创作生成'

            # 使用 font 标签设置颜色 (#808080 灰色)，div 标签控制对齐
            header_section += f' <div style="text-align: right;"> <font color="#808080">——{footer_text}</font> </div> \n\n'

            # 将标题、摘要和关键词插入到文件开头
            new_content = header_section + original_content

            # 写回文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(new_content)

            logger.info(f"成功将标题、摘要和关键词插入到文件开头: {file_path}")

        except Exception as e:
            logger.error(f"警告: 插入标题、摘要和关键词失败 - {str(e)}")

    def _normalize_heading_levels(self, content: str) -> str:
        """
        规范化Markdown标题层级，确保章节标题为二级标题，章节内标题依次递增

        增强功能：
        - 检测并转换非Markdown格式的章节标题
        - 确保第一个非空行是二级标题（章节标题）
        - 规范化所有子标题层级

        Args:
            content: 章节内容

        Returns:
            规范化后的内容
        """
        lines = content.split('\n')
        normalized_lines = []
        first_heading_found = False
        current_chapter_level = 0
        first_content_line = True

        for line in lines:
            stripped_line = line.strip()

            # 跳过空行，直到找到第一个内容行
            if first_content_line and not stripped_line:
                normalized_lines.append(line)
                continue

            # 检测是否是Markdown标题（以#开头）
            heading_match = re.match(r'^(#+)\s+(.+)$', stripped_line)

            # 检测是否是粗体标题（**xxx**格式，可能用作标题）
            bold_match = re.match(r'^\*\*(.+?)\*\*\s*$', stripped_line)

            if heading_match:
                hash_symbols = heading_match.group(1)
                heading_text = heading_match.group(2)
                current_level = len(hash_symbols)

                heading_text_clean = re.sub(r'^\*\*(.+?)\*\*$', r'\1', heading_text)
                numbered_heading_match = re.match(r'^(\d+(?:\.\d+)+)\s+', heading_text_clean)

                # 如果是第一个标题，将其设为二级标题（章节标题）
                if not first_heading_found:
                    first_heading_found = True
                    first_content_line = False
                    current_chapter_level = current_level
                    if numbered_heading_match:
                        number_depth = numbered_heading_match.group(1).count('.') + 1
                        new_level = min(number_depth + 1, 6)
                        normalized_lines.append(f"{'#' * new_level} {heading_text_clean}")
                    else:
                        normalized_lines.append(f"## {heading_text_clean}")
                else:
                    if numbered_heading_match:
                        number_depth = numbered_heading_match.group(1).count('.') + 1
                        new_level = min(number_depth + 1, 6)
                        normalized_lines.append(f"{'#' * new_level} {heading_text_clean}")
                    else:
                        level_diff = current_level - current_chapter_level
                        new_level = max(2, 2 + level_diff)
                        new_level = min(new_level, 6)
                        normalized_lines.append(f"{'#' * new_level} {heading_text_clean}")
            elif bold_match and not first_heading_found:
                # 如果第一个内容是粗体文本且还没有找到标题，将其转换为二级标题
                first_heading_found = True
                first_content_line = False
                current_chapter_level = 2
                heading_text = bold_match.group(1)
                normalized_lines.append(f"## {heading_text}")
            elif first_content_line and stripped_line and not first_heading_found:
                # 如果第一个非空行不是标题格式，也不是粗体，但看起来像标题（短文本，不以标点结尾）
                # 检查是否像标题：长度适中（<100字符）且不以句号、问号等结尾
                if len(stripped_line) < 100 and not re.search(r'[。？！.?!]$', stripped_line):
                    first_heading_found = True
                    first_content_line = False
                    current_chapter_level = 2
                    # 去除可能的粗体标记
                    heading_text = re.sub(r'^\*\*(.+?)\*\*$', r'\1', stripped_line)
                    normalized_lines.append(f"## {heading_text}")
                else:
                    # 不像标题，保持原样，但标记已经找到第一个内容
                    first_content_line = False
                    normalized_lines.append(line)
            else:
                # 非标题行保持不变
                first_content_line = False
                normalized_lines.append(line)

        return '\n'.join(normalized_lines)

    def merge_reports(self, section_contents, output_file, unique_id=None):
        """
        合并章节文件，生成最终报告

        Args:
            section_contents: 章节内容列表
            output_file: 输出文件路径

        Returns:
            Dict: 包含 abstract 和 keywords 的字典
        """
        # 如果没有提供unique_id，生成一个基于时间戳的唯一ID
        if unique_id is None:
            import time
            unique_id = f"msg-{int(time.time() * 1000)}"
        report_files = []
        for section_content in section_contents:
            # Handle both dict format (expected) and string format (fallback)
            if isinstance(section_content, dict):
                file_path = section_content.get('file_path')
            else:
                # Fallback for direct string paths
                file_path = section_content

            if file_path:
                full_path = self.workspace_path / file_path
                report_files.append(full_path)

        # 提取文件名中的数字索引并排序
        def extract_index(file_path):
            """从文件名中提取数字索引"""
            filename = os.path.basename(file_path)
            match = re.search(r'part_(\d+)\.md', filename)
            if match:
                return int(match.group(1))
            return None  # 不符合格式的文件

        # 创建(索引, 文件路径)元组列表并排序
        indexed_files = []
        for file_path in report_files:
            idx = extract_index(file_path)
            if idx is not None:
                indexed_files.append((idx, file_path))

        # 按索引排序
        indexed_files.sort(key=lambda x: x[0])

        if not indexed_files:
            logger.warning("警告: 未找到符合条件的part_*.md文件")
            return

        # 合并文件 - 关键修改：在写入前规范化标题层级
        try:
            merged_content = ""
            with open(output_file, 'w', encoding='utf-8') as outfile:
                for idx, file_path in indexed_files:
                    filename = os.path.basename(file_path)
                    try:
                        with open(file_path, 'r', encoding='utf-8') as infile:
                            # 读取文件内容
                            file_content = infile.read()

                            # 规范化标题层级
                            normalized_content = self._normalize_heading_levels(file_content)

                            # 写入规范化后的内容
                            outfile.write(normalized_content)
                            outfile.write("\n\n")
                            # 添加到合并内容用于后续处理
                            merged_content += normalized_content + "\n\n"
                            logger.info(f"已合并并规范化标题: {filename}")
                    except Exception as e:
                        logger.warning(f"警告: 无法读取文件 {filename} - {str(e)}")

            logger.info(f"\n合并完成! 结果保存在: {output_file}")
            logger.info(f"共合并了 {len(indexed_files)} 个文件")

            # 读取文件分析数据
            file_analysis_path = self.workspace_path / "doc_analysis" / "file_analysis.jsonl"
            file_analysis_data = {}
            file_num_to_path = {}

            # 【智能过滤】基于information_richness字段判断，而不是关键词匹配
            if file_analysis_path.exists():
                import json
                try:
                    # 【关键修复】使用连续编号而不是原始行号
                    # 因为LLM在生成报告时会重新编号（从1开始连续编号）
                    continuous_num = 0
                    with open(file_analysis_path, 'r', encoding='utf-8') as f:
                        for line_num, line in enumerate(f, 1):
                            try:
                                data = json.loads(line.strip())
                                if 'file_path' in data:
                                    file_path = data['file_path']
                                    doc_time = data.get('doc_time', '')
                                    info_richness = data.get('information_richness', '')
                                    
                                    # 【关键修复】过滤掉Processing failed的文件
                                    if doc_time == "Processing failed":
                                        logger.info(f"跳过处理失败的文件 [原始行号{line_num}]: {file_path}")
                                        continue
                                    
                                    # 【智能过滤】基于information_richness判断
                                    # 检查明确的负面表述：considered scarce, indicating scarcity, lacks substantive content
                                    info_richness_lower = info_richness.lower()
                                    negative_indicators = [
                                        'considered scarce', 'indicating scarcity', 'is scarce',
                                        'lacks substantive content', 'no substantive content',
                                        'very limited information', 'does not provide any substantive'
                                    ]
                                    if info_richness and any(indicator in info_richness_lower for indicator in negative_indicators):
                                        logger.info(f"跳过信息稀缺的文件 [原始行号{line_num}]: {file_path} (richness: {info_richness[:80]})")
                                        continue
                                    
                                    # 有效文件，使用连续编号
                                    continuous_num += 1
                                    file_analysis_data[file_path] = data

                                    # 【关键修复】使用连续编号作为映射，与报告中的引用编号一致
                                    # 报告中的引用编号是连续的 [1, 2, 3, ...]，不会跳过无效文件的编号
                                    file_num_to_path[continuous_num] = file_path
                                    logger.info(f"映射连续编号 {continuous_num} (原始行号{line_num}) 到文件路径: {file_path}")
                            except Exception as e:
                                logger.warning(f"警告: 无法解析分析数据行 {line_num} - {str(e)}")
                    logger.info(f"成功加载 {len(file_num_to_path)} 个序号到文件路径的映射（已过滤无效文件，使用连续编号）")
                except Exception as e:
                    logger.warning(f"警告: 无法读取文件分析数据 - {str(e)}")
            else:
                logger.warning(f"警告: 文件分析数据不存在: {file_analysis_path}")

            # 提取引用序号 - 匹配新格式 [数字]
            citation_pattern = r'\[(\d+)\]'
            citations_found = re.findall(citation_pattern, merged_content)

            # 调试：显示找到的原始引用标记
            logger.info(f"原始引用标记: {citations_found}")

            citation_numbers = list(set([int(num) for num in citations_found]))  # 去重
            citation_numbers.sort()

            logger.info(f"找到 {len(citation_numbers)} 个引用标记: {citation_numbers}")

            # 【关键修复】为所有有效文件生成参考文献（包括未直接引用的文献）
            # 获取所有file_analysis.jsonl中的文件序号
            all_file_numbers = sorted(file_num_to_path.keys())
            logger.info(f"file_analysis.jsonl中共有 {len(all_file_numbers)} 个有效文件")
            
            # 使用所有有效文件的序号
            reference_numbers = all_file_numbers
            logger.info(f"将生成 {len(reference_numbers)} 个参考文献条目（包含所有分析的有效文献）")

            # 构建引用列表 - 分离用户文件和其他文件
            user_file_references = []  # 用户文件引用
            other_references = []  # 其他引用

            # 为每个文件构建参考文献条目
            for num in reference_numbers:
                # 查找对应的文件路径
                file_path = file_num_to_path.get(num)

                if file_path:
                    analysis_data = file_analysis_data.get(file_path, {})
                    doc_time = analysis_data.get('doc_time', 'Unknown')
                    core_content = analysis_data.get('core_content', '')
                    task_relevance = analysis_data.get('task_relevance', '')
                    information_richness = analysis_data.get('information_richness', '')

                    # 【关键修复】跳过条件与 Writer 保持完全一致，避免序号错位
                    # 跳过处理失败的文件
                    if doc_time == "Processing failed":
                        logger.error(f"跳过处理失败的文件 {num}: {file_path}")
                        continue

                    # 跳过内容无效的文件（与 Writer 使用相同的关键词列表）
                    invalid_content_keywords = [
                        '安全验证', 'CAPTCHA', 'captcha', '验证码',
                        '404', '403', 'Forbidden', 'placeholder page', 'error page',
                        'no substantive content', 'lacks substantive content',
                        'does not provide any substantive', '没有实质性的信息',
                        'currently missing', 'content is missing', '内容缺失'
                    ]
                    
                    is_invalid_content = any(kw.lower() in core_content.lower() for kw in invalid_content_keywords)
                    if is_invalid_content:
                        logger.warning(f"跳过内容无效的文件 {num}: {file_path}")
                        continue

                    title = "Unknown Title"
                    url_source = "Unknown URL"

                    # 检查是否是用户上传的文件（路径包含 user_uploads）
                    if 'user_uploads' in file_path or file_path.startswith('./user_uploads/'):
                        # 用户上传的文件：使用文件名作为标题
                        user_file_path = self.workspace_path / file_path if not file_path.startswith(
                            './') else self.workspace_path / file_path[2:]
                        if user_file_path.exists():
                            filename = os.path.basename(user_file_path)
                            # 对于用户上传的文件，使用文件修改时间（上传时间）作为日期
                            try:
                                file_mtime = os.path.getmtime(user_file_path)
                                # 格式化为 "YYYY年MM月" 格式
                                upload_date = datetime.fromtimestamp(file_mtime)
                                doc_time = f"{upload_date.year}年{upload_date.month:02d}月"
                            except Exception as e:
                                # 如果获取文件时间失败，使用当前日期
                                try:
                                    current_date = datetime.now()
                                    doc_time = f"{current_date.year}年{current_date.month:02d}月"
                                except:
                                    doc_time = "Unknown"
                        else:
                            filename = os.path.basename(file_path)
                            # 文件不存在，使用当前日期
                            try:
                                current_date = datetime.now()
                                doc_time = f"{current_date.year}年{current_date.month:02d}月"
                            except:
                                doc_time = "Unknown"

                        # 使用辅助函数提取原始文件名和标题
                        original_filename = self._extract_original_filename(filename)
                        title = self._extract_title_from_filename(filename)

                        # 构建完整的来源信息（包含文件名）
                        full_source_info = f"用户上传文件: {original_filename}"
                        author = "[佚名]"  # 默认作者

                        try:
                            # 优先读取转换后的文本文件
                            txt_file_path = None
                            file_suffix = user_file_path.suffix.lower()
                            if file_suffix == '.pdf':
                                # 检查是否存在对应的.pdf.txt文件
                                txt_file_path = user_file_path.with_suffix('.pdf.txt')
                                if txt_file_path.exists():
                                    logger.info(f"找到PDF转换文本文件: {txt_file_path}")
                                    user_file_path = txt_file_path  # 使用文本文件
                                else:
                                    logger.warning(f"未找到PDF转换文本文件: {txt_file_path}")
                            elif file_suffix in ['.doc', '.docx']:
                                # 检查是否存在对应的.doc.txt或.docx.txt文件
                                txt_file_path = user_file_path.with_suffix(f'{file_suffix}.txt')
                                if txt_file_path.exists():
                                    logger.info(f"找到DOC/DOCX转换文本文件: {txt_file_path}")
                                    user_file_path = txt_file_path  # 使用文本文件
                                else:
                                    logger.warning(f"未找到DOC/DOCX转换文本文件: {txt_file_path}")

                            # 读取文件内容，尝试多种编码
                            file_content = None
                            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin1']

                            for encoding in encodings:
                                try:
                                    with open(user_file_path, 'r', encoding=encoding) as f:
                                        file_content = f.read()
                                    logger.info(f"成功使用 {encoding} 编码读取用户上传文件 {num}")
                                    break
                                except UnicodeDecodeError:
                                    continue

                            if file_content is None:
                                logger.error(f"无法读取用户上传文件 {num}，尝试了所有编码")
                                raise Exception("文件编码问题")

                            # 检查是否仍然是PDF二进制内容
                            if file_content.startswith('%PDF-'):
                                logger.error(f"用户上传文件 {num} 仍然是PDF二进制格式，需要使用转换后的文本文件")
                                raise Exception("PDF文件未转换")

                            # 调用extract_author_and_title_for_reference提取作者和标题
                            ref_info = self.extract_author_and_title_for_reference(
                                article_content=file_content,
                                source_info=full_source_info
                            )

                            # 使用提取的作者和标题
                            author = ref_info.get('author', '[佚名]')
                            extracted_title = ref_info.get('title', title)

                            # 如果提取的标题不是默认值，使用提取的标题
                            if extracted_title and extracted_title != "无题":
                                title = extracted_title

                            logger.info(
                                f"用户上传文件引用 {num}: 作者={author}, 标题={title}, 来源={full_source_info}, 日期={doc_time}")
                        except Exception as e:
                            logger.warning(f"提取作者和标题信息失败，使用默认值: {str(e)}")
                            import traceback
                            logger.error(traceback.format_exc())

                        # URL显示优化：使用完整的来源信息（包含文件名）
                        url_source = full_source_info
                    elif file_path.startswith('./library_refs/') or file_path.startswith('library_refs/'):
                        # 文档库文件：使用文件名作为标题
                        filename = os.path.basename(file_path)

                        # 使用辅助函数提取原始文件名和标题（保留原始扩展名）
                        original_filename = self._extract_original_filename(filename)
                        title = self._extract_title_from_filename(filename)

                        # 构建完整的来源信息（包含文件名）
                        full_source_info = f"用户文档库: {original_filename}"
                        author = "[佚名]"  # 默认作者

                        try:
                            # 构建完整文件路径
                            library_file_path = self.workspace_path / file_path if not file_path.startswith(
                                './') else self.workspace_path / file_path[2:]

                            if library_file_path.exists():
                                # 优先读取转换后的文本文件
                                file_suffix = library_file_path.suffix.lower()
                                if file_suffix == '.pdf':
                                    # 检查是否存在对应的.pdf.txt文件
                                    txt_file_path = library_file_path.with_suffix('.pdf.txt')
                                    if txt_file_path.exists():
                                        logger.info(f"找到PDF转换文本文件: {txt_file_path}")
                                        library_file_path = txt_file_path  # 使用文本文件
                                    else:
                                        logger.warning(f"未找到PDF转换文本文件: {txt_file_path}")
                                elif file_suffix in ['.doc', '.docx']:
                                    # 检查是否存在对应的.doc.txt或.docx.txt文件
                                    txt_file_path = library_file_path.with_suffix(f'{file_suffix}.txt')
                                    if txt_file_path.exists():
                                        logger.info(f"找到DOC/DOCX转换文本文件: {txt_file_path}")
                                        library_file_path = txt_file_path  # 使用文本文件
                                    else:
                                        logger.warning(f"未找到DOC/DOCX转换文本文件: {txt_file_path}")

                                # 读取文件内容，尝试多种编码
                                file_content = None
                                encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin1']

                                for encoding in encodings:
                                    try:
                                        with open(library_file_path, 'r', encoding=encoding) as f:
                                            file_content = f.read()
                                        logger.info(f"成功使用 {encoding} 编码读取文档库文件 {num}")
                                        break
                                    except UnicodeDecodeError:
                                        continue

                                if file_content is None:
                                    logger.error(f"无法读取文档库文件 {num}，尝试了所有编码")
                                    raise Exception("文件编码问题")

                                # 检查是否仍然是PDF二进制内容
                                if file_content.startswith('%PDF-'):
                                    logger.error(f"文档库文件 {num} 仍然是PDF二进制格式，需要使用转换后的文本文件")
                                    raise Exception("PDF文件未转换")

                                # 打印文件内容的前500个字符用于调试
                                logger.info(f"\n{'=' * 60}")
                                logger.info(f"文档库文件 {num} 内容预览（前500个字符）：")
                                logger.info(file_content[:500])
                                logger.info(f"{'=' * 60}\n")

                                # 调用extract_author_and_title_for_reference提取作者和标题
                                ref_info = self.extract_author_and_title_for_reference(
                                    article_content=file_content,
                                    source_info=full_source_info
                                )

                                # 使用提取的作者和标题
                                author = ref_info.get('author', '[佚名]')
                                extracted_title = ref_info.get('title', title)

                                # 如果提取的标题不是默认值，使用提取的标题
                                if extracted_title and extracted_title != "无题":
                                    title = extracted_title

                                logger.info(
                                    f"文档库文件引用 {num}: 作者={author}, 标题={title}, 来源={full_source_info}, 日期={doc_time}")
                            else:
                                logger.warning(f"文档库文件不存在: {library_file_path}，使用默认值")
                        except Exception as e:
                            logger.warning(f"提取作者和标题信息失败，使用默认值: {str(e)}")
                            import traceback
                            logger.error(traceback.format_exc())

                        # URL显示为完整的来源信息（包含文件名）
                        url_source = full_source_info
                    else:
                        # 通用路径处理：适用于所有其他类型的文件（research/, url_crawler_save_files/, arxiv/, 或未来新增的目录）
                        # 检查是否是 arXiv 文件（通过路径或文件名格式判断）
                        arxiv_filename = os.path.basename(file_path)
                        # arXiv paper_id 格式：YYMM.NNNNN[vN] (例如：2603.02208v1)
                        is_arxiv_file = (file_path.startswith('arxiv/') or file_path.startswith('./arxiv/') or
                                        re.match(r'^\d{4}\.\d{5}(v\d+)?\.txt$', arxiv_filename))
                        
                        # 检查是否是 PubMed 文件（通过路径或文件名格式判断）
                        pubmed_filename = os.path.basename(file_path)
                        pubmed_pmid_match = re.match(r'^(\d{5,10})(?:_abstract)?\.txt$', pubmed_filename)
                        is_pubmed_file = (
                            (file_path.startswith('pubmed/') or file_path.startswith('./pubmed/')) and
                            pubmed_pmid_match is not None
                        )

                        if is_arxiv_file:
                            # arXiv 文件特殊处理：从文件名构建 URL
                            # 提取 paper_id（例如：2603.02208v1.txt -> 2603.02208v1）
                            paper_id = arxiv_filename.replace('.txt', '')
                            
                            # 构建 arXiv URL
                            url_source = f"https://arxiv.org/abs/{paper_id}"
                            
                            # 从文件内容提取标题
                            direct_file_path = self.workspace_path / file_path
                            if direct_file_path.exists():
                                title, _ = self._extract_title_from_file_content(direct_file_path)
                                logger.info(f"arXiv 文件: {file_path}, URL: {url_source}")
                            else:
                                title = "Unknown Title"
                                logger.warning(f"arXiv 文件不存在: {direct_file_path}")
                        elif is_pubmed_file:
                            # PubMed 文件特殊处理：从文件名构建 URL
                            pmid = pubmed_pmid_match.group(1)
                            url_source = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
                            
                            # 从文件内容提取标题
                            direct_file_path = self.workspace_path / file_path
                            if direct_file_path.exists():
                                title, extracted_url = self._extract_title_from_file_content(direct_file_path)
                                # 如果文件内容中有更好的URL（如PMC URL），使用它
                                if extracted_url and extracted_url != "Unknown URL" and 'pmc' in extracted_url.lower():
                                    url_source = extracted_url
                                logger.info(f"PubMed 文件: {file_path}, PMID: {pmid}, URL: {url_source}")
                            else:
                                title = "Unknown Title"
                                logger.warning(f"PubMed 文件不存在: {direct_file_path}")
                        else:
                            # 其他文件：先尝试直接路径
                            direct_file_path = self.workspace_path / file_path
                            
                            if direct_file_path.exists():
                                # 文件存在，提取标题和URL
                                title, url_source = self._extract_title_from_file_content(direct_file_path)
                                logger.info(f"成功从文件提取标题: {file_path}")
                            else:
                                # 文件不存在，尝试在research目录查找（向后兼容）
                                research_file_path = self.workspace_path / "research" / os.path.basename(file_path)
                                if research_file_path.exists():
                                    title, url_source = self._extract_title_from_file_content(research_file_path)
                                    logger.info(f"从research目录找到文件: {research_file_path}")
                                else:
                                    # 两个位置都找不到，尝试从文件名提取标题
                                    title = self._extract_title_from_research_filename(file_path)
                                    logger.warning(f"警告: 文件不存在: {direct_file_path} 或 {research_file_path}，使用文件名提取标题")

                    # 【过滤无效引用】跳过 ResearchGate 的反爬虫页面
                    if title == "Just a moment..." or title.strip() == "Just a moment...":
                        logger.warning(f"跳过无效引用 {num}: ResearchGate 反爬虫页面 - {url_source}")
                        continue
                    
                    # 【修复 arXiv PDF 文件名作为标题】
                    # 如果标题是 PDF 文件名格式（如 1404.7828v4.pdf），尝试从文件内容重新提取
                    if title.endswith('.pdf') and re.match(r'^\d{4}\.\d{4,5}(v\d+)?\.pdf$', title):
                        logger.warning(f"检测到 PDF 文件名作为标题: {title}，尝试重新提取")
                        direct_file_path = self.workspace_path / file_path
                        if direct_file_path.exists():
                            # 重新提取标题，跳过 arXiv 元数据行
                            title, _ = self._extract_title_from_file_content(direct_file_path)
                            logger.info(f"重新提取的标题: {title}")
                    
                    # 【统一 arXiv URL 为 /abs/ 格式】
                    # 将 arxiv.org/pdf/ 转换为 arxiv.org/abs/
                    if 'arxiv.org/pdf/' in url_source:
                        # 提取 paper_id（例如：https://arxiv.org/pdf/1404.7828 -> 1404.7828）
                        pdf_match = re.search(r'arxiv\.org/pdf/(\d{4}\.\d{4,5}(?:v\d+)?)', url_source)
                        if pdf_match:
                            paper_id = pdf_match.group(1)
                            url_source = f"https://arxiv.org/abs/{paper_id}"
                            logger.info(f"统一 arXiv URL 为摘要页: {url_source}")
                    
                    # 清理标题中的特殊字符，避免显示问题
                    title_cleaned = title.replace('\u2013', '-').replace('\u2014', '-')  # 替换en-dash和em-dash为普通连字符
                    title_cleaned = title_cleaned.replace('\u201c', '"').replace('\u201d', '"')  # 替换智能引号
                    title_cleaned = title_cleaned.replace('\u2018', "'").replace('\u2019', "'")  # 替换智能单引号
                    title_cleaned = title_cleaned.replace('\xa0', ' ')  # 替换不间断空格
                    title_cleaned = ' '.join(title_cleaned.split())  # 规范化空格

                    # 转义Markdown特殊字符，避免被解析为格式符号
                    # 特别处理破折号：将 "- " 替换为 "—"（em-dash），避免被识别为列表
                    title_cleaned = title_cleaned.replace('- ', '— ')  # 破折号+空格 → em-dash+空格
                    title_cleaned = title_cleaned.replace(' -', ' —')  # 空格+破折号 → 空格+em-dash

                    # 优化时间信息显示
                    doc_time_cleaned = doc_time
                    show_time = True
                    if doc_time in ['Unknown', 'unable to determine the web page time', 'Unknown Time']:
                        show_time = False
                    elif 'unable to determine' in doc_time.lower():
                        show_time = False
                    elif '无法确定具体月份' in doc_time:
                        # 如果包含"无法确定具体月份"，只保留年份
                        # 例如: "2024年无法确定具体月份" -> "2024年"
                        match = re.search(r'(\d{4})年', doc_time)
                        if match:
                            doc_time_cleaned = f"{match.group(1)}年"
                        else:
                            doc_time_cleaned = doc_time.replace('无法确定具体月份', '').strip()

                    # 统一引用格式：根据是否有作者信息决定格式
                    # 如果是用户上传文件或文档库文件且提取了作者信息，使用参考文献格式
                    # 其他文件使用原有格式
                    if ('user_uploads' in file_path or file_path.startswith('./user_uploads/') or
                            file_path.startswith('./library_refs/') or file_path.startswith('library_refs/')):
                        # 用户上传文件和文档库文件：统一使用参考文献格式（带📎图标）
                        file_icon = '<font name="EmojiFont" style="font-style: normal;">📎</font>'
                        # 格式：[序号] 📎 作者. 标题[文献类型]. 来源, 时间
                        if 'author' in locals() and author and author != "[佚名]":
                            # 有作者信息
                            if show_time:
                                reference_entry = f"[{num}] {file_icon} {author}. {title_cleaned}[Z]. {url_source}, {doc_time_cleaned}"
                            else:
                                reference_entry = f"[{num}] {file_icon} {author}. {title_cleaned}[Z]. {url_source}"
                        else:
                            # 无作者信息，使用[佚名]
                            if show_time:
                                reference_entry = f"[{num}] {file_icon} [佚名]. {title_cleaned}[Z]. {url_source}, {doc_time_cleaned}"
                            else:
                                reference_entry = f"[{num}] {file_icon} [佚名]. {title_cleaned}[Z]. {url_source}"
                    else:
                        # 网络文档等其他文件：根据来源类型使用不同图标
                        # 格式：[num] 图标 标题，URL，时间
                        if url_source.startswith('http://') or url_source.startswith('https://'):
                            # 判断是否为学术论文（arXiv/PubMed/medRxiv/bioRxiv）
                            is_academic_paper = (
                                'arxiv.org' in url_source.lower() or 
                                'pubmed' in url_source.lower() or 
                                'ncbi.nlm.nih.gov' in url_source.lower() or
                                'medrxiv.org' in url_source.lower() or
                                'biorxiv.org' in url_source.lower()
                            )
                            
                            # 所有网络引用（包括学术论文和普通网页）都使用简洁格式
                            # 让 PDF 转换阶段的正则表达式统一添加图标（📄 或 🌐）和超链接
                            if show_time:
                                reference_entry = f"[{num}] {title_cleaned}，{url_source}，{doc_time_cleaned}"
                            else:
                                reference_entry = f"[{num}] {title_cleaned}，{url_source}"
                        else:
                            # 非 URL（如"用户上传文件"、"用户文档库"），使用传统格式带📎图标
                            file_icon = '<font name="EmojiFont" style="font-style: normal;">📎</font>'
                            if show_time:
                                reference_entry = f"[{num}] {file_icon} {title_cleaned}，{url_source}，{doc_time_cleaned}"
                            else:
                                reference_entry = f"[{num}] {file_icon} {title_cleaned}，{url_source}"

                    # 根据文件类型分类
                    if 'user_uploads' in file_path or file_path.startswith('./user_uploads/'):
                        user_file_references.append((num, reference_entry))
                        logger.info(f"添加用户文件引用 {num}: {reference_entry}")
                    else:
                        other_references.append((num, reference_entry))
                        logger.info(f"添加其他引用 {num}: {reference_entry}")
                else:
                    # 默认引用条目
                    default_entry = f"[{num}] Unknown Title，Unknown URL，Unknown Time"
                    other_references.append((num, default_entry))
                    logger.warning(f"警告: 找不到引用 {num} 对应的文件")

            # 重新排序：用户文件在前，其他文件在后
            # 先按序号排序用户文件
            user_file_references.sort(key=lambda x: x[0])
            # 再按序号排序其他文件
            other_references.sort(key=lambda x: x[0])

            # 创建原始序号到连续序号的映射，使参考文献序号连续
            # 用户文件在前，其他文件在后
            old_to_new_num = {}
            all_sorted_refs = user_file_references + other_references  # 用户文件优先

            for new_num, (old_num, _) in enumerate(all_sorted_refs, 1):
                old_to_new_num[old_num] = new_num

            logger.info(f"\n序号重新映射: {old_to_new_num}")

            # 替换正文中的引用序号为连续序号，并添加超链接
            # 【关键修复】只用负向前视 (?!</a>) 来避免匹配已在 <a> 标签内的引用
            # 这样可以正确处理连续引用如 [4][5]，同时避免重复替换
            for old_num, new_num in old_to_new_num.items():
                # 只检查 [数字] 后面是否紧跟 </a>，如果是则跳过（已被处理）
                # 这允许 </sup>[5] 中的 [5] 被正确替换
                merged_content = re.sub(
                    rf'\[{old_num}\](?!</a>)',
                    f'<sup><a href="#ref-{unique_id}-{new_num}" style="color: #04B5BB; text-decoration: none;">[{new_num}]</a></sup>',
                    merged_content
                )

            # 【兜底处理】检查并移除未匹配的纯文本引用
            # 查找所有未被转换为上标的 [数字] 引用（即没有对应参考文献的引用）
            remaining_plain_refs = re.findall(r'\[(\d+)\](?!</a>)', merged_content)
            if remaining_plain_refs:
                unique_remaining = sorted(set(int(r) for r in remaining_plain_refs))
                logger.warning(f"发现 {len(unique_remaining)} 个未匹配的引用，将被移除: {unique_remaining}")
                
                # 直接移除这些无效引用，只保留实际存在于参考文献列表中的引用
                for ref_num in unique_remaining:
                    merged_content = re.sub(
                        rf'\[{ref_num}\](?!</a>)',
                        '',
                        merged_content
                    )
                logger.info(f"已移除 {len(unique_remaining)} 个无效引用")

            # 重新写入更新后的正文
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(merged_content)
            logger.info(f"已将正文中的引用序号重新编号为连续序号")

            # 重新构建参考文献列表，使用连续序号，并添加锚点ID
            renumbered_references = []
            for new_num, (old_num, old_entry) in enumerate(all_sorted_refs, 1):
                # 替换引用条目中的序号: [OLD] -> <a id="ref-NEW"></a>[NEW]
                # 添加HTML锚点，使正文引用可以跳转到此处
                new_entry = re.sub(
                    rf'^\[{old_num}\]',
                    f'<a id="ref-{unique_id}-{new_num}"></a>[{new_num}]',
                    old_entry
                )
                renumbered_references.append(new_entry)

            # 使用重新编号的参考文献列表
            references = renumbered_references

            if user_file_references:
                logger.info(f"已优先排序：{len(user_file_references)} 个用户文件引用在前，{len(other_references)} 个其他引用在后")

            # ========== 新增：生成并插入标题、摘要和关键词 ==========
            logger.info("\n开始生成标题、摘要和关键词...")
            abstract_keywords_result = {"title": "", "abstract": "", "keywords": ""}
            try:
                # 读取完整的合并后文章内容
                with open(output_file, 'r', encoding='utf-8') as f:
                    complete_article = f.read()

                # 生成标题、摘要和关键词
                abstract_keywords_result = self.generate_abstract_and_keywords(complete_article)

                # 将标题、摘要和关键词插入到文件开头
                # 使用self.username（从MCPTools初始化时传入）
                self.insert_abstract_and_keywords_to_file(
                    output_file,
                    abstract_keywords_result.get("title", ""),
                    abstract_keywords_result.get("abstract", ""),
                    abstract_keywords_result.get("keywords", ""),
                    username=self.username
                )

                logger.info("标题、摘要和关键词已成功添加到文章开头！")
            except Exception as e:
                logger.warning(f"警告: 生成标题、摘要和关键词过程中出错 - {str(e)}")
                abstract_keywords_result = {
                    "title": "研究报告",
                    "abstract": "生成过程出错",
                    "keywords": "生成过程出错"
                }
            # ========== 标题、摘要和关键词生成完成 ==========

            # 确定参考文献标题语言
            ref_title = "参考来源"
            try:
                # 尝试检测文件内容语言
                if os.path.exists(output_file):
                    with open(output_file, 'r', encoding='utf-8') as f:
                        # 只读前5000个字符足够检测语言
                        sample_text = f.read(5000)

                    zh_count = len(re.findall(r'[\u4e00-\u9fff]', sample_text))
                    en_count = len(re.findall(r'[a-zA-Z]', sample_text))

                    # 判定逻辑：如果中文字符极少（<5%的英文字符量 或 绝对数量很少），则认为是英文
                    if zh_count < 10 or (en_count > 0 and zh_count / en_count < 0.05):
                        ref_title = "References"
            except Exception as e:
                logger.warning(f"语言检测失败，使用默认标题: {e}")

            # 将引用列表添加到报告末尾 - 确保即使没有找到引用也添加参考来源部分
            # 添加换页符、标题和分隔横线（间距更紧凑，线条更浅）
            references_section = f"\n\n<div style=\"page-break-before: always;\"></div>\n\n## {ref_title}\n<hr style=\"border: none; border-top: 1px solid #E5E7EB; margin: 0.5em 0;\" />\n\n"
            if references:
                references_section += "\n\n".join(references)
                with open(output_file, 'a', encoding='utf-8') as outfile:
                    outfile.write(references_section)
                logger.info(f"已添加 {len(references)} 个参考来源到报告末尾")
            else:
                # 即使没有找到引用，也要添加参考来源部分
                with open(output_file, 'a', encoding='utf-8') as outfile:
                    outfile.write(references_section)
                logger.info("已添加空的参考来源部分到报告末尾")

            # 生成PDF (确保文件已完全写入并关闭)
            try:
                with open(output_file, 'r', encoding='utf-8') as mdfile:
                    md_content = mdfile.read()

                # 使用 ReportLab 生成 PDF（黑体标题，宋体正文）
                pdf_path = self.workspace_path / "final_report.pdf"
                success = generate_pdf_with_reportlab(md_content, pdf_path)

                if success:
                    logger.info(f"保存PDF文件: {pdf_path}")

                else:
                    logger.error(f"警告: PDF 生成失败")

            except Exception as e:
                logger.warning(f"警告: 无法生成PDF文件 - {str(e)}")
                import traceback
                traceback.print_exc()

            # 返回摘要和关键词信息
            return abstract_keywords_result

        except Exception as e:
            logger.error(f"错误: 写入输出文件失败 - {str(e)}")
            # 清理可能创建的不完整输出文件
            if os.path.exists(output_file):
                try:
                    os.remove(output_file)
                except:
                    pass
            raise

    # 实现concat_section_files的具体内容
    def concat_section_files(
            self,
            section_files: List[Dict],
            final_file_path: str
    ) -> MCPToolResult:
        """
        Concatenate the content of the saved section files into a single file
        """
        try:
            logger.info(f"我现在开始调用concat_section_files了：{section_files}, {final_file_path}")

            # Ensure output directory exists
            import os
            output_dir = os.path.dirname(final_file_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)

            # Convert relative path to absolute if needed
            if not os.path.isabs(final_file_path):
                final_file_path = self.workspace_path / final_file_path

            # 读取section_files中的文件内容，并获取摘要和关键词
            # 生成基于时间戳的唯一ID，避免多轮对话中的锚点冲突
            import time
            unique_id = f"msg-{int(time.time() * 1000)}"
            abstract_keywords = self.merge_reports(section_files, final_file_path, unique_id)

            # 如果没有返回摘要和关键词，设置默认值
            if abstract_keywords is None:
                abstract_keywords = {"abstract": "", "keywords": ""}

            return MCPToolResult(
                success=True,
                data={
                    "merged_files": len(section_files),
                    "output_path": str(final_file_path),
                    "abstract": abstract_keywords.get("abstract", ""),
                    "keywords": abstract_keywords.get("keywords", "")
                },
                metadata={
                    'final_file_path': str(final_file_path),
                    'section_count': len(section_files),
                    'abstract': abstract_keywords.get("abstract", ""),
                    'keywords': abstract_keywords.get("keywords", "")
                }
            )
        except Exception as e:
            logger.error(f"Concatenate section files failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def _validate_file_allocation(
            self,
            classification_result: str,
            user_file_count: int,
            research_file_count: int,
            has_user_files: bool
    ) -> Dict:
        """
        验证LLM的文件分配结果是否符合要求

        Args:
            classification_result: LLM返回的分类结果
            user_file_count: 用户文件数量
            research_file_count: research文件数量
            has_user_files: 是否有用户文件

        Returns:
            验证结果字典，包含valid, message, expected, actual字段
        """
        try:
            # 统计分配的文件数量
            user_files_assigned = classification_result.count('user_uploads')
            research_files_assigned = classification_result.count('research/')

            # 计算预期值
            if has_user_files and research_file_count > 0:
                # 混合场景
                expected_user = user_file_count  # 100%
                if user_file_count >= 8:
                    expected_research_pct = 70
                elif user_file_count >= 5:
                    expected_research_pct = 80
                else:
                    expected_research_pct = 90
                expected_research_min = int(research_file_count * expected_research_pct / 100)

                # 验证
                user_valid = user_files_assigned >= expected_user * 0.8  # 允许80%容差
                research_valid = research_files_assigned >= expected_research_min * 0.5  # 允许50%容差

                if not user_valid or not research_valid:
                    return {
                        "valid": False,
                        "message": f"Mixed scenario: User files {user_files_assigned}/{expected_user}, Research files {research_files_assigned}/{expected_research_min}",
                        "expected": f"User: {expected_user} (100%), Research: {expected_research_min} ({expected_research_pct}%)",
                        "actual": f"User: {user_files_assigned}, Research: {research_files_assigned}"
                    }

            elif research_file_count > 0:
                # 只有research文件场景
                if research_file_count >= 30:
                    expected_research_pct = 85
                elif research_file_count >= 20:
                    expected_research_pct = 90
                elif research_file_count >= 10:
                    expected_research_pct = 95
                else:
                    expected_research_pct = 100

                expected_research_min = int(research_file_count * expected_research_pct / 100)
                research_valid = research_files_assigned >= expected_research_min * 0.7  # 允许70%容差

                if not research_valid:
                    return {
                        "valid": False,
                        "message": f"Research-only scenario: {research_files_assigned}/{expected_research_min} files assigned",
                        "expected": f"Research: {expected_research_min} ({expected_research_pct}%)",
                        "actual": f"Research: {research_files_assigned}"
                    }

            return {
                "valid": True,
                "message": "File allocation meets requirements",
                "expected": "N/A",
                "actual": f"User: {user_files_assigned}, Research: {research_files_assigned}"
            }

        except Exception as e:
            logger.error(f"Validation error: {e}")
            return {
                "valid": True,  # 验证失败时不阻止流程
                "message": f"Validation error: {e}",
                "expected": "N/A",
                "actual": "N/A"
            }

    def search_result_classifier(
            self,
            outline: str,
            key_files: List[Dict],
            model: str = "gpt-4o",
            temperature: float = 0.3,
            max_tokens: int = 4000
    ) -> MCPToolResult:
        """
        Classify and organize search result files according to a structured outline for comprehensive long-form content generation.

        Args:
            outline: Structured outline defining the sections and subsections for organizing the long-form content
            key_files: List of key files to classify
            model: AI model to use for classification and organization
            temperature: Creativity level for the AI classification (0-1)
            max_tokens: Maximum tokens for the AI response
        """
        try:
            logger.info(f"我现在开始调用search_result_classifier了：{outline}, {key_files}")
            # 处理输入的key_files - 使用四个分析维度
            # 获取本地的文件进行分析

            import os
            import json
            def load_json(file_path):
                res = []
                error_lines = 0
                with open(file_path, 'rb') as file:
                    for line in file:
                        try:
                            line.decode('utf-8')
                        except UnicodeDecodeError:
                            error_lines = error_lines + 1
                with open(file_path, "r", encoding="utf-8", errors='ignore') as f:
                    for idx, line in enumerate(f):
                        try:
                            ele = json.loads(line.strip())
                            res.append(ele)
                        except Exception as e:
                            logger.info(e)
                            continue
                return res

            key_files_dict = {}
            # Create full path relative to workspace
            full_analysis_path = self.workspace_path / "doc_analysis/file_analysis.jsonl"
            file_analysis_list = load_json(full_analysis_path)

            for file_info in file_analysis_list:
                if file_info.get('file_path'):
                    key_files_dict[file_info.get('file_path')] = file_info

            # 组装key_files
            prompt_files = ""
            if key_files:
                prompt_files += f"Key Files with Multi-Dimensional Analysis:\n"
                for i, file_info in enumerate(key_files, 1):
                    if file_info.get('file_path') in key_files_dict:
                        file_info = key_files_dict[file_info.get('file_path')]
                        prompt_files += f"\n{i}. File: {file_info.get('file_path', 'Unknown')}\n"
                        prompt_files += f"   Document Time: {file_info.get('doc_time', 'Not specified')}\n"
                        prompt_files += f"   Source Authority: {file_info.get('source_authority', 'Not specified')}\n"
                        prompt_files += f"   Core Content: {file_info.get('core_content', 'Not specified')}\n"
                        prompt_files += f"   Task Relevance: {file_info.get('task_relevance', 'Not specified')}\n"
                        prompt_files += f"   Information Richness: {file_info.get('information_richness', 'Not specified')}\n"
                        prompt_files += "\n"

                prompt_files += "\n"

            # 检查是否有用户上传的文件和research文件
            has_user_files = any('user_uploads' in f.get('file_path', '') for f in key_files)
            user_file_count = sum(1 for f in key_files if 'user_uploads' in f.get('file_path', ''))
            research_file_count = sum(1 for f in key_files if 'research' in f.get('file_path', ''))

            user_file_priority_note = ""
            if has_user_files and research_file_count > 0:
                # 动态计算每章节的文件分配目标
                # 策略：在保证用户文件优先的前提下，最大化总引用文献数量

                # 用户文件目标：强制使用所有文件，确保100%覆盖
                if user_file_count >= 8:
                    user_target_per_section = "ALL user files distributed across sections"
                    user_min_per_section = min(6, user_file_count)  # 提高最低要求
                    user_coverage_goal = f"Use ALL {user_file_count} user files (100% MANDATORY)"
                elif user_file_count >= 5:
                    user_target_per_section = f"ALL {user_file_count} user files"
                    user_min_per_section = min(5, user_file_count)  # 提高最少要求
                    user_coverage_goal = f"Use ALL {user_file_count} user files (100% MANDATORY)"
                elif user_file_count >= 3:
                    user_target_per_section = f"ALL {user_file_count} user files (distribute across sections)"
                    user_min_per_section = min(3, user_file_count)
                    user_coverage_goal = f"Use ALL {user_file_count} files (100% MANDATORY), assign to MULTIPLE sections"
                else:  # 1-2个用户文件
                    user_target_per_section = f"ALL {user_file_count} user file(s)"
                    user_min_per_section = user_file_count
                    user_coverage_goal = f"Use ALL {user_file_count} file(s) (100% MANDATORY), assign to MULTIPLE sections"

                # Research文件目标：最大化使用，确保更多文献被引用
                # 根据用户文件数量动态调整research文件的使用比例
                if user_file_count >= 8:
                    # 用户文件充足时，适度使用research文件
                    research_min_per_section = min(4, research_file_count)  # 提高最少要求从3到4
                    research_target_per_section = min(6, max(5, research_file_count // 3))  # 提高目标
                    research_coverage_pct = 80  # 提高覆盖率从70%到80%
                elif user_file_count >= 5:
                    # 用户文件中等时，增加research文件使用
                    research_min_per_section = min(4, research_file_count)  # 提高最少要求
                    research_target_per_section = min(7, max(5, research_file_count // 2))  # 提高目标
                    research_coverage_pct = 85  # 85%覆盖率
                else:
                    # 用户文件少时，大量使用research文件
                    research_min_per_section = min(5, research_file_count)  # 最少5个
                    research_target_per_section = min(8, max(6, research_file_count // 2))  # 提高目标
                    research_coverage_pct = 90  # 90%覆盖率

                user_file_priority_note = f"""
⚠️ **MIXED SCENARIO: USER FILES + RESEARCH FILES - MAXIMIZE TOTAL CITATIONS** ⚠️
**STRICT COMPLIANCE REQUIRED - BOTH FILE TYPES MUST BE USED**

📊 **INITIAL ASSESSMENT**:
- User-uploaded files: {user_file_count} files (HIGHEST PRIORITY)
- Research files: {research_file_count} files (MUST ALSO BE USED)
- **MANDATORY USER FILES**: ALL {user_file_count} user files MUST be used (100%)
- **MANDATORY RESEARCH FILES**: At least {int(research_file_count * research_coverage_pct / 100)} research files MUST be used ({research_coverage_pct}%)
- **MANDATORY PER SECTION**: {user_min_per_section}+ user files AND {research_min_per_section}+ research files

🎯 **STRICT ALLOCATION RULES (NON-NEGOTIABLE)**:

1. **USER FILES (HIGHEST PRIORITY)**:
   - Character quota: 15K per file (larger than research files)
   - Target per section: {user_target_per_section}
   - **MANDATORY**: Use ALL {user_file_count} user files (100% coverage)
   - **MANDATORY**: Each section has at least {user_min_per_section} user files
   - Can assign to MULTIPLE sections if content spans topics

2. **RESEARCH FILES (MUST ALSO BE USED)**:
   - Character quota: 8K per file
   - Target per section: {research_target_per_section}
   - **MANDATORY**: Use at least {int(research_file_count * research_coverage_pct / 100)} research files ({research_coverage_pct}%)
   - **MANDATORY**: Each section has at least {research_min_per_section} research files
   - **CRITICAL**: Do NOT ignore research files even when user files are available!

3. **MIXED REQUIREMENT**:
   - **FORBIDDEN**: Using only user files OR only research files in a section
   - **MANDATORY**: EVERY section must have BOTH user files AND research files
   - Optimal mix: {user_min_per_section}+ user + {research_min_per_section}+ research per section

4. **RELEVANCE CRITERIA (RELAXED)**:
   - **HIGH**: Directly discusses topic → ALWAYS include
   - **MEDIUM**: Related concepts/context → MUST include
   - **LOW**: Tangential/background → Include if space permits
   - ⚠️ When in doubt → INCLUDE the file!

📋 **ALLOCATION PRIORITY**:
  1st: High-relevance user files (HIGHEST)
  2nd: High-relevance research files (MUST include)
  3rd: Medium-relevance user files (MUST include for 100% coverage)
  4th: Medium-relevance research files (MUST include for quota)
  5th: Low-relevance user files (include if space permits)
  6th: Low-relevance research files (include if space permits)

✅ **MANDATORY PRE-SUBMISSION VERIFICATION**:

□ Step 1: Count user files assigned
   → My count: _____ user files
   → Required: {user_file_count} files (100%)
   → Status: □ PASS (= {user_file_count}) □ FAIL (< {user_file_count})

□ Step 2: Count research files assigned
   → My count: _____ research files
   → Required minimum: {int(research_file_count * research_coverage_pct / 100)} files ({research_coverage_pct}%)
   → Status: □ PASS (≥ {int(research_file_count * research_coverage_pct / 100)}) □ FAIL (< {int(research_file_count * research_coverage_pct / 100)})

□ Step 3: Check EACH section has BOTH types
   → Section 1: _____ user + _____ research □ PASS (both>0) □ FAIL
   → Section 2: _____ user + _____ research □ PASS (both>0) □ FAIL
   → (Continue for all sections)

□ Step 4: Verify minimums per section
   → Each section has ≥ {user_min_per_section} user files? □ YES □ NO
   → Each section has ≥ {research_min_per_section} research files? □ YES □ NO

⚠️ **IF ANY CHECK FAILS**:
   → DO NOT submit
   → GO BACK and assign more files
   → Ensure BOTH types in EVERY section
   → THEN verify again

🚫 **COMMON MISTAKES TO AVOID**:
❌ Using only user files (research files MUST be included!)
❌ Using only research files (user files have priority!)
❌ Not using all {user_file_count} user files (100% required!)
❌ Not meeting {research_coverage_pct}% research file quota
❌ Having sections with only one file type (BOTH required!)
❌ Being too strict about relevance (accept MEDIUM!)

✅ **SUCCESS CRITERIA**:
- User files: ALL {user_file_count} used (100%)
- Research files: {int(research_file_count * research_coverage_pct / 100)}+ used ({research_coverage_pct}%+)
- Each section: {user_min_per_section}+ user AND {research_min_per_section}+ research files
- Total citations: {user_file_count + int(research_file_count * research_coverage_pct / 100)}+ unique references
- File mix: BOTH types in EVERY section

**REMEMBER**: User files are priority, but research files MUST also be used extensively. Maximize total citations!

"""
            elif research_file_count > 0:
                # PRIMARY SCENARIO: Only research files, no user files
                # This is the MAIN use case - MAXIMIZE research file citations

                # 动态计算目标，根据research文件数量优化分配
                if research_file_count >= 30:
                    # 大量research文件（30+）
                    research_target_per_section = "10-11 research files"
                    research_min_per_section = 8
                    research_coverage_pct = 85  # 目标使用85%
                elif research_file_count >= 20:
                    # 中等数量research文件（20-29）
                    research_target_per_section = "9-11 research files"
                    research_min_per_section = 7
                    research_coverage_pct = 90  # 目标使用90%
                elif research_file_count >= 10:
                    # 较少research文件（10-19）
                    research_target_per_section = "8-10 research files"
                    research_min_per_section = 6
                    research_coverage_pct = 95  # 目标使用95%
                else:
                    # 很少research文件（<10）
                    research_target_per_section = f"ALL {research_file_count} research files"
                    research_min_per_section = min(research_file_count, 5)
                    research_coverage_pct = 100  # 目标使用100%

                user_file_priority_note = f"""
⚠️ **PRIMARY SCENARIO: RESEARCH FILES ONLY - MAXIMIZE CITATIONS** ⚠️
**THIS IS THE MAIN USE CASE - STRICT COMPLIANCE REQUIRED**

📊 **INITIAL ASSESSMENT**:
- Total research files available: {research_file_count} files
- Estimated sections in outline: ~7 sections
- **MANDATORY MINIMUM TOTAL**: You MUST assign at least {int(research_file_count * research_coverage_pct / 100)} files ({research_coverage_pct}%)
- **MANDATORY PER SECTION**: Each section MUST have at least {research_min_per_section} files

🎯 **STRICT ALLOCATION RULES (NON-NEGOTIABLE)**:

1. **MINIMUM FILES PER SECTION**: {research_min_per_section}-11 files
   - This is NOT a suggestion - it is MANDATORY
   - EVERY section must meet this minimum
   - Aim for 8-11 files per section when possible

2. **TOTAL COVERAGE REQUIREMENT**: ≥ {int(research_file_count * research_coverage_pct / 100)} files out of {research_file_count}
   - You MUST use at least {research_coverage_pct}% of all available files
   - Failure to meet this is NOT acceptable

3. **RELEVANCE CRITERIA (RELAXED - READ CAREFULLY)**:
   - **HIGH relevance**: Directly discusses section topic → ALWAYS include
   - **MEDIUM relevance**: Mentions related concepts, provides context, or offers supporting evidence → MUST include
   - **LOW relevance**: Tangentially related or provides general background → INCLUDE if space permits

   ⚠️ IMPORTANT: Do NOT be overly strict! If a file has ANY reasonable connection to the topic, INCLUDE it.
   When in doubt → INCLUDE the file (better to have more citations than fewer)

4. **MULTI-SECTION ASSIGNMENT**:
   - Files with broad topics MUST be assigned to MULTIPLE sections
   - This increases total file usage and citation diversity
   - A file can appear in 2-4 sections if its content spans multiple topics

📋 **ALLOCATION PRIORITY (FOLLOW THIS ORDER)**:
  1st: High-relevance files (assign to most relevant section)
  2nd: Medium-relevance files (assign extensively - these are CRITICAL for meeting quotas)
  3rd: Files that span multiple topics (assign to 2-3 sections)
  4th: Low-relevance files (use to fill remaining slots)

✅ **MANDATORY PRE-SUBMISSION VERIFICATION**:

Before you finalize your output, you MUST verify:

□ Step 1: Count total UNIQUE files assigned across ALL sections
   → My count: _____ files
   → Required minimum: {int(research_file_count * research_coverage_pct / 100)} files
   → Status: □ PASS (≥ {int(research_file_count * research_coverage_pct / 100)}) □ FAIL (< {int(research_file_count * research_coverage_pct / 100)})

□ Step 2: Check EACH section has minimum files
   → Section 1: _____ files (min: {research_min_per_section}) □ PASS □ FAIL
   → Section 2: _____ files (min: {research_min_per_section}) □ PASS □ FAIL
   → Section 3: _____ files (min: {research_min_per_section}) □ PASS □ FAIL
   → (Continue for all sections)

□ Step 3: Verify medium-relevance files are included
   → Have I included files with MEDIUM relevance? □ YES □ NO
   → If NO, GO BACK and add them!

⚠️ **IF ANY CHECK FAILS**: 
   → DO NOT submit your answer
   → GO BACK and assign more files
   → Lower your relevance threshold
   → Assign files to multiple sections
   → THEN verify again

🚫 **COMMON MISTAKES TO AVOID**:
❌ Being too strict about relevance (accept MEDIUM relevance!)
❌ Assigning too few files per section (need {research_min_per_section}+ per section)
❌ Not using multi-section assignments (files can go to multiple sections)
❌ Leaving files unused when they have ANY connection to topics
❌ Submitting without verification (ALWAYS count before submitting)

✅ **SUCCESS CRITERIA**:
- Total files assigned: {int(research_file_count * research_coverage_pct / 100)}+ files ({research_coverage_pct}%+)
- Each section: {research_min_per_section}-11 files
- Medium-relevance files: Extensively included
- Multi-section assignments: Used for broad-topic files
- Final report citations: 40+ unique references

**REMEMBER**: The goal is to MAXIMIZE citations. More files = better report. When in doubt, INCLUDE the file!

"""

            system_prompt = f"""You are an expert content organizer specializing in multi-dimensional file classification. 
Your task is to classify research files according to a given outline by analyzing four key dimensions.

ANALYSIS DIMENSIONS:
1. DOCUMENT TIME: Consider the temporal relevance and recency of the source
2. SOURCE AUTHORITY: Consider the credibility and expertise of the source
3. CORE CONTENT: Focus on the main themes and key insights
4. TASK RELEVANCE: Assess alignment with the specific research goals
5. INFORMATION RICHNESS: Determine the information richness of the document

CLASSIFICATION INSTRUCTIONS:
1. Read the outline and understand its structure
2. Analyze each file across the five dimensions provided
3. Accept files with moderate source authority, moderate information richness and reasonable task relevance
4. Use core content insights to find connections to outline sections, ensure the relevance of the assignment.
5. Files should be assigned to multiple sections when they contain information spanning different topics
6. Prioritize comprehensive coverage: Ensure every paragraph/chapter gets file assignments when any relevant content exists
7.Note that if the first chapter is an abstract or introduction, a corresponding file must be assigned and there cannot be an empty file.

CRITICAL REQUIREMENT - OUTLINE STRUCTURE PRESERVATION:
- The number of files assigned to each chapter cannot exceed 11, so if the number of files exceeds 11, you need the most relevant files.
- When the outline is divided into paragraphs/sections, you must preserve the original content and structure exactly. 
- Do NOT modify, rephrase, or alter any paragraph titles, headings, or structural elements from the original outline. 
- Keep all outline content completely intact in your output, including formatting and wording.
- You should split each paragraph according to the granularity of the first-level heading. You are not allowed to separate any second-level headings or smaller headings under the first-level heading into individual paragraphs.
- Formatting requirements for split paragraphs: 1. You are not allowed to modify the format of the paragraphs in the original outline, including all bold symbols, etc.; 2. If the given outline includes the title of the article, you must include the article title in the first paragraph, maintaining its original bold symbols.
{user_file_priority_note}
Strictly follow the following format for output:
paragraph 1: ...
file_path_list: file_path1, file_path2, ...

paragraph 2: ...
file_path_list: file_path3, file_path4, ...
..."""

            # 构建用户提示 - 仅包含输入数据
            user_prompt = f"""
OUTLINE TO ORGANIZE CONTENT:
{outline}

{prompt_files}
"""

            model_config = get_config().get_custom_llm_config()
            # PANGU 模型配置
            PANGU_URL = model_config.get('url') or os.getenv('MODEL_REQUEST_URL', '')
            model_name = model_config.get('model') or os.getenv("MODEL_NAME", "")
            headers = {'Content-Type': 'application/json'}

            import requests
            import litellm
            try:
                # Add retry logic for AI model call
                max_retries = 5
                response = None

                for attempt in range(max_retries):
                    try:
                        response = requests.post(
                            url=PANGU_URL,
                            headers=headers,
                            json={
                                "model": model_name,
                                "chat_template": "{% for message in messages %}{% if loop.first and messages[0]['role'] != 'system' %}{{ '[unused9]系统：[unused10]' }}{% endif %}{% if message['role'] == 'system' %}{{ message['content'] }}{% endif %}{% if message['role'] == 'assistant' %}{{'[unused9]助手：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'tool' %}{{'[unused9]工具：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'function' %}{{'[unused9]方法：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'user' %}{{'[unused9]用户：' + message['content'] + '[unused10]'}}{% endif %}{% endfor %}{% if add_generation_prompt %}{{ '[unused9]助手：' }}{% endif %}",
                                "messages": [
                                    {"role": "system", "content": "system 1"},
                                    {"role": "system", "content": system_prompt},
                                    {"role": "user", "content": user_prompt}
                                ],
                                "max_tokens": max_tokens,
                            },
                            verify=False
                        )
                        response = response.json()
                        logger.info(response)

                        break  # Success, exit retry loop
                    except Exception as e:
                        logger.warning(f"LLM API call attempt {attempt + 1} failed: {e}")
                        if attempt == max_retries - 1:
                            raise e  # Last attempt, re-raise the exception
                        time.sleep(5)  # Simple 1 second delay between retries

                if response is None:
                    raise Exception("Failed to get response after all retries")

                # ai_response = response.choices[0].message.content.strip()
                # 添加防御性检查，避免NoneType错误
                # 兼容 PANGU 模型：优先使用 content，如果为 None 则使用 reasoning_content
                message = response.get("choices", [{}])[0].get("message", {})
                content = message.get("content")

                # 如果 content 为 None，尝试使用 reasoning_content
                if content is None:
                    content = message.get("reasoning_content")
                    if content is not None:
                        logger.info("Using reasoning_content as content is None")

                if content is None:
                    raise Exception(f"AI model returned None content and reasoning_content. Response: {response}")
                ai_response = content.strip()

                import os
                import json
                log_dir = "./data_pangu"
                log_file = os.path.join(log_dir, "search_result_classifier_claude_cold_start.log")
                os.makedirs(log_dir, exist_ok=True)

                # 切换保存的方式
                conversation_history = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                    {"role": "assistant", "content": "[unused16][unused17]" + ai_response}
                ]
                with open(log_file, "a", encoding="utf-8") as f:
                    f.write(json.dumps({"messages": conversation_history}, ensure_ascii=False) + "\n")

                # 验证文件分配是否符合要求
                classification_result = ai_response.split('think>')[-1].strip()
                validation_result = self._validate_file_allocation(
                    classification_result,
                    user_file_count,
                    research_file_count,
                    has_user_files
                )

                if not validation_result["valid"]:
                    logger.warning(f"File allocation validation failed: {validation_result['message']}")
                    logger.warning(f"Expected: {validation_result['expected']}, Got: {validation_result['actual']}")

                return MCPToolResult(
                    success=True,
                    data=classification_result,
                )

            except Exception as e:
                logger.error(f"AI model call failed: {e}")
                return MCPToolResult(
                    success=False,
                    error=f"AI classification failed: {str(e)}"
                )

        except Exception as e:
            logger.error(f"Search result classifier failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    @staticmethod
    def _correct_title_format(content: str, overall_outline: str) -> str:
        """
        Correct title formats in content to match those in overall_outline.

        Args:
            content: The generated chapter content
            overall_outline: The overall outline containing correct title formats

        Returns:
            Content with corrected title formats
        """
        # Defensive: return content as-is if inputs are not valid strings
        if not content or not isinstance(content, str):
            return content or ""
        if not overall_outline or not isinstance(overall_outline, str):
            return content

        # Extract titles from overall_outline
        outline_titles = {}

        outline_titles_no_number = {}

        for line in overall_outline.split('\n'):
            line = line.strip()
            if line:
                # Extract core title content by removing various formatting symbols
                core_content = line

                # Remove leading symbols like **, -, etc.
                core_content = re.sub(r'^[\*\-\s]+', '', core_content)
                # Remove trailing symbols like **
                core_content = re.sub(r'[\*\s]+$', '', core_content)
                core_content = core_content.strip()

                if core_content:
                    # Store mapping from core content to the formatted line from outline
                    outline_titles[core_content.lower()] = line

                    # Also store a mapping with the number prefix removed
                    text_without_number = re.sub(r'^\d+(?:\.\d+)*\.?\s+', '', core_content).strip()
                    if text_without_number and text_without_number.lower() != core_content.lower():
                        outline_titles_no_number[text_without_number.lower()] = line

        # Process content line by line
        content_lines = content.split('\n')
        corrected_lines = []

        for line in content_lines:
            original_line = line
            line_stripped = line.strip()

            # Check if this line is a title (starts with # and typically has ** formatting)
            if line_stripped and re.match(r'^#+\s*[\*]*.*', line_stripped):
                # Extract core content from the title
                core_content = line_stripped

                # Remove markdown headers (#)
                core_content = re.sub(r'^#+\s*', '', core_content)
                # Remove ** formatting
                core_content = re.sub(r'^\*\*', '', core_content)
                core_content = re.sub(r'\*\*$', '', core_content)
                core_content = core_content.strip()

                # Look for exact matching title in outline_titles
                found_match = False
                core_content_lower = core_content.lower()

                for outline_core, outline_format in outline_titles.items():
                    if outline_core == core_content_lower:
                        # Replace with the correct format from overall_outline
                        corrected_lines.append(outline_format)
                        found_match = True
                        break

                if not found_match:
                    content_text_no_number = re.sub(r'^\d+(?:\.\d+)*\.?\s+', '', core_content).strip()
                    content_text_no_number_lower = content_text_no_number.lower()

                    for outline_text_no_num, outline_format in outline_titles_no_number.items():
                        if outline_text_no_num == core_content_lower or outline_text_no_num == content_text_no_number_lower:
                     
                            corrected_lines.append(outline_format)
                            found_match = True
                            break

                if not found_match:
                    # If still no match found, keep original line
                    corrected_lines.append(original_line)
            else:
                is_list_item = bool(re.match(r'^[\*\-]\s+', line_stripped))
                is_reference = line_stripped.startswith('[')
                if line_stripped and not is_list_item and not is_reference:
               
                    clean_text = re.sub(r'^\*\*(.+?)\*\*$', r'\1', line_stripped).strip()
                    clean_text_lower = clean_text.lower()
                    line_stripped_lower = line_stripped.lower()

                    found_match = False
                    for outline_text_no_num, outline_format in outline_titles_no_number.items():
                        if outline_text_no_num == clean_text_lower or outline_text_no_num == line_stripped_lower:
                     
                            corrected_lines.append(outline_format)
                            found_match = True
                            break

                    if not found_match:
                        corrected_lines.append(original_line)
                else:
                    corrected_lines.append(original_line)

        return '\n'.join(corrected_lines)

    def section_writer(
            self,
            written_chapters_summary: str,
            task_content: str,
            user_query: str,
            current_chapter_outline: str,
            overall_outline: str,
            target_file_path: str,
            key_files: List[Dict],
            model: str = "pangu_auto",
            temperature: Optional[float] = None,
            max_tokens: Optional[int] = None
    ) -> MCPToolResult:
        """
        Write the current chapter content based on given web information and chapter structure; also consider user questions, completed chapters, and overall outline to ensure content relevance while avoiding duplication or contradictions.

        Args:
            user_query: The user query, ensure the drafted content is highly relevant to the user's inquiry.
            current_chapter_outline: This field represents the current chapter structure to be drafted. When composing the chapter content, do not modify content and bold formatting symbols of the existing structure's titles!!!
            overall_outline: This field represents the overall outline of the article. When drafting the chapter content, you should consider the overall outline to ensure the chapter content is consistent with the overall outline.
            target_file_path: The path to save the chapter content
            key_files: These files are the source materials required for drafting the current chapter.
            model: AI model to use for writing the chapter content
            temperature: Creativity level for the AI response (0-1)
            max_tokens: Maximum tokens for the AI response
        """
        try:
            # Get configuration
            from config.config import get_model_config, get_storage_config
            model_config = get_model_config()
            storage_config = get_storage_config()

            # Use config values or defaults
            if temperature is None:
                temperature = model_config.get('temperature', 0.3)
            if max_tokens is None:
                max_tokens = model_config.get('max_tokens', 8192)

            key_files_dict = {}
            # Create full path relative to workspace using config
            analysis_path = storage_config.get('document_analysis_path', './doc_analysis')
            file_analysis_list = self.load_json(f"{analysis_path}/file_analysis.jsonl").data
            logger.debug("File analysis loaded successfully")

            for i, file_info in enumerate(file_analysis_list, 1):
                file_info['index'] = i  # 在这里给网页进行编号， 从1开始
                if file_info.get('file_path'):
                    key_files_dict[file_info.get('file_path')] = file_info

            prompt_files = ""
            if key_files:
                prompt_files += f"Web Information Source(s) As Follows::\n"
                for i, file_info in enumerate(key_files, 1):
                    if file_info.get('file_path') in key_files_dict:
                        file_info = key_files_dict[file_info.get('file_path')]
                        index = file_info.get('index')
                        # 通过file_path获取原始网页信息
                        file_path = file_info.get('file_path')

                        def get_file_head_content(file_path, max_length=10000):
                            try:
                                full_path = self._safe_join(file_path)
                                if not full_path.exists():
                                    return f"[Error: File does not exist: {file_path}]"
                                with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                                    content = f.read(max_length)
                                return content
                            except Exception as e:
                                return f"[Error reading file {file_path}: {str(e)}]"

                        file_content = get_file_head_content(file_path)
                        doc_time = file_info.get('doc_time', 'Not specified')
                        source_authority = file_info.get('source_authority', 'Not specified')
                        task_relevance = file_info.get('task_relevance', 'Not specified')
                        # 开始组装这些字段的值
                        prompt_files += f"\n[webpaeg{index} begin]网页时间: {doc_time}|||网页权威性：{source_authority}|||网页相关性：{task_relevance}|||网页内容：{file_content}[webpaeg{index} end]"

            # 设计system prompt
            # 检查是否有用户上传的文件
            has_user_files = any('user_uploads' in f.get('file_path', '') for f in key_files)
            user_file_priority_note = ""
            if has_user_files:
                # 统计用户文件数量
                user_file_count = sum(1 for f in key_files if 'user_uploads' in f.get('file_path', ''))
                user_file_priority_note = f"""
**CRITICAL: USER UPLOADED FILES PRIORITY AND CITATION REQUIREMENT**
- The FIRST {user_file_count} file(s) in the provided information sources are user-uploaded documents (marked with 'user_uploads' in the path).
- These user-uploaded files have HIGHEST PRIORITY and should be referenced FIRST when writing.
- **MANDATORY CITATION**: You MUST cite user-uploaded files EXTENSIVELY throughout your writing when they contain relevant information.
- **FILE ASSIGNMENT IMPLICATION**: If a user-uploaded file has been assigned to this chapter, it means the file contains at least some relevant content. You MUST carefully review the file content and find relevant parts to cite.
- **CITATION REQUIREMENT**: For each user-uploaded file assigned to this chapter that contains information relevant to the chapter topic, you MUST include at least 3-5 citations from that file. This ensures user files appear in the references section when they are actually used.
- **THOROUGH REVIEW**: Before deciding not to cite a user-uploaded file, carefully review its content to ensure you haven't missed any relevant information, background context, or supporting details that could enhance the chapter.
- **RELEVANCE-BASED CITATION**: Only cite user-uploaded files when they contain information relevant to the chapter topic. However, if a file has been assigned to this chapter, it likely contains relevant content - please review it thoroughly before deciding not to cite it.
- When you use ANY information from user-uploaded files, you MUST cite them using [X] format where X is the file's index number. User-uploaded files are typically indexed as [1], [2], etc.
- **EXAMPLE**: If you use information from the first user-uploaded file (index 1), you MUST add [1] at the end of that sentence or paragraph.
- When there are contradictions between sources, prioritize information from user-uploaded files.
- Even if the user file doesn't directly relate to the current chapter topic, it may contain important background context that should be considered.
- **IMPORTANT**: When user-uploaded files contain relevant information, failure to cite them will result in them not appearing in the references section, which is incorrect. You MUST ensure every user-uploaded file that provides relevant information receives citations.

"""

            system_prompt = f"""You are a writing master. Next, you will receive web page information, user questions, and the structure of the current chapter. You need to integrate the user's questions with the provided web content and write the chapter based on its given structure, and plan out tables based on the content of the current chapter to make the content more comprehensive and intuitive. Additionally, an overall outline and summaries of previously completed chapters will be provided for reference to avoid repetition or contradictions and ensure logical consistency within the broader framework. Specific requirements will be detailed below.

## 🌐 CRITICAL: Response Language Rules (MUST FOLLOW)
**Detect the language of the user's query and write this chapter accordingly:**
- **English query → Write this chapter in English**
- **Chinese query (中文) → Write this chapter in Chinese (中文撰写)**
- **Mixed Chinese-English query → Write this chapter in Chinese (中文撰写)**
This rule applies to ALL chapter content including: headings, body text, tables, and citations.
**IMPORTANT**: Maintain language consistency with other chapters. If previous chapters were written in Chinese, continue in Chinese. If in English, continue in English.

{user_file_priority_note}When drafting the current chapter content, strictly comply with the following requirements:
- ⚠️ **CRITICAL CITATION REQUIREMENT - CITE ALL FILES**: In the web page information I gave you, each result is in the format of [webpage X begin]...[webpage X end], where X represents the numerical index of each article. **YOU MUST cite ALL provided webpages at least once in your chapter**. This is NON-NEGOTIABLE. Please cite the context at the end of the sentence when appropriate. Please cite the context in the corresponding part of the answer in the format of the reference number [X]. If a sentence comes from multiple contexts, please list all relevant reference numbers, such as [3][5]. Remember not to collect the references at the end and return the reference numbers, but list them in the corresponding part of the answer. **MANDATORY VERIFICATION**: Before submitting your chapter, verify that you have cited EVERY webpage (1 through {len(key_files) if key_files else 0}) at least once. Count your citations: webpage 1 [✓/✗], webpage 2 [✓/✗], etc. If any webpage is not cited, GO BACK and find appropriate places to cite it. **SPECIAL EMPHASIS**: User-uploaded files (typically webpages 1-{user_file_count if has_user_files else 0}) MUST be cited multiple times (3-5 times each) when they contain relevant information.
- You can only use the provided web page information for writing, don't make up any content, ensure the accuracy of the facts. Note that when there are contradictions between the facts described in the above search results, you should use your internal knowledge to reasonably identify the correct information. If identification is impossible, you may select the most factual result based on the authority of the web pages and a voting mechanism (e.g., the description consistent with the majority of web pages). If judgment remains impossible using these methods, you may appropriately list possible differing statements, but you must not conflate different claims—prioritize ensuring factual accuracy!
- You are only permitted to write content strictly within the provided chapter framework. You are forbidden from creating additional subheadings or bullet points within the framework! However, there is a special exception: **You should proactively and actively use Markdown tables to present structured data**. When encountering data comparisons, technical parameters, multi-dimensional comparisons, statistical data, feature contrasts, timeline events, or any scenario where information can be organized in rows and columns, you MUST use tables instead of pure text narration. Tables greatly improve readability and information density. Furthermore, you are not allowed to use concise or summarizing language for narration! We must strictly ensure the information density of the writing and avoid excessive compression.
- You cannot make any changes to the structure of the chapter you are currently writing, such as the title content and the bold symbols in the title, you are not allowed to make any changes. **CRITICAL: Sub-heading numbers MUST match the chapter number.** For example, if the current chapter is "## 1. Title", then sub-headings MUST be "1.1 ...", "1.2 ...", NOT "2.1 ...". If the current chapter is "## 3. Title", sub-headings must be "3.1 ...", "3.2 ...", etc. Always derive the sub-heading prefix from the chapter number in current_chapter_outline. **Important Note:** When writing Chapter 1, if you find the chapter lacks article title, you must create one based on user query. However, this rule only applies to Chapter 1 - do not add any titles to any other chapters in the work. 
- Be careful to ensure that the narrative content is highly relevant and does not contain any common sense errors, note that although you are asked to ensure the richness of information when writing, you must ensure that the content you write is highly relevant and that the context is logically coherent and readable.
- Proceeding to explain the roles of other specified fields:
    * user_query: The user query, ensure the drafted content is highly relevant to the user's inquiry.
    * written_chapters: Reference written_chapters to avoid large amounts of repetitive or conflicting content
    * overall_outline: The purpose of giving an overall outline is to let you understand the summary of the article and avoid content inconsistent with other parts during your writing. In short, focus on writing the current chapter.
    * task_content: The task_content may provide the requirements for writing the current chapter as well as prompts for what to avoid. You can refer to this content when drafting.

**📊 TABLE USAGE STRATEGY (MUST FOLLOW):**
- **Extract and plan table data**: Before writing, extract the interrelated data from the provided web information and plan out tables. When there are comparisons of indicators, dataset sizes, ablation experiments, enumerations of hyperparameters, or any structured data, plan to present them in table form.
- **Proactively use tables**: Do NOT wait for data to "perfectly fit" a table format. Whenever you encounter 2+ items being compared, contrasted, or listed with multiple attributes, USE A TABLE.
- **Applicable scenarios include but are not limited to**: data comparison, technical parameter listing, feature/advantage/disadvantage contrast, chronological events, experimental results, statistical summaries, classification/categorization, multi-dimensional analysis, and policy/regulation comparisons.
- **Table title format**: Every table MUST have a title line above it in the format "**表X: 表格标题**" (Chinese) or "**Table X: Table Title**" (English), where X is the sequential table number within the chapter (starting from 1). For example: "**表1: 不同方法的性能对比**" or "**Table 1: Performance Comparison of Different Methods**".
- **Minimum expectation**: Each chapter (except Abstract/Introduction) should contain at least 1-2 tables when the source materials contain comparable or structured data. Actively look for opportunities to present information as tables.
- **Table quality**: Tables must have clear headers, consistent formatting, and meaningful content. Do not create tables with only 1 column or 1 data row.

Other points to note::
- If the first chapter is an **Abstract** or **Introduction**, do not include subheadings (level-2 or finer bullet points)—begin the content directly under the level-1 heading.  
- CONTENT LENGTH: Each section should contain approximately 2500 words to ensure comprehensive coverage.
- **CRITICAL HEADING FORMAT RULES:**
  * **Level 1 章节标题**: 使用 Markdown '##' 格式，如 "## 2. Molecular Mechanisms" 或 "## 2. 分子机制"
  * **Level 2 子标题**: 必须是**纯文本格式**，不能使用任何markdown符号（不要用 ###、**、* 等）
    - **错误格式**: "### 2.1 Title" 或 "**### 2.1 Title**" 或 "**2.1 Title**"（包含markdown符号）
    - **正确格式**: "2.1 Title" 或 "2.1 标题"（纯文本，只有数字+空格+标题）
  * 如果 current_chapter_outline 中的子标题包含markdown符号，你必须**移除这些符号**，只保留纯文本格式
  * 当 current_chapter_outline 中的标题符号与 overall_outline 不一致时，以 overall_outline 的标题符号为准，保持全文符号一致性
  * 这对于PDF目录识别至关重要
- Note that in Chapter 1, omit any mention of research objectives, methodology, or procedural details.
- **🌐 CRITICAL LANGUAGE RULES (MUST FOLLOW)**:
  * **English query (no Chinese characters) → Write ENTIRE chapter in English**, including title, all headings, and all content.
  * **Chinese query (contains ANY Chinese characters) → Write ENTIRE chapter in Chinese (中文)**, including title, all headings, and all content. Even if source materials are in English, translate and write in Chinese.
  * **Mixed Chinese-English query → Write ENTIRE chapter in Chinese (中文)**.
  * This rule applies to Chapter 1's title generation as well - the title MUST match the query language.
  * Technical terms may remain in English (e.g., "PINK1/Parkin pathway"), but all explanatory text must follow the language rule.

Strictly follow the following format for output:
<chapter_content>xxx</chapter_content>
"""

            user_prompt = f"""TASK CONTENT: {task_content}
    WEB PAGE INFORMATION: {prompt_files}
    OVERALL OUTLINE: {overall_outline}
    CURRENT CHAPTER OUTLINE: {current_chapter_outline}
    PREVIOUSLY WRITTEN CHAPTERS SUMMARY: {written_chapters_summary}
    USER QUERY: {user_query}"""

            # 调用AI模型进行分类
            # Get model URL and token from config
            config = get_config()
            model_config = config.get_custom_llm_config()

            model_url = model_config.get('url') or os.getenv('MODEL_REQUEST_URL', '')
            model_token = model_config.get('token') or os.getenv('MODEL_REQUEST_TOKEN', '')
            headers = {'Content-Type': 'application/json', 'csb-token': model_token}
            try:
                max_retries = 5
                response = None
                for attempt in range(max_retries):
                    try:
                        response = requests.post(
                            url=model_url,
                            headers=headers,
                            json={
                                "model": model_config.get('model', 'pangu_auto'),
                                "chat_template": "{% for message in messages %}{% if loop.first and messages[0]['role'] != 'system' %}{{ '<s>[unused9]系统：[unused10]' }}{% endif %}{% if message['role'] == 'system' %}{{'<s>[unused9]系统：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'assistant' %}{{'[unused9]助手：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'tool' %}{{'[unused9]工具：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'function' %}{{'[unused9]方法：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'user' %}{{'[unused9]用户：' + message['content'] + '[unused10]'}}{% endif %}{% endfor %}{% if add_generation_prompt %}{{ '[unused9]助手：' }}{% endif %}",
                                "messages": [
                                    {"role": "system", "content": system_prompt},
                                    {"role": "user", "content": user_prompt + " /no_think"}
                                ],
                                "spaces_between_special_tokens": False,
                                "max_tokens": max_tokens,
                                "temperature": temperature,
                            },
                            timeout=model_config.get("timeout", 180)
                        )
                        response = response.json()
                        logger.debug(f"API response received")

                        break  # Success, exit retry loop
                    except Exception as e:
                        logger.warning(f"LLM API call attempt {attempt + 1} failed: {e}")
                        if attempt == max_retries - 1:
                            raise e  # Last attempt, re-raise the exception
                        time.sleep(5)  # Simple 1 second delay between retries

                if response is None:
                    raise Exception("Failed to get response after all retries")

                # ai_response = response.choices[0].message.content.strip()
                ai_response = response["choices"][0]["message"]["content"].strip()

                # Extract content from first response
                content = ""
                if "<chapter_content>" in ai_response:
                    content = ai_response.split("<chapter_content>")[1].split("</chapter_content>")[0].strip()
                else:
                    content = ai_response

                logger.debug(f"Content before correction: {content[:200]}...")
                logger.debug(f"Overall outline: {overall_outline[:200]}...")
                content = self._correct_title_format(content, overall_outline)
                logger.debug(f"Content after correction: {content[:200]}...")
                # Second round: Request summary
                summary_prompt = "Please give a brief summary of the output chapter content. Be sure to ensure that the language of the summary is consistent with the language of the output chapter content. For example, if the chapter content is in Chinese, your summary should also be in Chinese."

                summary_response = None
                max_retries = 5
                for attempt in range(max_retries):
                    try:
                        summary_response = requests.post(
                            url=model_url,
                            headers=headers,
                            json={
                                "model": model_config.get('model', 'pangu_auto'),
                                "chat_template": "{% for message in messages %}{% if loop.first and messages[0]['role'] != 'system' %}{{ '<s>[unused9]系统：[unused10]' }}{% endif %}{% if message['role'] == 'system' %}{{'<s>[unused9]系统：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'assistant' %}{{'[unused9]助手：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'tool' %}{{'[unused9]工具：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'function' %}{{'[unused9]方法：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'user' %}{{'[unused9]用户：' + message['content'] + '[unused10]'}}{% endif %}{% endfor %}{% if add_generation_prompt %}{{ '[unused9]助手：' }}{% endif %}",
                                "messages": [
                                    {"role": "system", "content": system_prompt},
                                    {"role": "user", "content": user_prompt + " /no_think"},
                                    {"role": "assistant", "content": ai_response},
                                    {"role": "user", "content": summary_prompt + " /no_think"}
                                ],
                                "max_tokens": max_tokens,
                                "spaces_between_special_tokens": False,
                                "temperature": temperature,
                            },
                            timeout=model_config.get("timeout", 180)
                        )
                        summary_response = summary_response.json()
                        logger.debug(f"Summary API response received")

                        break  # Success, exit retry loop
                    except Exception as e:
                        logger.warning(f"Summary LLM API call attempt {attempt + 1} failed: {e}")
                        if attempt == max_retries - 1:
                            raise e  # Last attempt, re-raise the exception
                        time.sleep(5)  # Simple delay between retries

                if summary_response is None:
                    raise Exception("Failed to get summary response after all retries")

                # summary_ai_response = summary_response.choices[0].message.content.strip()
                summary_ai_response = summary_response["choices"][0]["message"]["content"].strip()
                summary = summary_ai_response

                session_context = self.get_session_context()
                session_id = session_context.get("session_id")

                # Save multi-round conversation history
                conversation_history = [
                    # {"role": "system", "content": "system 1"},
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt + " /no_think"},
                    {"role": "assistant", "content": ai_response},
                    {"role": "user", "content": summary_prompt + " /no_think"},
                    {"role": "assistant", "content": summary_ai_response}
                ]

                # 把当前内容写入到target_file_path中
                write_result = self.file_write(file_path=target_file_path,
                                               content=content,
                                               create_dirs=True)
                if not write_result.success:
                    raise Exception(f"File write failed: {write_result.error}")

                results = []
                return MCPToolResult(
                    success=True,
                    data=results.append({
                        "chapter_summary": summary,
                    }),
                    metadata={
                        'content_length': len(content),
                        'summary_length': len(summary)
                    }
                )

            except Exception as e:
                logger.error(f"AI model call failed: {e}")
                return MCPToolResult(
                    success=False,
                    error=f"section writer failed: {str(e)}"
                )

        except Exception as e:
            logger.error(f"section writer failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def document_extract(
            self,
            # save_analysis_file_path: str,
            tasks: List[Dict],
            model: str = "pangu_auto",
            temperature: Optional[float] = None,
            max_tokens: Optional[int] = None,
            max_workers: int = 5
    ) -> MCPToolResult:
        """
        Multi-dimensional analysis of locally stored files using AI models.
        Evaluates each file across four key dimensions: source authority, core content extraction,
        information richness, and query relevance scoring.

        Args:
            tasks: List of task dictionaries containing:
                - file_path: Relative path to the file (relative to workspace root) to read
                - task: task for relevance assessment
            model: AI model to use for multi-dimensional analysis
            temperature: Creativity level for the AI response (0-1)
            max_tokens: Maximum tokens for the AI response
            max_workers: Maximum number of concurrent model API requests
        """
        try:
            # Get configuration
            from config.config import get_model_config, get_storage_config
            model_config = get_model_config()
            storage_config = get_storage_config()

            # Use config values or defaults
            if temperature is None:
                temperature = model_config.get('temperature', 0.3)
            if max_tokens is None:
                max_tokens = model_config.get('max_tokens', 8192)
            logger.debug(f"Starting document extraction: tasks={tasks}")

            # 【关键修复】始终检查并自动补全 library_refs 和 user_uploads 目录的文件
            # 不再依赖 Agent 是否传入了这些目录的文件，而是主动扫描目录
            task_files = [t.get('file_path', '') for t in tasks]

            # 规范化路径：移除 ./ 前缀以便统一比较
            def normalize_path(path: str) -> str:
                return path[2:] if path.startswith('./') else path

            normalized_task_files = [normalize_path(f) for f in task_files]

            # 始终扫描 library_refs、user_uploads 和 arxiv 目录
            library_refs_dir = self.workspace_path / "library_refs"
            user_uploads_dir = self.workspace_path / "user_uploads"
            arxiv_dir = self.workspace_path / "arxiv"

            expected_files = []
            # 扫描三个目录中的所有可能的文件扩展名
            if library_refs_dir.exists():
                for ext in ['*.txt', '*.pdf', '*.doc', '*.docx']:
                    expected_files.extend([f"library_refs/{f.name}" for f in library_refs_dir.glob(ext)])
            if user_uploads_dir.exists():
                for ext in ['*.txt', '*.pdf', '*.doc', '*.docx']:
                    expected_files.extend([f"user_uploads/{f.name}" for f in user_uploads_dir.glob(ext)])
            if arxiv_dir.exists():
                for ext in ['*.txt', '*.pdf']:
                    expected_files.extend([f"arxiv/{f.name}" for f in arxiv_dir.glob(ext)])

            # 如果这些目录有文件，则进行补全检查
            if expected_files:
                # 智能匹配：基于文件主体名称（保留原始扩展名，移除 .txt 后缀）
                def get_core_name(path: str) -> str:
                    """
                    获取文件的核心名称，移除 .txt 后缀但保留原始扩展名
                    例如：
                    - file.pdf.txt -> file.pdf
                    - file.pdf -> file.pdf
                    - file.doc.txt -> file.doc
                    """
                    import os
                    name = os.path.basename(path)
                    # 如果以 .txt 结尾，移除它
                    if name.endswith('.txt'):
                        name = name[:-4]
                    return name

                # 构建核心名称到文件路径的映射（支持一对多）
                from collections import defaultdict
                expected_core_map = defaultdict(list)
                for f in expected_files:
                    expected_core_map[get_core_name(f)].append(f)

                task_core_names = {get_core_name(f) for f in normalized_task_files}

                # 找出真正缺失的文件（核心名称不在传入列表中）
                missing_core_names = set(expected_core_map.keys()) - task_core_names
                # 对于缺失的核心名称，选择第一个匹配的文件（通常是 .txt 版本）
                missing_files = set()
                for core_name in missing_core_names:
                    missing_files.add(expected_core_map[core_name][0])

                # 验证结果日志
                if missing_files:
                    # 分类统计缺失文件
                    missing_library = [f for f in missing_files if f.startswith('library_refs/')]
                    missing_uploads = [f for f in missing_files if f.startswith('user_uploads/')]

                    # 统一输出警告信息
                    missing_info = []
                    if missing_library:
                        missing_info.append(f"文档库 {len(missing_library)} 个")
                    if missing_uploads:
                        missing_info.append(f"⚠️ 用户上传 {len(missing_uploads)} 个")

                    logger.warning(f"⚠️ 检测到 {len(missing_files)} 个文件未包含在分析任务中 ({', '.join(missing_info)})")
                    logger.warning(f"预期: {len(expected_files)} 个，实际: {len(tasks)} 个 | 🔧 自动补全中...")

                    # 获取第一个任务的task内容作为默认任务描述
                    default_task = tasks[0].get('task', '文档分析') if tasks else '文档分析'

                    for missing_file in sorted(missing_files):
                        tasks.append({
                            'file_path': missing_file,
                            'task': default_task
                        })

                    logger.info(f"✅ 已补全，总任务数: {len(tasks)}")
                else:
                    logger.info(f"✅ 所有 library_refs/user_uploads/arxiv 文件已包含")
            else:
                # library_refs、user_uploads 和 arxiv 目录都没有文件，无需补全
                logger.info(f"跳过补全验证（library_refs/user_uploads/arxiv 目录无文件）")

            # 【关键修复】过滤掉二进制源文件，优先使用转换后的 .txt 文件
            # 如果同时存在 xxx.pdf 和 xxx.pdf.txt，只保留 .pdf.txt（可读取）
            # 如果同时存在 xxx.docx 和 xxx.docx.txt，只保留 .docx.txt
            # 如果同时存在 xxx.doc 和 xxx.doc.txt，只保留 .doc.txt
            # 如果只存在 xxx.pdf（无 .txt 版本），仍然保留（虽然可能读取失败）
            filtered_tasks = []
            txt_converted_files = set()  # 存储所有转换后的 .txt 文件对应的源文件名

            # 定义需要检查的源文件扩展名
            source_extensions = ['.pdf', '.docx', '.doc']

            # 第一遍：收集所有转换后的 .txt 文件对应的源文件
            for task in tasks:
                file_path = task.get('file_path', '')
                normalized_path = file_path.lstrip('./')

                # 检查是否是转换后的 .txt 文件（xxx.pdf.txt, xxx.docx.txt, xxx.doc.txt）
                if normalized_path.endswith('.txt'):
                    for ext in source_extensions:
                        potential_source = normalized_path[:-4]  # 去掉 .txt
                        if potential_source.endswith(ext):
                            txt_converted_files.add(potential_source)
                            break

            # 第二遍：过滤任务 - 如果源文件有对应的 .txt 版本，则跳过源文件
            for task in tasks:
                file_path = task.get('file_path', '')
                normalized_path = file_path.lstrip('./')

                # 检查是否是二进制源文件（pdf, docx, doc）
                should_skip = False
                for ext in source_extensions:
                    if normalized_path.endswith(ext):
                        # 检查是否有对应的 .txt 转换文件
                        if normalized_path in txt_converted_files:
                            # 找到对应的 .txt 文件，跳过这个源文件
                            logger.info(f"⏭️ 跳过 {file_path}（已有转换后的 {normalized_path}.txt）")
                            should_skip = True
                        break

                if not should_skip:
                    filtered_tasks.append(task)

            if len(filtered_tasks) < len(tasks):
                logger.info(f"✅ 过滤后任务数: {len(filtered_tasks)} (原始: {len(tasks)})")
                tasks = filtered_tasks

            def process_single_task(task: Dict) -> Dict:
                file_path = task['file_path']
                task_content = task['task']

                # 1. 读取文件内容
                read_result = self.file_read(file_path)
                if not read_result.success:
                    return {
                        'file_path': file_path,
                        'task': task_content,
                        'success': False,
                        'error': f"File read error: {read_result.error}",
                        'answer': None
                    }

                content = read_result.data
                system_prompt = (
                    "You are a text expert. Next, you will be given a document and task content. You need to analyze this document carefully and then provide multiple dimensional information for this document.\n\n"

                    "The following are some dimensional information extracted:\n"
                    "1. Web page time: According to the content of the document, extract the web page time of the document content. If it cannot be judged, it is expressed as \"unable to determine the web page time\"; otherwise, the time of the web page is output, accurate to the month, in the format of \"YYYY year MM month\", such as \"2023 June\";\n"
                    "2. Authority: According to the information of the document, judge the source of the web page to confirm the credibility of the web page.\n"
                    "3. Relevance: According to the current task (task_content) and the given document, judge whether the current document is related to the current task.\n"
                    "4. Core content: Based on this document, you make a core content summary to ensure the richness of information, with a word count of about 200 words.\n"
                    "Information richness: Estimate the total word count of substantive content in the document. Less than 200 words indicates scarcity; over 800 words suggests high richness; between these thresholds denotes moderate richness. Be careful not to just give the word count results, but also give a corresponding text description of how informative the content is.\n"

                    "Note:\n1. Ensure the document's language aligns with the extracted dimensions (e.g., Chinese content requires Chinese extraction).\n2. For **source_authority** and **task_relevance**, first provide a brief description before concluding.  \n"
                    "- **Authority**: Briefly assess the source's credibility (e.g., expertise, reputation). *Conclusion*: [High/Medium/Low].  \n"
                    "- **Relevance**: Summarize content alignment with the topic. *Conclusion*: [High/Medium/Low].\n"

                    "The final output format must be a valid JSON object:\n"
                    "{\n"
                    "  \"doc_time\": \"xxx\",\n"
                    "  \"source_authority\": \"xxx\",\n"
                    "  \"task_relevance\": \"xxx\",\n"
                    "  \"core_content\": \"xxx\",\n"
                    "  \"information_richness\": \"xxx\"\n"
                    "}\n\n"
                    "Important: Return ONLY the JSON object, no additional text or formatting."
                )

                # 构建用户提示
                user_prompt = (
                    f"DOCUMENT CONTENT:\n{content}\n"
                    # f"DOCUMENT LEN: The length of the file content is{len(content)}\n"
                    f"TASK FOR RELEVANCE ASSESSMENT: {task_content}"
                )

                # Get model URL and token from config
                config = get_config()
                model_config = config.get_custom_llm_config()

                model_url = model_config.get('url') or os.getenv('MODEL_REQUEST_URL', '')
                model_token = model_config.get('token') or os.getenv('MODEL_REQUEST_TOKEN', '')
                headers = {'Content-Type': 'application/json', 'csb-token': model_token}

                try:
                    # Add retry logic for AI model call
                    max_retries = 5
                    response = None

                    for attempt in range(max_retries):
                        try:
                            response = requests.post(
                                url=model_url,
                                headers=headers,
                                json={
                                    "model": model_config.get('model', 'pangu_auto'),
                                    "chat_template": "{% for message in messages %}{% if loop.first and messages[0]['role'] != 'system' %}{{ '<s>[unused9]系统：[unused10]' }}{% endif %}{% if message['role'] == 'system' %}{{'<s>[unused9]系统：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'assistant' %}{{'[unused9]助手：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'tool' %}{{'[unused9]工具：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'function' %}{{'[unused9]方法：' + message['content'] + '[unused10]'}}{% endif %}{% if message['role'] == 'user' %}{{'[unused9]用户：' + message['content'] + '[unused10]'}}{% endif %}{% endfor %}{% if add_generation_prompt %}{{ '[unused9]助手：' }}{% endif %}",
                                    "messages": [
                                        {"role": "system", "content": system_prompt},
                                        {"role": "user", "content": user_prompt + " /no_think"}
                                    ],
                                    "max_tokens": max_tokens,
                                    "spaces_between_special_tokens": False,
                                    "temperature": temperature,
                                },
                                timeout=model_config.get("timeout", 180)
                            )
                            response = response.json()
                            logger.info(f"LLM API response: {response}")

                            break  # Success, exit retry loop
                        except Exception as e:
                            logger.warning(f"LLM API call attempt {attempt + 1} failed: {e}")
                            if attempt == max_retries - 1:
                                raise e  # Last attempt, re-raise the exception
                            time.sleep(4)  # Simple 1 second delay between retries

                    if response is None:
                        raise Exception("Failed to get response after all retries")

                    # answer = response.choices[0].message.content
                    answer = response["choices"][0]["message"]["content"]

                    session_context = self.get_session_context()
                    session_id = session_context.get("session_id")
                    # 切换保存的方式
                    conversation_history = [
                        # {"role": "system", "content": "system 1"},
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt + " /no_think"},
                        {"role": "assistant", "content": answer}
                    ]

                    return {
                        'file_path': file_path,
                        'task': task_content,
                        'success': True,
                        'answer': answer,
                        'metadata': {
                            'file_size': len(content),
                            'line_count': len(content.splitlines())
                        }
                    }

                except Exception as e:
                    logger.error(f"Model API call failed for file '{file_path}': {e}")
                    return {
                        'file_path': file_path,
                        'task': task_content,
                        'success': False,
                        'error': f"Model API error: {str(e)}"
                    }

            # 4. 并发处理所有任务
            results = []
            with ThreadPoolExecutor(max_workers=min(max_workers, len(tasks))) as executor:
                future_to_task = {executor.submit(process_single_task, task): task for task in tasks}

                for future in as_completed(future_to_task):
                    try:
                        result = future.result()
                        results.append(result)
                    except Exception as e:
                        task = future_to_task[future]
                        logger.error(f"Task processing failed for file '{task['file_path']}': {e}")
                        results.append({
                            'file_path': task['file_path'],
                            'task': task['task'],
                            'success': False,
                            'error': f"Task processing exception: {str(e)}"
                        })

            # 5. 保持原始任务顺序
            task_order = {task['file_path']: i for i, task in enumerate(tasks)}
            results.sort(key=lambda x: task_order.get(x['file_path'], float('inf')))

            # 保存结果到文件
            def parse_answer_to_structured_data(answer_text: str, file_path: str) -> Dict[str, str]:
                """Parse the AI JSON response into structured data"""
                # Default structure
                structured_data = {
                    "file_path": file_path,
                    "doc_time": "Unknown",
                    "source_authority": "Unknown",
                    "task_relevance": "Unknown",
                    "information_richness": "Unknown",
                    "core_content": "Unknown"
                }

                if not answer_text:
                    return structured_data

                try:
                    # Try to parse as JSON
                    answer_text = answer_text.strip()

                    # Remove any markdown code blocks if present
                    if answer_text.startswith('```'):
                        lines = answer_text.split('\n')
                        # Find the start and end of JSON content
                        start_idx = 0
                        end_idx = len(lines)
                        for i, line in enumerate(lines):
                            if line.strip().startswith('{'):
                                start_idx = i
                                break
                        for i in range(len(lines) - 1, -1, -1):
                            if lines[i].strip().endswith('}'):
                                end_idx = i + 1
                                break
                        answer_text = '\n'.join(lines[start_idx:end_idx])

                    # Parse JSON
                    parsed_data = json.loads(answer_text)

                    # Update structured_data with parsed values
                    if isinstance(parsed_data, dict):
                        structured_data.update({
                            "file_path": file_path,
                            "doc_time": parsed_data.get("doc_time", "Unknown"),
                            "source_authority": parsed_data.get("source_authority", "Unknown"),
                            "task_relevance": parsed_data.get("task_relevance", "Unknown"),
                            "core_content": parsed_data.get("core_content", "Unknown"),
                            "information_richness": parsed_data.get("information_richness", "Unknown")
                        })

                    return structured_data

                except json.JSONDecodeError as e:
                    # If JSON parsing fails, return default with error info
                    structured_data[
                        "core_content"] = f"JSON parsing error: {str(e)}. Raw response: {answer_text[:200]}..."
                    return structured_data
                except Exception as e:
                    # Handle any other parsing errors
                    structured_data["core_content"] = f"Parsing error: {str(e)}"
                    return structured_data

            # Transform results into the desired format
            structured_results = []
            for result in results:
                if result.get('success', False) and result.get('answer'):
                    structured_data = parse_answer_to_structured_data(
                        result['answer'],
                        result['file_path']
                    )
                    structured_results.append(structured_data)
                else:
                    # For failed results, still include basic info
                    structured_results.append({
                        "file_path": result.get('file_path', 'Unknown'),
                        "doc_time": "Processing failed",
                        "source_authority": "Processing failed",
                        "task_relevance": "Processing failed",
                        "information_richness": "Unknown",
                        "core_content": f"Error: {result.get('error', 'Unknown error')}"
                    })

            # Save structured results to JSON file
            # Create full path relative to workspace
            analysis_path = storage_config.get('document_analysis_path', './doc_analysis')
            full_save_path = self.workspace_path / analysis_path / "file_analysis.jsonl"
            full_save_path.parent.mkdir(parents=True, exist_ok=True)

            # 【关键修复】使用核心名称作为 key 进行去重，避免重复记录
            def get_core_key(path: str) -> str:
                """
                获取文件的核心 key，用于去重
                - 移除 ./ 前缀
                - 移除转换后缀 .pdf.txt -> .pdf, .docx.txt -> .docx, .doc.txt -> .doc
                例如：./user_uploads/xxx.pdf 和 user_uploads/xxx.pdf.txt 都会映射到 user_uploads/xxx.pdf
                """
                # 移除 ./ 前缀
                normalized = path.lstrip('./')
                # 移除转换后的 .txt 后缀
                for ext in ['.pdf.txt', '.docx.txt', '.doc.txt']:
                    if normalized.endswith(ext):
                        normalized = normalized[:-4]  # 移除 .txt
                        break
                return normalized

            def is_successful_result(data: dict) -> bool:
                """判断分析结果是否成功"""
                return data.get('doc_time') != 'Processing failed'

            # Read existing data to avoid duplicates (智能去重：优先保留成功记录)
            existing_data = {}
            if full_save_path.exists():
                try:
                    with open(full_save_path, "r", encoding='utf-8', errors='ignore') as f:
                        for line in f:
                            try:
                                data = json.loads(line.strip())
                                if data.get('file_path'):
                                    core_key = get_core_key(data['file_path'])
                                    existing = existing_data.get(core_key)
                                    
                                    # 智能覆盖：优先保留成功的记录
                                    should_replace = True
                                    if existing:
                                        existing_success = is_successful_result(existing)
                                        new_success = is_successful_result(data)
                                        
                                        if existing_success and not new_success:
                                            # 已有成功的结果，不用失败的覆盖
                                            should_replace = False
                                    
                                    if should_replace:
                                        existing_data[core_key] = data
                            except json.JSONDecodeError:
                                continue  # Skip malformed lines
                except Exception as e:
                    logger.error(f"Warning: Failed to read existing analysis file: {e}")

            # Merge new results (智能覆盖：成功覆盖失败，不用失败覆盖成功)
            for result in structured_results:
                if result.get('file_path'):
                    core_key = get_core_key(result['file_path'])
                    existing = existing_data.get(core_key)
                    
                    # 决定是否覆盖
                    should_replace = True
                    if existing:
                        existing_success = is_successful_result(existing)
                        new_success = is_successful_result(result)
                        
                        if existing_success and not new_success:
                            # 已有成功的结果，不用失败的覆盖
                            logger.info(f"保留已有成功结果: {core_key}")
                            should_replace = False
                        elif new_success and not existing_success:
                            # 新结果成功，覆盖旧的失败结果
                            logger.info(f"用成功结果覆盖失败记录: {core_key}")
                    
                    if should_replace:
                        existing_data[core_key] = result

            # Write back all data (overwrite mode to ensure no duplicates)
            # 【关键修复】用户上传文件优先排序，文档库和网络检索由LLM自行判断
            # 这样确保用户上传文件在file_analysis.jsonl中排在最前面，Writer会强制优先引用
            user_uploaded_files = []
            other_files = []

            for file_data in existing_data.values():
                file_path = file_data.get('file_path', '')
                if 'user_uploads' in file_path or file_path.startswith('./user_uploads/'):
                    user_uploaded_files.append(file_data)
                else:
                    # 文档库和网络检索不做优先级区分，保持原有顺序
                    other_files.append(file_data)

            # 用户上传文件在前，其他文件在后（保持原有顺序）
            sorted_data = user_uploaded_files + other_files

            with open(full_save_path, mode="w", encoding='utf-8') as f:
                for data in sorted_data:
                    f.write(json.dumps(data, ensure_ascii=False) + "\n")

            # 统计各类文件数量
            library_count = sum(1 for d in other_files if 'library_refs' in d.get('file_path', ''))
            research_count = sum(1 for d in other_files if 'research' in d.get('file_path', ''))
            if user_uploaded_files:
                logger.info(
                    f"✅ 文件分析结果已保存: 用户上传({len(user_uploaded_files)}个,强制优先) + 文档库({library_count}个) + 网络检索({research_count}个,由LLM判断)")

            # 6. 统计结果
            successful_tasks = len([r for r in results if r.get('success', False)])

            return MCPToolResult(
                success=True,
                data=results,
                metadata={
                    'total_tasks': len(tasks),
                    'successful_tasks': successful_tasks,
                    'failed_tasks': len(tasks) - successful_tasks,
                    'model': model,
                    'concurrent_workers': min(max_workers, len(tasks))
                }
            )

        except Exception as e:
            logger.error(f"Context-based document extraction failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def document_qa(
            self,
            tasks: List[Dict],
            model: str = "gpt-4o",
            temperature: float = 0.3,
            max_tokens: int = 8192,
            max_workers: int = 5
    ) -> MCPToolResult:
        """
        Answer questions based on content stored in local files.
        Each task contains a file path and a question to be answered using that file's content.

        Args:
            tasks: List of task dictionaries containing:
                - file_path: Relative path to the file (relative to workspace root) to read
                - question: Question to ask about this file
            model: AI model to use for generating answers
            temperature: Creativity level for the AI response (0-1)
            max_tokens: Maximum tokens for the AI response
            max_workers: Maximum number of concurrent model API requests
        """
        try:
            logger.info(f"我现在开始调用document_qa了：{tasks}")

            # 处理单个任务
            def process_single_task(task: Dict) -> Dict:
                file_path = task['file_path']
                question = task['question']

                # 1. 读取文件内容
                read_result = self.file_read(file_path)
                if not read_result.success:
                    return {
                        'file_path': file_path,
                        'question': question,
                        'success': False,
                        'error': f"File read error: {read_result.error}",
                        'answer': None
                    }

                content = read_result.data

                # 2. 构建系统提示
                system_prompt = (
                    "You are an expert document analyst. Answer the user's question "
                    "based ONLY on the provided context. If the answer cannot be found "
                    "in the context, say 'I don't know'.\n\n"
                    "CONTEXT:\n{context}"
                ).format(context=content)

                # 3. 调用大模型API
                import litellm
                try:
                    # Add retry logic for AI model call
                    max_retries = 5
                    response = None

                    for attempt in range(max_retries):
                        try:
                            response = litellm.completion(
                                model=model,
                                messages=[
                                    {"role": "system", "content": system_prompt},
                                    {"role": "user", "content": question}
                                ],
                                temperature=temperature,
                                # temperature=1,
                                max_tokens=max_tokens,
                                proxy=proxy
                            )
                            break  # Success, exit retry loop
                        except Exception as e:
                            logger.warning(f"LLM API call attempt {attempt + 1} failed: {e}")
                            if attempt == max_retries - 1:
                                raise e  # Last attempt, re-raise the exception
                            time.sleep(1)  # Simple 1 second delay between retries

                    if response is None:
                        raise Exception("Failed to get response after all retries")

                    answer = response.choices[0].message.content
                    return {
                        'file_path': file_path,
                        'question': question,
                        'success': True,
                        'answer': answer,
                        'metadata': {
                            'file_size': len(content),
                            'line_count': len(content.splitlines())
                        }
                    }

                except Exception as e:
                    logger.error(f"Model API call failed for file '{file_path}': {e}")
                    return {
                        'file_path': file_path,
                        'question': question,
                        'success': False,
                        'error': f"Model API error: {str(e)}"
                    }

            # 4. 并发处理所有任务
            results = []
            with ThreadPoolExecutor(max_workers=min(max_workers, len(tasks))) as executor:
                future_to_task = {executor.submit(process_single_task, task): task for task in tasks}

                for future in as_completed(future_to_task):
                    try:
                        result = future.result()
                        results.append(result)
                    except Exception as e:
                        task = future_to_task[future]
                        logger.error(f"Task processing failed for file '{task['file_path']}': {e}")
                        results.append({
                            'file_path': task['file_path'],
                            'question': task['question'],
                            'success': False,
                            'error': f"Task processing exception: {str(e)}"
                        })

            # 5. 保持原始任务顺序
            task_order = {task['file_path']: i for i, task in enumerate(tasks)}
            results.sort(key=lambda x: task_order.get(x['file_path'], float('inf')))

            # 6. 统计结果
            successful_tasks = len([r for r in results if r.get('success', False)])

            return MCPToolResult(
                success=True,
                data=results,
                metadata={
                    'total_tasks': len(tasks),
                    'successful_tasks': successful_tasks,
                    'failed_tasks': len(tasks) - successful_tasks,
                    'model': model,
                    'concurrent_workers': min(max_workers, len(tasks))
                }
            )

        except Exception as e:
            logger.error(f"Context-based QA batch processing failed: {e}")
            logger.error(f"document qa error: '{task['file_path']}': {e}")
            return MCPToolResult(success=False, error=str(e))

    # ================ FILE DOWNLOAD TOOLS ================

    def download_files(
            self,
            urls: List[str],
            target_directory: str = None,
            overwrite: bool = False,
            max_file_size_mb: int = 100
    ) -> MCPToolResult:
        """
        Download human-readable research files such as PDFs, documents, and data files.
        
        Use this tool for downloading research papers, documentation, reports, data files (CSV, JSON, XML),
        academic publications, and other human-readable content that you can analyze.
        
        WARNING: Do NOT use this tool for downloading web pages (HTML/HTM files) or other non-readable formats.
        For web page content extraction, use the url_crawler tool instead.
        
        Args:
            urls: List of URLs to download (PDFs, DOCs, research papers, data files, etc.)
            target_directory: Directory to save files (relative to session workspace)
            overwrite: Whether to overwrite existing files
            max_file_size_mb: Maximum file size in MB
        """
        try:
            if target_directory:
                # Ensure target_directory is relative to session workspace for security
                download_dir = self._safe_join(target_directory)
            else:
                download_dir = self.workspace_path / "downloads"

            download_dir.mkdir(parents=True, exist_ok=True)

            def download_single_file(url: str) -> Dict[str, Any]:
                """Download a single file"""
                try:
                    # Parse URL to get filename
                    parsed_url = urlparse(url)
                    filename = os.path.basename(parsed_url.path) or 'downloaded_file'

                    # Ensure filename has extension
                    if '.' not in filename:
                        filename += '.html'  # Default extension

                    if os.path.isabs(filename):
                        raise Exception(f"Path '{filename}' is absolute. Only relative paths are allowed.")
                    # 检测是否为PDF文件（网页下载）
                    is_pdf_from_web = filename.lower().endswith('.pdf')
                    # 如果是PDF文件，修改扩展名为.txt以避免PDF格式问题
                    if is_pdf_from_web:
                        filename = filename[:-4] + '.txt'
                        logger.info(f"网页下载的PDF文件将转换为文本格式: {filename}")

                    file_path = download_dir / filename
                    if not os.path.realpath(file_path).startswith(self.full_workspace_path):
                        raise Exception(f"Path '{filename}' is outside workspace directory.")

                    # Check if file exists
                    if file_path.exists() and not overwrite:
                        return {
                            'url': url,
                            'success': False,
                            'error': 'File already exists',
                            'file_path': str(file_path)
                        }

                    # Download file
                    response = requests.get(url, stream=True, timeout=30)
                    response.raise_for_status()

                    # Check file size
                    content_length = response.headers.get('content-length')
                    if content_length and int(content_length) > max_file_size_mb * 1024 * 1024:
                        return {
                            'url': url,
                            'success': False,
                            'error': f'File too large (>{max_file_size_mb}MB)',
                            'file_path': None
                        }

                    # Save file
                    with open(file_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)

                    return {
                        'url': url,
                        'success': True,
                        'file_path': str(file_path),
                        'file_size': file_path.stat().st_size
                    }

                except Exception as e:
                    return {
                        'url': url,
                        'success': False,
                        'error': str(e),
                        'file_path': None
                    }

            # Process downloads concurrently
            results = []
            max_concurrent_downloads = min(5, len(urls))  # Limit concurrent downloads to avoid overwhelming servers
            with ThreadPoolExecutor(max_workers=max_concurrent_downloads) as executor:
                # Submit all download tasks
                future_to_url = {executor.submit(download_single_file, url): url for url in urls}

                # Collect results as they complete
                for future in as_completed(future_to_url):
                    try:
                        result = future.result()
                        results.append(result)
                    except Exception as e:
                        url = future_to_url[future]
                        logger.error(f"Download task failed for '{url}': {e}")
                        results.append({
                            'url': url,
                            'success': False,
                            'error': f"Download task exception: {str(e)}",
                            'file_path': None
                        })

            # Sort results to maintain original URL order
            url_order = {urls[i]: i for i in range(len(urls))}
            results.sort(key=lambda x: url_order.get(x['url'], float('inf')))

            # Generate status message following context management philosophy
            successful_downloads = len([r for r in results if r.get('success', False)])
            failed_downloads = len(results) - successful_downloads

            status_msg = f"File download task completed. Processed {len(urls)} URLs with {successful_downloads} successful downloads and {failed_downloads} failures. Files saved to {download_dir.relative_to(self.workspace_path)}. Use file reading tools to examine the downloaded files."

            return MCPToolResult(
                success=True,
                data=status_msg,
                metadata={
                    'download_directory': str(download_dir),
                    'total_urls': len(urls),
                    'successful_downloads': successful_downloads,
                    'failed_downloads': failed_downloads
                }
            )

        except Exception as e:
            logger.error(f"Download files failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def process_user_uploaded_files(
            self,
            file_ids: List[str],
            backend_url: str = "http://localhost:5000",
            target_subdir: str = "user_uploads"
    ) -> MCPToolResult:
        """
        处理用户上传的文件，下载到工作区并解析内容（标记为高优先级）

        Args:
            file_ids: 用户上传文件的ID列表
            backend_url: Flask 后端的 URL
            target_subdir: 目标子目录名称，默认为 "user_uploads"
        """
        try:
            if not file_ids:
                return MCPToolResult(success=False, error="No file IDs provided")

            user_files_dir = self.workspace_path / target_subdir
            user_files_dir.mkdir(parents=True, exist_ok=True)

            # 调用 Flask 后端 API
            response = requests.post(
                f"{backend_url}/api/user_files/download_and_parse",
                json={"file_ids": file_ids},
                timeout=60
            )

            if not response.ok:
                return MCPToolResult(success=False, error=f"HTTP {response.status_code}")

            result_data = response.json()
            processed_files = []
            # 用于跟踪文件名，避免重复文件名冲突
            used_filenames = set()

            for file_info in result_data.get('files', []):
                if not file_info.get('success'):
                    continue

                file_id = file_info.get('file_id', '')
                filename = file_info.get('filename', 'unknown.txt')

                # 安全化文件名：保留中文和特殊字符，只移除Windows不允许的字符
                # Windows保留字符和路径分隔符
                forbidden_chars = r'<>:"/\|?*'
                forbidden_chars += ''.join(chr(i) for i in range(32))  # 控制字符
                safe_filename = "".join(c for c in filename if c not in forbidden_chars)
                safe_filename = safe_filename.strip(' .')  # 移除首尾空格和点号
                if not safe_filename or safe_filename == '.':
                    safe_filename = f'file_{file_id[:8]}.pdf' if filename.endswith(
                        '.pdf') else f'file_{file_id[:8]}.txt'

                # 处理文件名冲突：如果文件名已存在，添加 file_id 前缀
                if safe_filename in used_filenames:
                    # 提取文件扩展名
                    if '.' in safe_filename:
                        name_part, ext_part = safe_filename.rsplit('.', 1)
                        safe_filename = f"{file_id[:8]}_{name_part}.{ext_part}"
                    else:
                        safe_filename = f"{file_id[:8]}_{safe_filename}"

                # 如果还是冲突（极罕见情况），添加时间戳
                if safe_filename in used_filenames:
                    import time
                    timestamp = str(int(time.time() * 1000))[-6:]
                    if '.' in safe_filename:
                        name_part, ext_part = safe_filename.rsplit('.', 1)
                        safe_filename = f"{name_part}_{timestamp}.{ext_part}"
                    else:
                        safe_filename = f"{safe_filename}_{timestamp}"

                used_filenames.add(safe_filename)
                target_path = user_files_dir / safe_filename

                # 根据文件类型决定如何处理
                file_type = file_info.get('file_type', '').lower()
                source_path = file_info.get('source_path')  # 原始文件路径（用于复制）

                # 二进制文件（PDF、Word、Excel等）需要直接复制原始文件
                binary_extensions = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.rtf', '.odt', '.epub'}
                is_binary = file_type in binary_extensions or any(
                    filename.lower().endswith(ext) for ext in binary_extensions)

                logger.info(
                    f"Processing file: {filename}, file_type={file_type}, is_binary={is_binary}, source_path={source_path}")

                if is_binary:
                    if not source_path:
                        logger.error(f"No source_path provided for binary file: {filename}")
                        continue

                    # 复制二进制文件
                    try:
                        source_file = Path(source_path)
                        if not source_file.exists():
                            logger.error(f"Source file not found: {source_path} (absolute: {source_file.absolute()})")
                            continue

                        # 执行复制
                        shutil.copy2(source_file, target_path)
                        copied_size = target_path.stat().st_size
                        source_size = source_file.stat().st_size
                        logger.info(
                            f"Copied binary file: {source_file} ({source_size} bytes) -> {target_path} ({copied_size} bytes)")

                        # 验证复制是否成功
                        if copied_size != source_size:
                            logger.warning(f"File size mismatch: source={source_size}, target={copied_size}")

                        # 对于PDF和DOCX文件，尝试提取并缓存文本版本（可选优化）
                        if file_type == '.pdf':
                            try:
                                # 提取PDF文本
                                extracted_text = self._read_pdf_text(target_path)
                                if extracted_text and len(extracted_text.strip()) > 100:
                                    # 保存文本缓存（文件名.pdf.txt）
                                    text_cache_path = target_path.with_suffix('.pdf.txt')
                                    with open(text_cache_path, 'w', encoding='utf-8') as f:
                                        f.write(extracted_text)
                                    logger.info(
                                        f"Created text cache for PDF: {text_cache_path} ({len(extracted_text)} chars)")
                                    # 更新返回路径，优先使用文本版本（Agent可以直接读取文本）
                                    # 注意：PDF原文仍然保留，文本缓存是可选的
                            except Exception as e:
                                logger.debug(f"Failed to create text cache for PDF {filename}: {e}")
                                # 缓存失败不影响主流程，继续使用PDF原文

                        elif file_type == '.docx':
                            try:
                                # 提取DOCX文本
                                extracted_text = self._read_docx_text(target_path)
                                if extracted_text and len(extracted_text.strip()) > 100:
                                    # 保存文本缓存（文件名.docx.txt）
                                    text_cache_path = target_path.with_suffix('.docx.txt')
                                    with open(text_cache_path, 'w', encoding='utf-8') as f:
                                        f.write(extracted_text)
                                    logger.info(
                                        f"Created text cache for DOCX: {text_cache_path} ({len(extracted_text)} chars)")
                            except Exception as e:
                                logger.debug(f"Failed to create text cache for DOCX {filename}: {e}")
                                # 缓存失败不影响主流程，继续使用DOCX原文

                        elif file_type == '.doc':
                            try:
                                # 提取DOC文本
                                extracted_text = self._read_doc_text(target_path)
                                if extracted_text and len(extracted_text.strip()) > 100:
                                    # 保存文本缓存（文件名.doc.txt）
                                    text_cache_path = target_path.with_suffix('.doc.txt')
                                    with open(text_cache_path, 'w', encoding='utf-8') as f:
                                        f.write(extracted_text)
                                    logger.info(
                                        f"Created text cache for DOC: {text_cache_path} ({len(extracted_text)} chars)")
                            except Exception as e:
                                logger.debug(f"Failed to create text cache for DOC {filename}: {e}")
                                # 缓存失败不影响主流程，继续使用DOC原文

                    except Exception as e:
                        logger.error(f"Error copying binary file {filename}: {e}", exc_info=True)
                        continue
                else:
                    # 文本文件：写入文本内容
                    try:
                        content = file_info.get('content', '')
                        with open(target_path, 'w', encoding='utf-8') as f:
                            f.write(content)
                        logger.info(f"Written text file: {target_path} ({len(content)} chars)")
                    except Exception as e:
                        logger.error(f"Error writing text file {filename}: {e}", exc_info=True)
                        continue

                processed_files.append({
                    'file_id': file_id,
                    'filename': filename,
                    'local_path': f"./{target_subdir}/{safe_filename}",
                    'content_length': file_info.get('content_length', 0),
                    'is_user_uploaded': True,
                    'priority': 'high',
                    'success': True
                })

            return MCPToolResult(
                success=True,
                data={
                    'files': processed_files,
                    'total_files': len(processed_files),
                    'user_files_directory': str(user_files_dir)
                }
            )

        except Exception as e:
            logger.error(f"Error processing user files: {e}")
            return MCPToolResult(success=False, error=str(e))

    def process_library_files(
            self,
            file_ids: List[str],
            backend_url: str = "http://localhost:5000",
            target_subdir: str = "library_refs"
    ) -> MCPToolResult:
        """
        处理用户从文档库选择的文件，下载到工作区（无优先级标记，由LLM自行判断）

        与 process_user_uploaded_files 的区别：
        - 保存到 library_refs 目录
        - 不标记为高优先级
        - 与网络检索文档平等，由 LLM 根据 task_relevance 等维度自行判断是否引用

        Args:
            file_ids: 文档库文件的ID列表
            backend_url: Flask 后端的 URL
            target_subdir: 目标子目录名称，默认为 "library_refs"
        """
        try:
            if not file_ids:
                return MCPToolResult(success=False, error="No file IDs provided")

            library_files_dir = self.workspace_path / target_subdir
            library_files_dir.mkdir(parents=True, exist_ok=True)

            # 调用 Flask 后端 API（与 user_uploads 使用相同的 API）
            response = requests.post(
                f"{backend_url}/api/user_files/download_and_parse",
                json={"file_ids": file_ids},
                timeout=60
            )

            if not response.ok:
                return MCPToolResult(success=False, error=f"HTTP {response.status_code}")

            result_data = response.json()
            processed_files = []
            used_filenames = set()

            for file_info in result_data.get('files', []):
                if not file_info.get('success'):
                    continue

                file_id = file_info.get('file_id', '')
                filename = file_info.get('filename', 'unknown.txt')

                # 安全化文件名
                forbidden_chars = r'<>:"/\|?*'
                forbidden_chars += ''.join(chr(i) for i in range(32))
                safe_filename = "".join(c for c in filename if c not in forbidden_chars)
                safe_filename = safe_filename.strip(' .')
                if not safe_filename or safe_filename == '.':
                    safe_filename = f'file_{file_id[:8]}.pdf' if filename.endswith(
                        '.pdf') else f'file_{file_id[:8]}.txt'

                # 处理文件名冲突
                if safe_filename in used_filenames:
                    if '.' in safe_filename:
                        name_part, ext_part = safe_filename.rsplit('.', 1)
                        safe_filename = f"{file_id[:8]}_{name_part}.{ext_part}"
                    else:
                        safe_filename = f"{file_id[:8]}_{safe_filename}"

                if safe_filename in used_filenames:
                    import time
                    timestamp = str(int(time.time() * 1000))[-6:]
                    if '.' in safe_filename:
                        name_part, ext_part = safe_filename.rsplit('.', 1)
                        safe_filename = f"{name_part}_{timestamp}.{ext_part}"
                    else:
                        safe_filename = f"{safe_filename}_{timestamp}"

                used_filenames.add(safe_filename)
                target_path = library_files_dir / safe_filename

                # 根据文件类型决定如何处理
                file_type = file_info.get('file_type', '').lower()
                source_path = file_info.get('source_path')

                binary_extensions = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.rtf', '.odt', '.epub'}
                is_binary = file_type in binary_extensions or any(
                    filename.lower().endswith(ext) for ext in binary_extensions)

                if is_binary and source_path:
                    # 复制二进制文件
                    try:
                        source_file = Path(source_path)
                        if source_file.exists():
                            shutil.copy2(source_file, target_path)

                            # 创建文本缓存（与 user_uploads 相同）
                            if file_type == '.pdf':
                                try:
                                    extracted_text = self._read_pdf_text(target_path)
                                    if extracted_text and len(extracted_text.strip()) > 100:
                                        text_cache_path = target_path.with_suffix('.pdf.txt')
                                        with open(text_cache_path, 'w', encoding='utf-8') as f:
                                            f.write(extracted_text)
                                        logger.info(f"Created text cache for library PDF: {text_cache_path}")
                                except Exception as e:
                                    logger.debug(f"Failed to create text cache for PDF {filename}: {e}")

                            elif file_type == '.docx':
                                try:
                                    extracted_text = self._read_docx_text(target_path)
                                    if extracted_text and len(extracted_text.strip()) > 100:
                                        text_cache_path = target_path.with_suffix('.docx.txt')
                                        with open(text_cache_path, 'w', encoding='utf-8') as f:
                                            f.write(extracted_text)
                                        logger.info(f"Created text cache for library DOCX: {text_cache_path}")
                                except Exception as e:
                                    logger.debug(f"Failed to create text cache for DOCX {filename}: {e}")
                    except Exception as e:
                        logger.error(f"Failed to copy library file {filename}: {e}")
                        continue
                else:
                    # 文本文件直接写入
                    content = file_info.get('content', '')
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(content)

                processed_files.append({
                    'file_id': file_id,
                    'filename': safe_filename,
                    'file_path': str(target_path.relative_to(self.workspace_path)),
                    'file_type': file_type,
                    'is_library_file': True,
                    'priority': 'normal',  # 关键：不标记为高优先级
                    'success': True
                })

            logger.info(f"✅ 文档库文件处理完成: {len(processed_files)} 个文件已保存到 {target_subdir}/")

            return MCPToolResult(
                success=True,
                data={
                    'files': processed_files,
                    'total_files': len(processed_files),
                    'library_files_directory': str(library_files_dir)
                }
            )

        except Exception as e:
            logger.error(f"Error processing library files: {e}")
            return MCPToolResult(success=False, error=str(e))

    # ================ WORKSPACE TOOLS ================

    def list_workspace(
            self,
            path: str = None,
            recursive: bool = False,
            include_hidden: bool = False,
            max_depth: int = 3
    ) -> MCPToolResult:
        """
        List files and directories in workspace with tree structure visualization
        
        Args:
            path: Specific path to list (relative to session workspace)
            recursive: Whether to list recursively
            include_hidden: Whether to include hidden files
            max_depth: Maximum recursion depth
        """
        try:
            if path:
                # Ensure path is relative to session workspace for security
                target_path = self._safe_join(path)
            else:
                target_path = self.workspace_path

            if not target_path.exists():
                return MCPToolResult(
                    success=False,
                    error=f"Path does not exist: {target_path}"
                )

            if not target_path.is_dir():
                return MCPToolResult(
                    success=False,
                    error=f"Path is not a directory: {target_path}"
                )

            items = []
            tree_structure = []

            def _list_items(current_path: Path, current_depth: int = 0):
                if current_depth > max_depth:
                    return

                try:
                    # Get all items and sort them (directories first, then files, both alphabetically)
                    all_items = list(current_path.iterdir())
                    if not include_hidden:
                        all_items = [item for item in all_items if not item.name.startswith('.')]

                    # Sort: directories first, then files, both alphabetically
                    all_items.sort(key=lambda x: (not x.is_dir(), x.name.lower()))

                    for item in all_items:
                        item_info = {
                            'name': item.name,
                            'path': str(item.relative_to(self.workspace_path)),
                            'type': 'directory' if item.is_dir() else 'file',
                            'size': item.stat().st_size if item.is_file() else None,
                            'modified': item.stat().st_mtime,
                            'depth': current_depth
                        }

                        items.append(item_info)

                        # Recurse into directories
                        if recursive and item.is_dir():
                            _list_items(item, current_depth + 1)

                except PermissionError:
                    pass  # Skip directories we can't read

            def _generate_tree_structure(current_path: Path, prefix: str = "", is_last: bool = True,
                                         current_depth: int = 0):
                """Generate ASCII tree structure recursively"""
                if current_depth > max_depth:
                    return

                try:
                    # Get all items and sort them (directories first, then files, both alphabetically)
                    all_items = list(current_path.iterdir())
                    if not include_hidden:
                        all_items = [item for item in all_items if not item.name.startswith('.')]

                    # Sort: directories first, then files, both alphabetically
                    all_items.sort(key=lambda x: (not x.is_dir(), x.name.lower()))

                    for i, item in enumerate(all_items):
                        is_last_item = i == len(all_items) - 1

                        # Choose the appropriate tree symbols
                        if is_last_item:
                            current_symbol = "└── "
                            extension = "    "
                        else:
                            current_symbol = "├── "
                            extension = "│   "

                        # Add file/directory indicator
                        if item.is_dir():
                            name_with_indicator = f"📁 {item.name}/"
                        else:
                            # Add file size for files
                            try:
                                size = item.stat().st_size
                                if size < 1024:
                                    size_str = f"{size}B"
                                elif size < 1024 * 1024:
                                    size_str = f"{size / 1024:.1f}KB"
                                else:
                                    size_str = f"{size / (1024 * 1024):.1f}MB"
                                name_with_indicator = f"📄 {item.name} ({size_str})"
                            except:
                                name_with_indicator = f"📄 {item.name}"

                        tree_line = prefix + current_symbol + name_with_indicator
                        tree_structure.append(tree_line)

                        # Recurse into directories
                        if recursive and item.is_dir():
                            _generate_tree_structure(
                                item,
                                prefix + extension,
                                is_last_item,
                                current_depth + 1
                            )

                except PermissionError:
                    tree_structure.append(prefix + "└── [Permission Denied]")

            # Generate both flat list and tree structure
            _list_items(target_path)

            # Generate tree structure
            root_name = target_path.name if target_path.name else "workspace"
            tree_structure.append(f"📁 {root_name}/")

            if recursive:
                _generate_tree_structure(target_path)
            else:
                # For non-recursive, just show immediate children
                try:
                    all_items = list(target_path.iterdir())
                    if not include_hidden:
                        all_items = [item for item in all_items if not item.name.startswith('.')]

                    all_items.sort(key=lambda x: (not x.is_dir(), x.name.lower()))

                    for i, item in enumerate(all_items):
                        is_last_item = i == len(all_items) - 1
                        symbol = "└── " if is_last_item else "├── "

                        if item.is_dir():
                            name_with_indicator = f"📁 {item.name}/"
                        else:
                            try:
                                size = item.stat().st_size
                                if size < 1024:
                                    size_str = f"{size}B"
                                elif size < 1024 * 1024:
                                    size_str = f"{size / 1024:.1f}KB"
                                else:
                                    size_str = f"{size / (1024 * 1024):.1f}MB"
                                name_with_indicator = f"📄 {item.name} ({size_str})"
                            except:
                                name_with_indicator = f"📄 {item.name}"

                        tree_structure.append(symbol + name_with_indicator)

                except PermissionError:
                    tree_structure.append("└── [Permission Denied]")

            # Create the tree string
            tree_string = "\n".join(tree_structure)

            return MCPToolResult(
                success=True,
                data={
                    'items': items,
                    'tree_structure': tree_string,
                    'tree_lines': tree_structure
                },
                metadata={
                    'target_path': str(target_path.relative_to(self.workspace_path)) if path else '.',
                    'total_items': len(items),
                    'recursive': recursive,
                    'max_depth': max_depth,
                    'include_hidden': include_hidden
                }
            )

        except Exception as e:
            logger.error(f"List workspace failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    # ================ FILE EDITING TOOLS ================
    def str_replace_based_edit_tool(
            self,
            action: str,
            file_path: str,
            content: str = None,
            old_str: str = None,
            new_str: str = None,
            line_number: int = None,
            max_char_len: int = 10000,
    ) -> MCPToolResult:
        """
        Comprehensive file editing tool
        
        Args:
            action: 'create', 'view', 'str_replace', 'insert', 'append', 'delete'
            file_path: Path to the file
            content: Content for create action
            old_str: String to replace (for str_replace)
            new_str: Replacement string (for str_replace)
            line_number: Line number for insert action
        """
        try:
            full_path = self._safe_join(file_path)

            if action == 'create':
                if full_path.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"File already exists: {file_path}"
                    )

                full_path.parent.mkdir(parents=True, exist_ok=True)
                full_path.write_text(content or '', encoding='utf-8')

                return MCPToolResult(
                    success=True,
                    data=f"File created: {file_path}",
                    metadata={'file_size': full_path.stat().st_size}
                )

            elif action == 'view':
                if not full_path.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"File does not exist: {file_path}"
                    )

                content = full_path.read_text(encoding='utf-8')
                if len(content) > max_char_len:
                    content = ("Due to the content being too long, only the first 10,000 characters are returned. "
                               "It is recommended to use other tools such as `document_qa` to extract the required content from the file. "
                               "Below is the returned portion of the file content: \n\n") + content[:max_char_len]

                return MCPToolResult(
                    success=True,
                    data=content,
                    metadata={
                        'file_size': len(content),
                        'line_count': len(content.splitlines())
                    }
                )

            elif action == 'str_replace':
                if not full_path.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"File does not exist: {file_path}"
                    )

                if not old_str or new_str is None:
                    return MCPToolResult(
                        success=False,
                        error="Both old_str and new_str are required for str_replace"
                    )

                original_content = full_path.read_text(encoding='utf-8')

                if old_str not in original_content:
                    return MCPToolResult(
                        success=False,
                        error=f"String not found: {old_str[:50]}..."
                    )

                new_content = original_content.replace(old_str, new_str)
                full_path.write_text(new_content, encoding='utf-8')

                return MCPToolResult(
                    success=True,
                    data=f"Replaced {original_content.count(old_str)} occurrence(s)",
                    metadata={
                        'old_size': len(original_content),
                        'new_size': len(new_content)
                    }
                )

            elif action == 'insert':
                if not full_path.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"File does not exist: {file_path}"
                    )

                if line_number is None or content is None:
                    return MCPToolResult(
                        success=False,
                        error="Both line_number and content are required for insert"
                    )

                lines = full_path.read_text(encoding='utf-8').splitlines()

                if line_number < 0 or line_number > len(lines):
                    return MCPToolResult(
                        success=False,
                        error=f"Invalid line number: {line_number}"
                    )

                lines.insert(line_number, content)
                full_path.write_text('\n'.join(lines), encoding='utf-8')

                return MCPToolResult(
                    success=True,
                    data=f"Inserted content at line {line_number}",
                    metadata={'new_line_count': len(lines)}
                )

            elif action == 'append':
                if not full_path.exists():
                    full_path.touch()

                with open(full_path, 'a', encoding='utf-8') as f:
                    f.write(content or '')

                return MCPToolResult(
                    success=True,
                    data=f"Appended content to {file_path}",
                    metadata={'file_size': full_path.stat().st_size}
                )

            elif action == 'delete':
                if not full_path.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"File does not exist: {file_path}"
                    )

                full_path.unlink()

                return MCPToolResult(
                    success=True,
                    data=f"Deleted file: {file_path}"
                )

            else:
                return MCPToolResult(
                    success=False,
                    error=f"Unknown action: {action}"
                )

        except Exception as e:
            logger.error(f"File edit failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    # ================ BASIC FILE TOOLS ================

    def file_read(self, file_path: str, encoding: str = 'utf-8', max_char_len: int = 10000) -> MCPToolResult:
        """Read file content"""
        try:
            full_path = self._safe_join(file_path)

            if not full_path.exists():
                return MCPToolResult(
                    success=False,
                    error=f"File does not exist: {file_path}"
                )

            content = full_path.read_text(encoding=encoding)
            # Differential quota based on file source for optimal context usage
            # Dynamic quota strategy:
            # - With user files: user=15K, research=8K (prioritize user files)
            # - Without user files: research=15K (no competition, larger quota)
            is_user_file = 'user_uploads' in str(full_path)
            user_uploads_dir = self.workspace_path / "user_uploads"

            # Safe check for user files existence
            has_user_files = False
            try:
                has_user_files = user_uploads_dir.exists() and any(user_uploads_dir.iterdir())
            except Exception:
                has_user_files = False

            if is_user_file:
                max_chars = 15000
            else:
                max_chars = 15000 if not has_user_files else 8000

            if len(content) > max_chars:
                file_type = "user-uploaded file" if is_user_file else "research file"
                content = f"Due to the {file_type} content being too long, only the first {max_chars} characters are returned. It is recommended to use other tools such as `document_qa` to extract the required content. Below is the first {max_chars} characters:\n\n" + content[
                                                                                                                                                                                                                                                                             :max_chars]

            return MCPToolResult(
                success=True,
                data=content,
                metadata={
                    'file_size': len(content),
                    'line_count': len(content.splitlines()),
                    'encoding': encoding
                }
            )

        except Exception as e:
            logger.error(f"File read failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def _read_pdf_text(self, path: Path) -> str:
        """
        读取PDF文本内容（多方案回退机制）
        优先级：PyMuPDF > pdfminer.six > PyPDF2
        """
        # 验证文件是否存在且不为空
        if not path.exists():
            logger.error(f"PDF file does not exist: {path}")
            return ''

        file_size = path.stat().st_size
        if file_size == 0:
            logger.error(f"PDF file is empty: {path}")
            return ''

        if file_size < 100:  # PDF 文件至少应该有几百字节
            logger.warning(f"PDF file suspiciously small ({file_size} bytes): {path}")

        # 检查文件头是否为 PDF 格式（%PDF-）
        try:
            with open(path, 'rb') as f:
                header = f.read(5)
                if header != b'%PDF-':
                    logger.error(f"File is not a valid PDF (header: {header}): {path}")
                    return ''
        except Exception as e:
            logger.error(f"Failed to read file header for {path}: {e}")
            return ''

        # 1. 优先使用 PyMuPDF
        try:
            import fitz  # PyMuPDF
            text = []
            with fitz.open(str(path)) as doc:
                for page in doc:
                    text.append(page.get_text())
            result = "\n".join(text)
            if result.strip():
                logger.info(f"Successfully extracted PDF text using PyMuPDF: {path}")
                return result
        except Exception as e:
            logger.warning(f"PyMuPDF failed for {path}: {e}")

        # 2. 其次使用 pdfminer.six
        try:
            from pdfminer.high_level import extract_text
            result = extract_text(str(path)) or ''
            if result.strip():
                logger.info(f"Successfully extracted PDF text using pdfminer.six: {path}")
                return result
        except Exception as e:
            logger.debug(f"pdfminer.six failed for {path}: {e}")

        # 3. 再次使用 PyPDF2（文本质量一般）
        try:
            import PyPDF2
            text = []
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text() or '')
            result = "\n".join(text)
            if result.strip():
                logger.info(f"Successfully extracted PDF text using PyPDF2: {path}")
                return result
        except Exception as e:
            logger.debug(f"PyPDF2 failed for {path}: {e}")

        logger.warning(f"All PDF extraction methods failed for {path}")
        return ''

    def _read_docx_text(self, path: Path) -> str:
        """
        读取Word文档文本内容（支持.docx格式）
        """
        try:
            from docx import Document
            doc = Document(str(path))

            # 提取所有段落文本
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # 提取表格中的文本
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(' | '.join(row_text))

            # 合并段落和表格文本
            all_text = paragraphs + tables_text
            result = '\n'.join(all_text)

            if result.strip():
                logger.info(f"Successfully extracted text from DOCX: {path} ({len(result)} chars)")
                return result
            else:
                logger.warning(f"DOCX file appears to be empty: {path}")
                return ''

        except ImportError:
            logger.error(f"python-docx library not installed. Please install: pip install python-docx")
            return ''
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {path}: {e}")
            return ''

    def _read_doc_text(self, path: Path) -> str:
        """
        读取旧版Word文档文本内容（支持.doc格式）
        多方案回退机制（纯Python优先，适配openEuler等环境）：
        1. olefile (纯Python) - 无需系统依赖，推荐
        2. win32com (Windows) - 需要MS Word
        3. antiword (系统工具) - 如果已安装
        4. textract (Python库) - 依赖较多
        """
        import sys
        import subprocess

        # 方案1: 使用olefile纯Python解析（推荐，无需系统依赖）
        try:
            text = self._extract_doc_with_olefile(path)
            if text and text.strip():
                logger.info(f"Successfully extracted text from DOC using olefile: {path} ({len(text)} chars)")
                return text.strip()
        except Exception as e:
            logger.debug(f"olefile extraction failed for {path}: {e}")

        # 方案2: Windows优先尝试win32com（如果安装了MS Word，效果最好）
        if sys.platform == 'win32':
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(path.absolute()))
                text = doc.Content.Text
                doc.Close()
                word.Quit()

                if text.strip():
                    logger.info(f"Successfully extracted text from DOC using win32com: {path} ({len(text)} chars)")
                    return text.strip()
            except ImportError:
                logger.debug(f"win32com not installed (pip install pywin32)")
            except Exception as e:
                logger.debug(f"win32com failed for {path}: {e}")

        # 方案3: 使用antiword（如果系统已安装）
        try:
            result = subprocess.run(
                ['antiword', str(path)],
                capture_output=True,
                text=True,
                timeout=30,
                check=False
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                logger.info(f"Successfully extracted text from DOC using antiword: {path} ({len(text)} chars)")
                return text
            else:
                logger.debug(f"antiword returned code {result.returncode}, stderr: {result.stderr}")
        except FileNotFoundError:
            logger.debug(f"antiword not found")
        except Exception as e:
            logger.debug(f"antiword failed for {path}: {e}")

        # 方案4: 使用textract（Python库，依赖较多）
        try:
            import textract
            text = textract.process(str(path)).decode('utf-8')
            if text.strip():
                logger.info(f"Successfully extracted text from DOC using textract: {path} ({len(text)} chars)")
                return text.strip()
        except ImportError:
            logger.debug(f"textract not installed (pip install textract)")
        except Exception as e:
            logger.debug(f"textract failed for {path}: {e}")

        # 所有方法都失败
        logger.warning(f"All DOC extraction methods failed for {path}. Recommendations:")
        logger.warning(f"  1. Install olefile: pip install olefile (Pure Python, recommended)")
        logger.warning(f"  2. Or convert the file to .docx/.pdf format manually")
        return ''

    def _extract_doc_with_olefile(self, path: Path) -> str:
        """
        使用olefile纯Python库解析.doc文件提取文本
        .doc文件是OLE Compound Document格式
        """
        import re
        
        try:
            import olefile
        except ImportError:
            logger.debug("olefile not installed. Install with: pip install olefile")
            raise ImportError("olefile not installed")

        ole = olefile.OleFileIO(str(path))
        
        try:
            # 检查是否是Word文档
            if not ole.exists('WordDocument'):
                raise ValueError("Not a valid Word document")

            # 读取WordDocument流
            word_stream = ole.openstream('WordDocument').read()
            
            # 提取文本
            text = self._extract_text_from_word_stream(word_stream)
            
            if text:
                return text
            
            # 备选：尝试从Data流提取
            if ole.exists('Data'):
                data_stream = ole.openstream('Data').read()
                text = self._decode_ole_stream(data_stream)
                if text:
                    return text
            
            return ''
        finally:
            ole.close()

    def _extract_text_from_word_stream(self, stream: bytes) -> str:
        """从WordDocument流中提取文本"""
        import re
        
        text_parts = []
        
        # 尝试UTF-16LE解码（Windows Word默认编码）
        try:
            raw_text = stream.decode('utf-16-le', errors='ignore')
            # 过滤出可打印字符
            decoded = ''.join(c for c in raw_text if c.isprintable() or c in '\n\r\t')
            if len(decoded) > 50:
                text_parts.append(decoded)
        except Exception:
            pass
        
        # 尝试提取ASCII/GBK文本（中文文档）
        try:
            # 匹配连续的可打印字符序列
            ascii_pattern = re.compile(b'[\x20-\x7e\x0a\x0d\t]{4,}')
            matches = ascii_pattern.findall(stream)
            if matches:
                ascii_text = b' '.join(matches).decode('gbk', errors='ignore')
                if len(ascii_text) > len(''.join(text_parts)):
                    text_parts = [ascii_text]
        except Exception:
            pass
        
        # 合并并清理文本
        full_text = '\n'.join(text_parts)
        return self._clean_doc_text(full_text)

    def _decode_ole_stream(self, stream: bytes) -> str:
        """解码OLE流数据"""
        try:
            text = stream.decode('utf-16-le', errors='ignore')
            text = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
            return self._clean_doc_text(text)
        except Exception:
            return ''

    def _clean_doc_text(self, text: str) -> str:
        """清理提取的文本"""
        import re
        
        if not text:
            return ''
        
        # 移除控制字符
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        # 移除Word内部标记
        text = re.sub(r'(Times New Roman|Arial|Calibri|宋体|黑体|微软雅黑)\s*', '', text)
        text = re.sub(r'HYPERLINK\s*"[^"]*"', '', text)
        # 合并多个空格和换行
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()

    def _normalize_report_part_path(self, file_path: str) -> str:
        """Normalize misformatted report chapter filenames.

        This fixes occasional LLM-generated paths like:
        - ./report/part_2_1.md
        - ./report/part_2.1.md
        - ./report/part_2.1..md

        and rewrites them to the canonical chapter filename:
        - ./report/part_2.md
        """
        try:
            path_obj = Path(file_path)
            filename = path_obj.name
            parts = path_obj.parts
            # Only normalize files under a report directory to avoid
            # touching unrelated paths.
            if "report" not in parts and path_obj.parent.name != "report":
                return file_path

            match = re.match(r'^part_(\d+)[\._]\d+.*\.md$', filename)
            if not match:
                return file_path

            chapter_index = match.group(1)
            normalized = path_obj.with_name(f"part_{chapter_index}.md")
            return normalized.as_posix()
        except Exception:
            # On any error, fall back to the original path
            return file_path

    def _clean_report_artifacts(self, content: str) -> str:
        """Clean internal marker tokens from final report content.

        This removes control tokens like [unused17] that come from the
        Pangu chat template, and normalizes accidental leakage of
        [webpaeg22] / [webpage22] markers into standard numeric
        citations [22].
        """
        if not content:
            return content

        try:
            # Remove any [unusedXX] style control tokens
            content = re.sub(r"\[unused\d+\]", "", content)

            # Normalize [webpaeg22] or [webpage22] -> [22]
            content = re.sub(r"\[webp(?:aeg|age)(\d+)\]", r"[\1]", content)
        except Exception:
            # On regex errors, return original content to be safe
            return content

        return content

    # ================ ENHANCED FILE ANALYSIS TOOLS ================

    def file_stats(self, file_path: str) -> MCPToolResult:
        """
        Get comprehensive file statistics without reading full content.
        Perfect for deciding whether to read full file or use targeted extraction.
        
        Args:
            file_path: Path to the file (relative to workspace)
        """
        try:
            full_path = self._safe_join(file_path)

            if not full_path.exists():
                return MCPToolResult(
                    success=False,
                    error=f"File does not exist: {file_path}"
                )

            if not full_path.is_file():
                return MCPToolResult(
                    success=False,
                    error=f"Path is not a file: {file_path}"
                )

            # Get basic file stats
            stat_info = full_path.stat()
            file_size = stat_info.st_size

            # Quick content analysis without loading full file
            encoding = 'utf-8'
            line_count = 0
            word_count = 0
            char_count = 0
            first_lines = []
            last_lines = []

            try:
                with open(full_path, 'r', encoding=encoding, errors='ignore') as f:
                    # Read first few lines for preview
                    for i, line in enumerate(f):
                        line_count += 1
                        if i < 5:  # First 5 lines
                            first_lines.append(line.rstrip())

                        char_count += len(line)
                        word_count += len(line.split())

                        # For efficiency, stop detailed counting after reasonable limit
                        if line_count > 10000:
                            # Estimate remaining based on average
                            remaining_size = file_size - f.tell()
                            if remaining_size > 0:
                                avg_line_size = f.tell() / line_count
                                estimated_remaining_lines = int(remaining_size / avg_line_size)
                                line_count += estimated_remaining_lines

                                # Estimate words and chars
                                avg_chars_per_line = char_count / min(line_count, 10000)
                                avg_words_per_line = word_count / min(line_count, 10000)
                                char_count += int(remaining_size)
                                word_count += int(estimated_remaining_lines * avg_words_per_line)
                            break

                # Get last few lines if file is reasonable size
                if file_size < 1024 * 1024:  # Less than 1MB
                    with open(full_path, 'r', encoding=encoding, errors='ignore') as f:
                        lines = f.readlines()
                        last_lines = [line.rstrip() for line in lines[-5:]]
                        if line_count <= 10000:  # Recalculate if we estimated
                            line_count = len(lines)
                            char_count = sum(len(line) for line in lines)
                            word_count = sum(len(line.split()) for line in lines)

            except Exception as e:
                # Try binary mode to at least get size info
                encoding = 'binary'
                char_count = file_size

            # Determine file type
            file_extension = full_path.suffix.lower()
            file_type = self._detect_file_type(full_path, file_extension)

            # Reading recommendation
            reading_recommendation = self._get_reading_recommendation(
                file_size, line_count, word_count, file_type
            )

            stats = {
                'file_path': file_path,
                'file_size_bytes': file_size,
                'file_size_human': self._format_file_size(file_size),
                'line_count': line_count,
                'word_count': word_count,
                'character_count': char_count,
                'encoding': encoding,
                'file_type': file_type,
                'file_extension': file_extension,
                'modified_time': stat_info.st_mtime,
                'is_large_file': file_size > 1024 * 1024,  # > 1MB
                'is_very_large_file': file_size > 10 * 1024 * 1024,  # > 10MB
                'first_lines_preview': first_lines,
                'last_lines_preview': last_lines,
                'reading_recommendation': reading_recommendation
            }

            return MCPToolResult(
                success=True,
                data=stats,
                metadata={
                    'analysis_method': 'efficient_sampling' if line_count > 10000 else 'full_analysis'
                }
            )

        except Exception as e:
            logger.error(f"File stats failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    @staticmethod
    def _detect_file_type(file_path: Path, extension: str) -> str:
        """Detect file type based on extension and content"""

        # Extension-based detection
        type_map = {
            '.py': 'python_code',
            '.js': 'javascript_code',
            '.ts': 'typescript_code',
            '.java': 'java_code',
            '.cpp': 'cpp_code',
            '.c': 'c_code',
            '.html': 'html_markup',
            '.css': 'css_stylesheet',
            '.json': 'json_data',
            '.xml': 'xml_data',
            '.yaml': 'yaml_config',
            '.yml': 'yaml_config',
            '.md': 'markdown_document',
            '.txt': 'plain_text',
            '.csv': 'csv_data',
            '.sql': 'sql_code',
            '.sh': 'shell_script',
            '.dockerfile': 'docker_config',
            '.env': 'environment_config'
        }

        if extension in type_map:
            return type_map[extension]

        # Content-based detection for unknown extensions
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                first_line = f.readline().strip()

                if first_line.startswith('#!'):
                    return 'executable_script'
                elif first_line.startswith('<?xml'):
                    return 'xml_data'
                elif first_line.startswith('{') or first_line.startswith('['):
                    return 'json_data'
                elif 'DOCTYPE html' in first_line or '<html' in first_line:
                    return 'html_markup'
        except:
            pass

        return 'unknown_text'

    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human readable format"""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        elif size_bytes < 1024 * 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
        else:
            return f"{size_bytes / (1024 * 1024 * 1024):.1f} GB"

    def _get_reading_recommendation(self, file_size: int, line_count: int,
                                    word_count: int, file_type: str) -> Dict[str, Any]:
        """Provide intelligent recommendations for how to read the file"""

        recommendations = {
            'strategy': 'full_read',
            'reason': 'File is small enough for full reading',
            'alternatives': []
        }

        # Large file strategies
        if file_size > 1024 * 1024:  # > 1MB
            recommendations['strategy'] = 'selective_read'
            recommendations['reason'] = 'File is large, consider targeted approaches'
            recommendations['alternatives'] = [
                'Use file_grep_with_context to search for specific content',
                'Use content_preview to get overview before full read',
                'Use file_read_lines to read specific sections',
                'Content indexing has been disabled'
            ]

        elif line_count > 1000:
            recommendations['strategy'] = 'preview_first'
            recommendations['reason'] = 'Many lines, preview recommended before full read'
            recommendations['alternatives'] = [
                'Use content_preview for quick overview',
                'Use file_grep_with_context for specific searches'
            ]

        # File type specific recommendations
        if file_type in ['json_data', 'xml_data']:
            recommendations['alternatives'].append('Consider parsing structure instead of full text read')
        elif file_type.endswith('_code'):
            recommendations['alternatives'].append('Use grep to find specific functions/classes')
        elif file_type == 'csv_data':
            recommendations['alternatives'].append('Consider reading headers first with file_read_lines')

        return recommendations

    # ================ BASIC FILE TOOLS ================
    def file_write(
            self,
            file_path: str,
            content: str,
            encoding: str = 'utf-8',
            create_dirs: bool = True
    ) -> MCPToolResult:
        """Write content to file"""
        try:
            try:
                path_obj = Path(file_path)
                parts = [p for p in path_obj.parts if p not in ('.',)]
                if parts and parts[0] == 'research' and path_obj.suffix.lower() == '.pdf':
                    new_path_obj = path_obj.with_suffix('.txt')
                    logger.info(f"Rewrite research PDF path to TXT: {file_path} -> {new_path_obj.as_posix()}")
                    file_path = new_path_obj.as_posix()
            except Exception:
                pass

            # Normalize misformatted report chapter filenames such as
            # ./report/part_2_1.md or ./report/part_2.1..md to
            # the canonical ./report/part_2.md format.
            try:
                file_path = self._normalize_report_part_path(file_path)
            except Exception:
                pass

            # For report markdown files, strip internal control markers like
            # [unused17] and normalize leaked [webpaeg22]/[webpage22] tokens
            # to standard numeric citations [22].
            try:
                path_obj = Path(file_path)
                parts = [p for p in path_obj.parts if p not in ('.',)]
                if ("report" in parts or path_obj.parent.name == "report") and path_obj.suffix.lower() == ".md":
                    content = self._clean_report_artifacts(content)
            except Exception:
                pass

            full_path = self._safe_join(file_path)

            if create_dirs:
                full_path.parent.mkdir(parents=True, exist_ok=True)

            # full_path.write_text(content, encoding=encoding)
            with open(full_path, "a", encoding=encoding) as f:
                f.write(content)

            return MCPToolResult(
                success=True,
                data=f"Written {len(content)} characters to {file_path}",
                metadata={
                    'file_size': full_path.stat().st_size,
                    'encoding': encoding
                }
            )

        except Exception as e:
            logger.error(f"File write failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    # ================ SEARCH TOOLS ================

    def file_grep_search(
            self,
            pattern: str,
            file_pattern: str = "*",
            recursive: bool = True,
            ignore_case: bool = False,
            max_matches: int = 100
    ) -> MCPToolResult:
        """Search for pattern in files using grep-like functionality"""
        try:
            import fnmatch

            flags = re.IGNORECASE if ignore_case else 0
            regex = re.compile(pattern, flags)

            matches = []
            search_path = self.workspace_path

            def _search_file(file_path: Path):
                try:
                    content = file_path.read_text(encoding='utf-8', errors='ignore')
                    lines = content.splitlines()

                    for line_num, line in enumerate(lines, 1):
                        if regex.search(line):
                            matches.append({
                                'file': str(file_path.relative_to(self.workspace_path)),
                                'line_number': line_num,
                                'line_content': line.strip(),
                                'match_start': regex.search(line).start() if regex.search(line) else 0
                            })

                            if len(matches) >= max_matches:
                                return False  # Stop searching

                    return True

                except Exception:
                    return True  # Continue searching other files

            # Search files
            if recursive:
                for file_path in search_path.rglob(file_pattern):
                    if file_path.is_file():
                        if not _search_file(file_path):
                            break
            else:
                for file_path in search_path.glob(file_pattern):
                    if file_path.is_file():
                        if not _search_file(file_path):
                            break

            return MCPToolResult(
                success=True,
                data=matches,
                metadata={
                    'pattern': pattern,
                    'total_matches': len(matches),
                    'truncated': len(matches) >= max_matches
                }
            )

        except Exception as e:
            logger.error(f"Grep search failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def file_find_by_name(
            self,
            name_pattern: str,
            recursive: bool = True,
            case_sensitive: bool = False,
            max_results: int = 100
    ) -> MCPToolResult:
        """Find files by name pattern"""
        try:
            import fnmatch

            if not case_sensitive:
                name_pattern = name_pattern.lower()

            matches = []
            search_path = self.workspace_path

            def _match_name(file_path: Path) -> bool:
                name = file_path.name
                if not case_sensitive:
                    name = name.lower()

                return fnmatch.fnmatch(name, name_pattern)

            # Search files
            if recursive:
                for file_path in search_path.rglob("*"):
                    if _match_name(file_path):
                        matches.append({
                            'name': file_path.name,
                            'path': str(file_path.relative_to(self.workspace_path)),
                            'type': 'directory' if file_path.is_dir() else 'file',
                            'size': file_path.stat().st_size if file_path.is_file() else None
                        })

                        if len(matches) >= max_results:
                            break
            else:
                for file_path in search_path.iterdir():
                    if _match_name(file_path):
                        matches.append({
                            'name': file_path.name,
                            'path': str(file_path.relative_to(self.workspace_path)),
                            'type': 'directory' if file_path.is_dir() else 'file',
                            'size': file_path.stat().st_size if file_path.is_file() else None
                        })

                        if len(matches) >= max_results:
                            break

            return MCPToolResult(
                success=True,
                data=matches,
                metadata={
                    'pattern': name_pattern,
                    'total_matches': len(matches),
                    'truncated': len(matches) >= max_results
                }
            )

        except Exception as e:
            logger.error(f"File find failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def file_read_lines(
            self,
            file_path: str,
            start_line: int = 1,
            end_line: int = None,
            max_lines: int = 1000
    ) -> MCPToolResult:
        """
        Read specific line ranges from a file without loading the entire file.
        Perfect for reading specific sections after grep or for large files.
        
        Args:
            file_path: Path to the file
            start_line: Starting line number (1-based)
            end_line: Ending line number (1-based, None for end of file)
            max_lines: Maximum number of lines to read (safety limit)
        """
        try:
            full_path = self._safe_join(file_path)

            if not full_path.exists():
                return MCPToolResult(
                    success=False,
                    error=f"File does not exist: {file_path}"
                )

            if start_line < 1:
                return MCPToolResult(
                    success=False,
                    error="start_line must be >= 1"
                )

            lines_read = []
            current_line = 0

            with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    current_line += 1

                    # Skip lines before start_line
                    if current_line < start_line:
                        continue

                    # Stop if we've reached end_line
                    if end_line and current_line > end_line:
                        break

                    # Safety check for max_lines
                    if len(lines_read) >= max_lines:
                        break

                    lines_read.append({
                        'line_number': current_line,
                        'content': line.rstrip('\n\r')
                    })

            # Calculate actual end line
            actual_end_line = lines_read[-1]['line_number'] if lines_read else start_line - 1

            return MCPToolResult(
                success=True,
                data={
                    'file_path': file_path,
                    'start_line': start_line,
                    'end_line': actual_end_line,
                    'lines': lines_read,
                    'line_count': len(lines_read)
                },
                metadata={
                    'total_lines_read': len(lines_read),
                    'truncated_due_to_max_lines': len(lines_read) >= max_lines
                }
            )

        except Exception as e:
            logger.error(f"File read lines failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    # ================ BASH TOOLS ================

    def bash(
            self,
            command: str,
            timeout: int = 30,
            capture_output: bool = True,
            working_directory: str = None
    ) -> MCPToolResult:
        """Execute bash command"""
        try:
            # Security check - prevent dangerous commands
            dangerous_patterns = [
                r'rm\s+-rf\s+/',
                r'sudo\s+rm',
                r'mkfs',
                r'dd\s+if=.*of=/dev/',
                r'>\s*/dev/sd[a-z]',
                r'cat\s+.*>\s*/dev/sd[a-z]'
            ]

            for pattern in dangerous_patterns:
                if re.search(pattern, command, re.IGNORECASE):
                    return MCPToolResult(
                        success=False,
                        error=f"Potentially dangerous command blocked: {command}"
                    )

            # Set working directory
            cwd = self.workspace_path
            if working_directory:
                cwd = Path(working_directory)
                if not cwd.exists():
                    return MCPToolResult(
                        success=False,
                        error=f"Working directory does not exist: {working_directory}"
                    )

            # Execute command
            result = subprocess.run(
                command,
                shell=True,
                cwd=str(cwd),
                capture_output=capture_output,
                text=True,
                timeout=timeout
            )

            return MCPToolResult(
                success=result.returncode == 0,
                data={
                    'stdout': result.stdout if capture_output else None,
                    'stderr': result.stderr if capture_output else None,
                    'returncode': result.returncode,
                    'command': command
                },
                metadata={
                    'execution_time': timeout,
                    'working_directory': str(cwd)
                }
            )

        except subprocess.TimeoutExpired:
            return MCPToolResult(
                success=False,
                error=f"Command timed out after {timeout} seconds"
            )
        except Exception as e:
            logger.error(f"Bash command failed: {e}")
            return MCPToolResult(success=False, error=str(e))

    def assign_multi_tasks_to_info_seeker(
            self,
            tasks: List[Dict[str, str]],
            max_workers: int = 4
    ) -> MCPToolResult:
        """
        MCP Tool: Assign multiple tasks to InformationSeekerAgents for parallel execution

        Creates multiple TaskInput objects and routes them to info_seeker agents for concurrent execution.
        This tool enables the PlannerAgent to assign multiple research tasks through the MCP tool interface.

        Args:
            tasks: List of task dictionaries with the following keys:
                - task_content (required): The specific task content
                - task_steps_for_reference: Optional reference steps for execution
                - deliverable_contents: Format of expected deliverable
                - acceptance_checking_criteria: Criteria for task completion and quality
                - workspace_id: Workspace ID for stored files and memory
                - current_task_status: Description of current task status

            max_workers: Maximum concurrent threads (default=4)

        Returns:
            MCPToolResult with execution results for all tasks
        """
        try:
            # Validate task count (1-4 tasks)
            if not (1 <= len(tasks) <= 4):
                return MCPToolResult(
                    success=False,
                    error=f"Invalid task count ({len(tasks)}). Must assign 1-4 tasks."
                )

            # Import here to avoid circular imports
            try:
                from agents import TaskInput, create_information_seeker
            except ImportError:
                from ..agents import TaskInput, create_information_seeker

            results = []
            lock = threading.Lock()

            def process_task(task: Dict[str, str]):
                """Process a single task with thread-safe result collection"""
                try:
                    # Get workspace_id from task or set default
                    task_workspace_id = task.get("workspace_id")
                    if not task_workspace_id:
                        task_workspace_id = f"info_seeker_task_{int(time.time())}"

                    # Get current_task_status from task or set default
                    task_status = task.get("current_task_status")
                    if not task_status:
                        task_status = "Task assigned to InformationSeekerAgent for execution"

                    # Create TaskInput object
                    task_input = TaskInput(
                        task_content=task["task_content"],
                        task_steps_for_reference=task.get("task_steps_for_reference"),
                        deliverable_contents=task.get("deliverable_contents"),
                        current_task_status=task_status,
                        workspace_id=task_workspace_id,
                        acceptance_checking_criteria=task.get("acceptance_checking_criteria")
                    )

                    # Create and execute with info seeker agent
                    info_seeker = create_information_seeker(
                        workspace_path=str(self.workspace_path),
                    )

                    logger.info(f"Assigning task to InformationSeekerAgent: {task['task_content'][:800]}...")

                    # Execute the task
                    result = info_seeker.execute_task(task_input)

                    # Prepare response data
                    response_data = {
                        "task_assignment": {
                            "task_content": task["task_content"],
                            "task_executor": "info_seeker",
                            "workspace_id": task_workspace_id,
                            "acceptance_criteria": task.get("acceptance_checking_criteria")
                        },
                        "execution_result": {
                            "success": result.success,
                            "iterations": result.iterations,
                            "execution_time": result.execution_time,
                            "agent_name": result.agent_name
                        }
                    }

                    # Include result data if successful
                    if result.success and result.result:
                        response_data["task_result"] = result.result

                    # Include error if failed
                    if not result.success and result.error:
                        response_data["execution_result"]["error"] = result.error

                    # Include reasoning trace summary
                    if result.reasoning_trace:
                        response_data["execution_result"]["reasoning_steps"] = len([
                            step for step in result.reasoning_trace if step.get("type") == "reasoning"
                        ])
                        response_data["execution_result"]["action_steps"] = len([
                            step for step in result.reasoning_trace if step.get("type") == "action"
                        ])

                    # Thread-safe result collection
                    with lock:
                        results.append(response_data)

                    return response_data

                except Exception as e:
                    error_msg = f"Task processing failed: {str(e)}"
                    logger.error(error_msg)
                    with lock:
                        results.append({
                            "task_content": task.get("task_content", "Unknown task"),
                            "success": False,
                            "error": error_msg
                        })
                    return None

            # Execute tasks in parallel with thread pool
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = [executor.submit(process_task, task) for task in tasks]
                # Wait for all tasks to complete
                for future in futures:
                    future.result()  # Raise exceptions if any

            # Check overall success
            all_success = all(task_result.get("success", False) for task_result in results)

            return MCPToolResult(
                success=all_success,
                data={"tasks": results},
                error=None if all_success else "Some tasks failed",
                metadata={
                    "tool_name": "assign_multi_tasks_to_info_seeker",
                    "task_count": len(tasks),
                    "success_count": sum(1 for r in results if r.get("success")),
                    "failure_count": sum(1 for r in results if not r.get("success"))
                }
            )

        except Exception as e:
            logger.error(f"Multi-task assignment failed: {e}")
            return MCPToolResult(
                success=False,
                error=f"Multi-task assignment failed: {str(e)}"
            )

    def assign_task_to_info_seeker(
            self,
            task_content: str,
            task_steps_for_reference: str = None,
            deliverable_contents: str = None,
            acceptance_checking_criteria: str = None,
            workspace_id: str = None,
            current_task_status: str = None
    ) -> MCPToolResult:
        """
        MCP Tool: Assign a task to the InformationSeekerAgent

        Creates a TaskInput object and routes it to the info_seeker agent for execution.
        This tool enables the PlannerAgent to assign research and information gathering tasks
        through the standard MCP tool interface.

        Args:
            task_content: The specific task content for the info seeker
            task_steps_for_reference: Optional reference steps for execution
            deliverable_contents: Format of expected deliverable
            acceptance_checking_criteria: Criteria for task completion and quality
            workspace_id: Workspace ID for stored files and memory
            current_task_status: Description of current task status

        Returns:
            MCPToolResult with task execution results
        """
        try:
            # Import here to avoid circular imports
            try:
                from agents import TaskInput, create_information_seeker
            except ImportError:
                from ..agents import TaskInput, create_information_seeker

            # Set default workspace if not provided
            if not workspace_id:
                workspace_id = f"info_seeker_task_{int(time.time())}"

            # Set default status if not provided
            if not current_task_status:
                current_task_status = "Task assigned to InformationSeekerAgent for execution"

            # Create TaskInput object
            task_input = TaskInput(
                task_content=task_content,
                task_steps_for_reference=task_steps_for_reference,
                deliverable_contents=deliverable_contents,
                current_task_status=current_task_status,
                task_executor="info_seeker",
                workspace_id=workspace_id,
                acceptance_checking_criteria=acceptance_checking_criteria
            )

            # Create and execute with info seeker agent
            info_seeker = create_information_seeker(workspace_path=str(self.workspace_path))

            logger.info(f"Assigning task to InformationSeekerAgent: {task_content[:100]}...")

            # Execute the task
            result = info_seeker.execute_task(task_input)

            # Prepare response data
            response_data = {
                "task_assignment": {
                    "task_content": task_content,
                    "task_executor": "info_seeker",
                    "workspace_id": workspace_id,
                    "acceptance_criteria": acceptance_checking_criteria
                },
                "execution_result": {
                    "success": result.success,
                    "iterations": result.iterations,
                    "execution_time": result.execution_time,
                    "agent_name": result.agent_name
                }
            }

            # Include result data if successful
            if result.success and result.result:
                response_data["task_result"] = result.result

            # Include error if failed
            if not result.success and result.error:
                response_data["execution_result"]["error"] = result.error

            # Include reasoning trace summary
            if result.reasoning_trace:
                response_data["execution_result"]["reasoning_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "reasoning"
                ])
                response_data["execution_result"]["action_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "action"
                ])

            return MCPToolResult(
                success=result.success,
                data=response_data,
                error=result.error if not result.success else None,
                metadata={
                    "tool_name": "assign_task_to_info_seeker",
                    "task_executor": "info_seeker",
                    "workspace_id": workspace_id,
                    "execution_time": result.execution_time
                }
            )

        except Exception as e:
            logger.error(f"Error assigning task to info seeker: {e}")
            return MCPToolResult(
                success=False,
                error=f"Failed to assign task to InformationSeekerAgent: {str(e)}"
            )

    def assign_task_to_writer(
            self,
            # save_analysis_file_path: str,
            task_content: str,
            user_query: str,
            key_files: List[Dict[str, str]]
    ) -> MCPToolResult:
        """
        MCP Tool: Assign a task to the WriterAgent

        Creates a WriterAgentTaskInput object and routes it to the writer agent for execution.
        This tool enables the PlannerAgent to assign content creation and writing tasks
        through the standard MCP tool interface.

        Args:
            task_content: Detailed description of the writing task to be performed
            summary: Consolidated summary from information seeker results to guide outline generation and content structure
            key_files: Curated list of relevant files with file_path and desc for each file

        Returns:
            MCPToolResult with task execution results
        """
        try:
            # Import here to avoid circular imports
            try:
                from agents.base_agent import WriterAgentTaskInput
                from agents import create_writer_agent
            except ImportError:
                from ..agents.base_agent import WriterAgentTaskInput
                from ..agents import create_writer_agent

            # Generate workspace ID using timestamp
            workspace_id = f"writer_task_{int(time.time())}"

            # Create WriterAgentTaskInput object
            task_input = WriterAgentTaskInput(
                # save_analysis_file_path=save_analysis_file_path,
                user_query=user_query,
                task_content=task_content,
                key_files=key_files,
                workspace_id=workspace_id,
            )

            # Create and execute with writer agent
            writer = create_writer_agent(workspace_path=str(self.workspace_path))

            logger.info(f"Assigning task to WriterAgent: {task_content[:800]}...")

            # Execute the task
            result = writer.execute_task(task_input)

            # Prepare response data
            response_data = {
                "task_assignment": {
                    "task_content": task_content,
                    "task_executor": "writer",
                    "workspace_id": workspace_id,
                    "user_query": user_query,
                    "key_files_count": len(key_files)
                },
                "execution_result": {
                    "success": result.success,
                    "iterations": result.iterations,
                    "execution_time": result.execution_time,
                    "agent_name": result.agent_name
                }
            }

            # Include result data if successful
            if result.success and result.result:
                response_data["task_result"] = result.result

            # Include error if failed
            if not result.success and result.error:
                response_data["execution_result"]["error"] = result.error

            # Include reasoning trace summary
            if result.reasoning_trace:
                response_data["execution_result"]["reasoning_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "reasoning"
                ])
                response_data["execution_result"]["action_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "action"
                ])

            return MCPToolResult(
                success=result.success,
                data=response_data,
                error=result.error if not result.success else None,
                metadata={
                    "tool_name": "assign_task_to_writer",
                    "task_executor": "writer",
                    "workspace_id": workspace_id,
                    "execution_time": result.execution_time
                }
            )

        except Exception as e:
            logger.error(f"Error assigning task to writer: {e}")
            return MCPToolResult(
                success=False,
                error=f"Failed to assign task to WriterAgent: {str(e)}"
            )

    def assign_task_to_section_writer(
            self,
            task_content: str,
            write_file_path: str,
            user_query: str,
            overall_outline: str,
            current_chapter_outline: str,
            key_files: List[Dict[str, str]]
    ) -> MCPToolResult:
        """
        MCP Tool: Assign a task to the SectionWriterAgent

        Creates a SectionWriterTaskInput object and routes it to the section writer agent for execution.
        This tool enables the PlannerAgent to assign specific section writing tasks
        through the standard MCP tool interface.

        Args:
            write_file_path: The path where the section content should be written
            current_chapter_outline: The outline for the current chapter/section to be written
            key_files: Curated list of relevant files with file_path and desc for each file

        Returns:
            MCPToolResult with task execution results
        """
        try:
            # Import here to avoid circular imports
            try:
                from agents.base_agent import SectionWriterTaskInput
                from agents.section_writer import create_section_writer
            except ImportError:
                from ..agents.base_agent import SectionWriterTaskInput
                from ..agents.section_writer import create_section_writer

            # Generate workspace ID using timestamp
            workspace_id = f"section_writer_task_{int(time.time())}"

            # Create SectionWriterTaskInput object
            task_input = SectionWriterTaskInput(
                task_content=task_content,
                write_file_path=write_file_path,
                current_chapter_outline=current_chapter_outline,
                overall_outline=overall_outline,
                user_query=user_query,
                key_files=key_files,
                workspace_id=workspace_id,
            )

            # Create and execute with section writer agent
            section_writer = create_section_writer(workspace_path=str(self.workspace_path))

            logger.info(f"Assigning task to SectionWriterAgent: {write_file_path}")

            # Execute the task
            result = section_writer.execute_task(task_input, write_file_path)

            # Prepare response data
            response_data = {
                "task_assignment": {
                    "write_file_path": write_file_path,
                    "current_chapter_outline": current_chapter_outline,
                    "task_executor": "section_writer",
                    "workspace_id": workspace_id,
                    "key_files_count": len(key_files)
                },
                "execution_result": {
                    "success": result.success,
                    "iterations": result.iterations,
                    "execution_time": result.execution_time,
                    "agent_name": result.agent_name
                }
            }

            # Include result data if successful
            if result.success and result.result:
                response_data["task_result"] = result.result

            # Include error if failed
            if not result.success and result.error:
                response_data["execution_result"]["error"] = result.error

            # Include reasoning trace summary
            if result.reasoning_trace:
                response_data["execution_result"]["reasoning_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "reasoning"
                ])
                response_data["execution_result"]["action_steps"] = len([
                    step for step in result.reasoning_trace if step.get("type") == "action"
                ])

            return MCPToolResult(
                success=result.success,
                data=response_data,
                error=result.error if not result.success else None,
                metadata={
                    "tool_name": "assign_task_to_section_writer",
                    "task_executor": "section_writer",
                    "workspace_id": workspace_id,
                    "execution_time": result.execution_time
                }
            )

        except Exception as e:
            logger.error(f"Error assigning task to section writer: {e}")
            return MCPToolResult(
                success=False,
                error=f"Failed to assign task to SectionWriterAgent: {str(e)}"
            )

    def assign_task_to_agent(
            self,
            task_content: str,
            task_executor: str,
            task_steps_for_reference: str = None,
            deliverable_contents: str = None,
            acceptance_checking_criteria: str = None,
            workspace_id: str = None,
            current_task_status: str = None
    ) -> MCPToolResult:
        """
        MCP Tool: Generic task assignment to any agent

        Routes a task to the specified agent (info_seeker or writer).
        This provides a unified interface for task assignment.

        Args:
            task_content: The specific task content
            task_executor: Which agent to route to ("info_seeker" or "writer")
            task_steps_for_reference: Optional reference steps for execution
            deliverable_contents: Format of expected deliverable
            acceptance_checking_criteria: Criteria for task completion and quality
            workspace_id: Workspace ID for stored files and memory
            current_task_status: Description of current task status

        Returns:
            MCPToolResult with task execution results
        """
        try:
            # Validate task_executor
            if task_executor not in ["info_seeker", "writer"]:
                return MCPToolResult(
                    success=False,
                    error=f"Invalid task_executor '{task_executor}'. Must be 'info_seeker' or 'writer'"
                )

            # Route to specific assignment method
            if task_executor == "info_seeker":
                return self.assign_task_to_info_seeker(
                    task_content=task_content,
                    task_steps_for_reference=task_steps_for_reference,
                    deliverable_contents=deliverable_contents,
                    acceptance_checking_criteria=acceptance_checking_criteria,
                    workspace_id=workspace_id,
                    current_task_status=current_task_status
                )
            else:  # writer
                return MCPToolResult(
                    success=False,
                    error="Writer agent assignment via assign_task_to_agent is no longer supported. Use assign_task_to_writer directly with summary and key_files parameters."
                )

        except Exception as e:
            logger.error(f"Error in generic task assignment: {e}")
            return MCPToolResult(
                success=False,
                error=f"Failed to assign task: {str(e)}"
            )

    def semantic_search(self, **kwargs) -> MCPToolResult:
        """
        Search semantically through system-maintained knowledge index.
        Uses high-performance Faiss when available, fallback to JSON-based search.
        """
        try:
            # Extract parameters
            query = kwargs.get('query', '')
            max_tokens = kwargs.get('max_tokens', 2000)
            max_results = kwargs.get('max_results', 5)
            similarity_threshold = kwargs.get('similarity_threshold', 0.7)
            filters = kwargs.get('filters', {})

            if not query:
                return MCPToolResult(
                    success=False,
                    error="query is required for semantic search"
                )

            # Check OpenAI availability
            if not hasattr(self.config, 'get_openai_client') or not self.config.get_openai_client():
                return MCPToolResult(
                    success=False,
                    error="OpenAI API key required for embeddings. Please set OPENAI_API_KEY."
                )

            # Use Faiss-based system if available for high performance
            if FAISS_AVAILABLE and get_optimized_knowledge_manager:
                try:
                    manager = get_optimized_knowledge_manager(self.config)
                    search_result = manager.search(
                        query=query,
                        max_tokens=max_tokens,
                        max_results=max_results,
                        similarity_threshold=similarity_threshold,
                        filters=filters
                    )

                    if search_result['success']:
                        return MCPToolResult(
                            success=True,
                            data=search_result
                        )
                    else:
                        # Fallback to JSON-based search on error
                        logger.warning(
                            f"Faiss search failed: {search_result.get('error')}, falling back to JSON search")
                except Exception as e:
                    logger.warning(f"Faiss search error: {e}, falling back to JSON search")

            # Fallback to JSON-based search
            logger.info("Using JSON-based search (install faiss-cpu for better performance)")
            client = self.config.get_openai_client()

            # Use system-managed index (session_knowledge.json)
            index_file = "session_knowledge.json"
            try:
                with open(index_file, 'r', encoding='utf-8') as f:
                    index_data = json.load(f)
            except FileNotFoundError:
                return MCPToolResult(
                    success=True,
                    data={
                        'query': query,
                        'results': [],
                        'total_matches': 0,
                        'message': 'No knowledge index found yet. System will build index as agents complete tasks.',
                        'search_metadata': {
                            'similarity_threshold': similarity_threshold,
                            'max_tokens_requested': max_tokens,
                            'embedding_model': 'text-embedding-3-small',
                            'vector_store': 'JSON-based (fallback)'
                        }
                    }
                )

            if not index_data:
                return MCPToolResult(
                    success=True,
                    data={
                        'query': query,
                        'results': [],
                        'total_matches': 0,
                        'message': 'Knowledge index is empty.',
                        'search_metadata': {
                            'similarity_threshold': similarity_threshold,
                            'max_tokens_requested': max_tokens,
                            'embedding_model': 'text-embedding-3-small'
                        }
                    }
                )

            # Generate query embedding
            response = client.embeddings.create(
                input=query,
                model="text-embedding-3-small"
            )
            query_embedding = response.data[0].embedding

            # Calculate similarities and apply filters
            candidate_results = []
            for item in index_data:
                # Apply filters
                if filters.get('task_name') and item['task_name'] != filters['task_name']:
                    continue
                if filters.get('file_path') and filters['file_path'] not in item['file_path']:
                    continue
                if filters.get('is_final_output') is not None and item['is_final_output'] != filters['is_final_output']:
                    continue

                # Calculate cosine similarity
                import numpy as np
                item_embedding = np.array(item['embedding'])
                query_emb = np.array(query_embedding)

                similarity = np.dot(query_emb, item_embedding) / (
                        np.linalg.norm(query_emb) * np.linalg.norm(item_embedding)
                )

                if similarity >= similarity_threshold:
                    candidate_results.append({
                        'task_name': item['task_name'],
                        'file_path': item['file_path'],
                        'file_desc': item['file_desc'],
                        'is_final_output': item['is_final_output'],
                        'chunk_index': item['chunk_index'],
                        'content': item['chunk_content'],
                        'similarity_score': float(similarity),
                        'token_count': item.get('token_count', len(item['chunk_content'].split()))
                    })

            # Sort by similarity
            candidate_results.sort(key=lambda x: x['similarity_score'], reverse=True)

            # Apply token limit - intelligent selection
            selected_results = []
            total_tokens = 0

            for result in candidate_results:
                result_tokens = result['token_count']
                if total_tokens + result_tokens <= max_tokens and len(selected_results) < max_results:
                    selected_results.append(result)
                    total_tokens += result_tokens
                elif len(selected_results) < max_results:
                    # Try to fit a shorter excerpt if we have space
                    remaining_tokens = max_tokens - total_tokens
                    if remaining_tokens > 100:  # Minimum meaningful excerpt
                        words = result['content'].split()
                        excerpt_words = words[:remaining_tokens]
                        excerpt_content = ' '.join(excerpt_words) + '...' if len(words) > remaining_tokens else result[
                            'content']

                        result_copy = result.copy()
                        result_copy['content'] = excerpt_content
                        result_copy['token_count'] = len(excerpt_words)
                        result_copy['is_excerpt'] = True

                        selected_results.append(result_copy)
                        total_tokens += len(excerpt_words)
                        break

            return MCPToolResult(
                success=True,
                data={
                    'query': query,
                    'results': selected_results,
                    'total_matches': len(candidate_results),
                    'tokens_used': total_tokens,
                    'search_metadata': {
                        'similarity_threshold': similarity_threshold,
                        'max_tokens_requested': max_tokens,
                        'max_results_requested': max_results,
                        'filters_applied': filters,
                        'embedding_model': 'text-embedding-3-small',
                        'total_candidates_found': len(candidate_results)
                    }
                }
            )

        except Exception as e:
            return MCPToolResult(
                success=False,
                error=f"Semantic search failed: {str(e)}"
            )

    def _create_text_chunks(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """
        Split text into overlapping chunks for better search coverage.
        Uses word-based chunking to maintain readability.
        """
        words = text.split()
        chunks = []

        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            chunks.append(chunk_text)

            # Stop if we've reached the end
            if i + chunk_size >= len(words):
                break

        return chunks

    def knowledge_status(self, **kwargs) -> MCPToolResult:
        """
        Get status and statistics about the system-managed knowledge index.
        Shows which vector store system is active (Faiss vs JSON).
        """
        try:
            # Try Faiss-based system first for optimal performance
            if FAISS_AVAILABLE and get_optimized_knowledge_manager:
                try:
                    manager = get_optimized_knowledge_manager(self.config)
                    stats = manager.get_stats()

                    # Add system information
                    stats['vector_store_system'] = 'Faiss (High Performance)'
                    stats['performance'] = 'O(log n) search complexity'
                    stats['recommendation'] = 'Using optimal vector storage system'

                    return MCPToolResult(
                        success=True,
                        data=stats
                    )
                except Exception as e:
                    logger.warning(f"Faiss status error: {e}, checking JSON fallback")

            # Fallback to JSON-based system
            try:
                from knowledge.knowledge_manager import get_knowledge_manager
            except ImportError:
                from ..knowledge.knowledge_manager import get_knowledge_manager

            manager = get_knowledge_manager(self.config)
            stats = manager.get_index_stats()

            # Add performance information
            stats['vector_store_system'] = 'JSON-based (Fallback)'
            stats['performance'] = 'O(n) linear search'
            stats['recommendation'] = 'Install faiss-cpu for better performance: pip install faiss-cpu'

            return MCPToolResult(
                success=True,
                data=stats
            )

        except Exception as e:
            return MCPToolResult(
                success=False,
                error=f"Failed to get knowledge status: {str(e)}"
            )

    def search_pubmed_key_words(self, keywords, max_results=10) -> MCPToolResult:
        # 生成搜索 URL
        try:
            search_url = generate_pubmed_search_url(term=keywords, num_results=max_results)
            logger.info(f"Generated URL: {search_url}")

            # 获取并解析搜索结果
            pmids = search_pubmed(search_url)

            articles = []
            for pmid in pmids:
                metadata = get_pubmed_metadata(pmid)
                if metadata:
                    articles.append(metadata)

            return MCPToolResult(success=True, data=articles)
        except Exception as e:
            logger.error(e)
            return MCPToolResult(success=False, error=f"获取pubmed信息失败!{e}")

    def search_pubmed_advanced(self, term, title, author, journal, start_date, end_date, num_results) -> MCPToolResult:
        # 生成搜索 URL
        try:
            search_url = generate_pubmed_search_url(term=term, title=title, author=author,
                                                    journal=journal, start_date=start_date,
                                                    end_date=end_date, num_results=num_results)
            logger.info(f"Generated URL: {search_url}")

            # 获取并解析搜索结果
            pmids = search_pubmed(search_url)

            articles = []
            for pmid in pmids:
                metadata = get_pubmed_metadata(pmid)
                if metadata:
                    articles.append(metadata)

            return MCPToolResult(success=True, data=articles)
        except Exception as e:
            logger.error(e)
            return MCPToolResult(success=False, error=f"获取pubmed信息失败!{e}")

    def _extract_pmc_fulltext_xml(self, pmc_id_num: str) -> str:
        """
        通过PMC efetch API获取全文XML并提取纯文本
        pmc_id_num: 纯数字PMC ID（不含'PMC'前缀）
        """
        efetch_pmc_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?db=pmc&id={pmc_id_num}&rettype=full&retmode=xml"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        }
        response = requests.get(efetch_pmc_url, headers=headers, verify=False, proxies=proxy, timeout=30)
        if response.status_code != 200:
            logger.warning(f"PMC efetch API returned status {response.status_code} for PMC ID {pmc_id_num}")
            return ""

        try:
            root = ET.fromstring(response.content)
        except ET.ParseError as e:
            logger.warning(f"Failed to parse PMC XML for {pmc_id_num}: {e}")
            return ""

        article_parts = []

        # 提取标题
        title_el = root.find(".//article-title")
        if title_el is not None:
            title_text = "".join(title_el.itertext()).strip()
            if title_text:
                article_parts.append(f"# {title_text}\n")

        # 提取摘要
        abstract_el = root.find(".//abstract")
        if abstract_el is not None:
            article_parts.append("## Abstract\n")
            for sec in abstract_el.findall(".//sec"):
                sec_title = sec.find("title")
                if sec_title is not None:
                    article_parts.append(f"### {sec_title.text.strip()}\n")
                for p in sec.findall("p"):
                    p_text = "".join(p.itertext()).strip()
                    if p_text:
                        article_parts.append(p_text + "\n")
            # 也处理没有sec包裹的abstract段落
            for p in abstract_el.findall("p"):
                p_text = "".join(p.itertext()).strip()
                if p_text:
                    article_parts.append(p_text + "\n")

        # 提取正文
        body_el = root.find(".//body")
        if body_el is not None:
            for sec in body_el.findall(".//sec"):
                sec_title = sec.find("title")
                if sec_title is not None:
                    sec_title_text = "".join(sec_title.itertext()).strip()
                    if sec_title_text:
                        article_parts.append(f"\n## {sec_title_text}\n")
                for p in sec.findall("p"):
                    p_text = "".join(p.itertext()).strip()
                    if p_text:
                        article_parts.append(p_text + "\n")
            # 也处理没有sec包裹的body段落
            for p in body_el.findall("./p"):
                p_text = "".join(p.itertext()).strip()
                if p_text:
                    article_parts.append(p_text + "\n")

        return "\n".join(article_parts)

    def get_pubmed_article(self, pmid) -> MCPToolResult:
        logger.info(f"Attempting to access full text for PMID: {pmid}")
        try:
            # 获取PMC ID
            efetch_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?db=pubmed&id={pmid}&retmode=xml"
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
            }
            response = requests.get(efetch_url, headers=headers, verify=False, proxies=proxy)

            if response.status_code != 200:
                return MCPToolResult(success=False, error=f"Unable to fetch article data (status code: {response.status_code})")

            root = ET.fromstring(response.content)
            pmc_id = root.find(".//ArticleId[@IdType='pmc']")

            # 预先获取元数据（标题等），用于嵌入到保存的文件中
            metadata = get_pubmed_metadata(pmid)
            article_title = metadata.get("Title", "Unknown Title") if metadata else "Unknown Title"
            pubmed_url = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"

            # 无PMC ID，使用Abstract
            if pmc_id is None:
                logger.info(f"No PMC ID found for PMID: {pmid}, falling back to abstract only")
                if metadata and metadata.get("Abstract"):
                    file_path = f"./pubmed/{pmid}_abstract.txt"
                    # 嵌入元数据头，确保引用时能提取标题和URL
                    content_with_header = f"# {article_title}\nURL Source: {pubmed_url}\n\n{metadata['Abstract']}"
                    self.file_write(file_path=file_path, content=content_with_header, create_dirs=True)
                    return MCPToolResult(success=True, data={
                        "content": metadata["Abstract"],
                        "file_path": file_path,
                        "url": pubmed_url,
                        "source_type": "pubmed_abstract",
                        "pmid": pmid
                    })
                else:
                    return MCPToolResult(success=False, error=f"No PMC ID or Abstract found for PMID: {pmid}")
                
            pmc_id = pmc_id.text  # e.g. "PMC7096724"
            pmc_id_num = pmc_id.replace("PMC", "")  # e.g. "7096724"
            pmc_url = f"https://pmc.ncbi.nlm.nih.gov/articles/{pmc_id}/"
            logger.info(f"Found PMC ID: {pmc_id} for PMID: {pmid}")

            save_path = "./pubmed"
            os.makedirs(self.workspace_path / save_path, exist_ok=True)

            # ====== 方案1: PMC efetch XML API 获取全文（最可靠） ======
            try:
                logger.info(f"[方案1] Trying PMC efetch XML API for {pmc_id}")
                xml_fulltext = self._extract_pmc_fulltext_xml(pmc_id_num)
                if xml_fulltext and len(xml_fulltext.strip()) > 500:
                    txt_file = f"{save_path}/{pmid}.txt"
                    # 如果XML全文不是以#标题开头，则添加元数据头
                    if not xml_fulltext.strip().startswith('# '):
                        xml_fulltext = f"# {article_title}\nURL Source: {pmc_url}\n\n{xml_fulltext}"
                    else:
                        # 在标题后插入URL Source行
                        lines = xml_fulltext.split('\n', 1)
                        xml_fulltext = f"{lines[0]}\nURL Source: {pmc_url}\n{lines[1] if len(lines) > 1 else ''}"
                    with open(self.workspace_path / txt_file, 'w', encoding='utf-8') as f:
                        f.write(xml_fulltext)
                    logger.info(f"[方案1 成功] PMC XML full text saved as {txt_file}, length: {len(xml_fulltext)}")
                    return MCPToolResult(success=True, data={
                        "content": xml_fulltext,
                        "file_path": txt_file,
                        "url": pmc_url,
                        "source_type": "pubmed_pmc_xml",
                        "pmid": pmid,
                        "pmc_id": pmc_id
                    })
                else:
                    logger.warning(f"[方案1 失败] PMC XML content too short ({len(xml_fulltext.strip()) if xml_fulltext else 0} chars)")
            except Exception as e:
                logger.warning(f"[方案1 异常] PMC efetch XML failed for {pmc_id}: {e}")

            # ====== 方案2: 下载PDF并提取文本 ======
            try:
                logger.info(f"[方案2] Trying PDF download for {pmc_id}")
                pdf_url = f"https://pmc.ncbi.nlm.nih.gov/articles/{pmc_id}/pdf"
                pdf_response = requests.get(pdf_url, headers=headers, verify=False, proxies=proxy, timeout=30, allow_redirects=True)
                logger.info(f"[方案2] PDF response status: {pdf_response.status_code}, content-type: {pdf_response.headers.get('content-type', 'unknown')}")

                if pdf_response.status_code == 200 and 'pdf' in pdf_response.headers.get('content-type', '').lower():
                    temp_pdf = self.workspace_path / save_path / f"{pmid}_temp.pdf"
                    with open(temp_pdf, 'wb') as f:
                        f.write(pdf_response.content)
                    
                    pdf_text = self._read_pdf_text(temp_pdf)
                    
                    if temp_pdf.exists():
                        temp_pdf.unlink()
                    
                    if pdf_text and len(pdf_text.strip()) > 100:
                        txt_file = f"{save_path}/{pmid}.txt"
                        # 添加元数据头
                        pdf_text_with_header = f"# {article_title}\nURL Source: {pmc_url}\n\n{pdf_text}"
                        with open(self.workspace_path / txt_file, 'w', encoding='utf-8') as f:
                            f.write(pdf_text_with_header)
                        logger.info(f"[方案2 成功] PDF text saved as {txt_file}, length: {len(pdf_text)}")
                        return MCPToolResult(success=True, data={
                            "content": pdf_text,
                            "file_path": txt_file,
                            "url": pmc_url,
                            "source_type": "pubmed_pmc_pdf",
                            "pmid": pmid,
                            "pmc_id": pmc_id
                        })
                    else:
                        logger.warning(f"[方案2 失败] PDF text extraction too short or empty")
                else:
                    logger.warning(f"[方案2 失败] PDF not available or wrong content-type")
            except Exception as e:
                logger.warning(f"[方案2 异常] PDF download failed for {pmc_id}: {e}")

            # ====== 方案3: 抓取PMC HTML页面 ======
            article = ""
            try:
                logger.info(f"[方案3] Trying PMC HTML scraping for {pmc_id}")
                pmc_response = requests.get(pmc_url, headers=headers, verify=False, proxies=proxy, timeout=30)
                logger.info(f"[方案3] HTML response status: {pmc_response.status_code}")

                if pmc_response.status_code == 200:
                    soup = BeautifulSoup(pmc_response.content, 'html.parser')
                    
                    # 提取标题 - 尝试多种选择器
                    title = soup.find('h1', class_='content-title')
                    if not title:
                        title = soup.find('h1', id='article-title')
                    if not title:
                        title = soup.select_one('.article-title, .head-title, h1.heading-title')
                    if title:
                        article += f"# {title.get_text(strip=True)}\n\n"
                    
                    # 提取摘要
                    abstract = soup.find('div', class_='abstract')
                    if not abstract:
                        abstract = soup.find('section', class_='abstract')
                    if not abstract:
                        abstract = soup.select_one('#abstract, .abstract-content')
                    if abstract:
                        article += "## Abstract\n\n"
                        article += abstract.get_text(separator='\n', strip=True) + "\n\n"
                    
                    # 提取正文 - 尝试多种选择器
                    main_content = soup.find('div', class_='jig-ncbiinpagenav')
                    if not main_content:
                        main_content = soup.find('div', id='mc')
                    if not main_content:
                        main_content = soup.find('div', class_='article-body')
                    if not main_content:
                        main_content = soup.find('main')
                    if not main_content:
                        main_content = soup.find('article')
                    if not main_content:
                        # 尝试所有 tsec class 的 div
                        tsec_divs = soup.find_all('div', class_='tsec')
                        if tsec_divs:
                            for tsec in tsec_divs:
                                for el in tsec.find_all(['p', 'h2', 'h3', 'h4']):
                                    text = el.get_text(strip=True)
                                    if text:
                                        if el.name in ['h2', 'h3', 'h4']:
                                            article += f"\n## {text}\n\n"
                                        else:
                                            article += text + "\n\n"
                    
                    if main_content:
                        for section in main_content.find_all(['p', 'h2', 'h3', 'h4']):
                            text = section.get_text(strip=True)
                            if text:
                                if section.name in ['h2', 'h3', 'h4']:
                                    article += f"\n## {text}\n\n"
                                else:
                                    article += text + "\n\n"
                    
                    logger.info(f"[方案3] Extracted HTML content length: {len(article)}")
            except Exception as e:
                logger.warning(f"[方案3 异常] HTML scraping failed for {pmc_id}: {e}")

            if article and len(article.strip()) > 200:
                txt_file = f"{save_path}/{pmid}.txt"
                # 如果HTML内容不是以#标题开头，添加元数据头
                if not article.strip().startswith('# '):
                    article = f"# {article_title}\nURL Source: {pmc_url}\n\n{article}"
                else:
                    lines = article.split('\n', 1)
                    article = f"{lines[0]}\nURL Source: {pmc_url}\n{lines[1] if len(lines) > 1 else ''}"
                with open(self.workspace_path / txt_file, 'w', encoding='utf-8') as f:
                    f.write(article)
                logger.info(f"[方案3 成功] HTML content saved as {txt_file}")
                return MCPToolResult(success=True, data={
                    "content": article,
                    "file_path": txt_file,
                    "url": pmc_url,
                    "source_type": "pubmed_pmc_html",
                    "pmid": pmid,
                    "pmc_id": pmc_id
                })

            # ====== 所有方案失败，回退到摘要 ======
            logger.warning(f"All full-text methods failed for PMID {pmid} ({pmc_id}), falling back to abstract")
            if metadata and metadata.get("Abstract"):
                file_path = f"./pubmed/{pmid}_abstract.txt"
                content_with_header = f"# {article_title}\nURL Source: {pmc_url}\n\n{metadata['Abstract']}"
                self.file_write(file_path=file_path, content=content_with_header, create_dirs=True)
                return MCPToolResult(success=True, data={
                    "content": metadata["Abstract"],
                    "file_path": file_path,
                    "url": pmc_url,
                    "source_type": "pubmed_abstract",
                    "pmid": pmid,
                    "pmc_id": pmc_id
                })
            else:
                return MCPToolResult(success=False, error=f"Failed to get any content for PMID: {pmid}")
            
        except Exception as e:
            logger.error(f"Error in get_pubmed_article: {e}")
            return MCPToolResult(success=False, error=f"获取pubmed论文内容失败!{e}")

    def arxiv_search(self, query: str, max_results: int = 10) -> MCPToolResult:
        BASE_URL = "http://export.arxiv.org/api/query"
        params = {
            'search_query': query,
            'max_results': max_results,
            'sortBy': 'submittedDate',
            'sortOrder': 'descending'
        }
        response = requests.get(BASE_URL, params=params, verify=False, proxies=proxy)
        feed = feedparser.parse(response.content)
        papers = []
        for entry in feed.entries:
            try:
                authors = [author.name for author in entry.authors]
                published = datetime.strptime(entry.published, '%Y-%m-%dT%H:%M:%SZ')
                updated = datetime.strptime(entry.updated, '%Y-%m-%dT%H:%M:%SZ')
                pdf_url = next((link.href for link in entry.links if link.type == 'application/pdf'), '')
                papers.append(Paper(
                    paper_id=entry.id.split('/')[-1],
                    title=entry.title,
                    authors=authors,
                    abstract=entry.summary,
                    url=entry.id,
                    pdf_url=pdf_url,
                    published_date=published,
                    updated_date=updated,
                    source='arxiv',
                    categories=[tag.term for tag in entry.tags],
                    keywords=[],
                    doi=entry.get('doi', '')
                ).to_dict())
            except Exception as e:
                return MCPToolResult(success=False, error=f"获取arxiv论文信息失败!{e}")
        return MCPToolResult(success=True, data={"papers": papers})

    def download_pdf(self, paper_id: str, save_path: str) -> str:
        pdf_url = f"https://arxiv.org/pdf/{paper_id}.pdf"
        response = requests.get(pdf_url, verify=False, proxies=proxy)
        
        # 确保目录存在
        os.makedirs(save_path, exist_ok=True)
        
        # 先保存 PDF 文件到临时位置
        temp_pdf_path = Path(save_path) / f"{paper_id}.pdf"
        with open(temp_pdf_path, 'wb') as f:
            f.write(response.content)
        
        # 提取 PDF 文本内容
        try:
            extracted_text = self._read_pdf_text(temp_pdf_path)
            if not extracted_text or len(extracted_text.strip()) < 100:
                logger.warning(f"PDF text extraction failed or content too short for {paper_id}, keeping PDF file")
                # 如果提取失败，保留 PDF 文件
                return str(temp_pdf_path)
            
            # 保存提取的文本
            output_file = Path(save_path) / f"{paper_id}.txt"
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(extracted_text)
            
            # 删除临时 PDF 文件
            temp_pdf_path.unlink()
            
            logger.info(f"Successfully extracted text from arXiv paper {paper_id}")
            return str(output_file)
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF {paper_id}: {e}")
            # 提取失败时保留 PDF 文件
            return str(temp_pdf_path)

    def arxiv_read_paper(self, paper_id: str, save_path: str = "./arxiv") -> MCPToolResult:
        """Read a paper and convert it to text format.

        Args:
            paper_id: arXiv paper ID
            save_path: Directory where the PDF is/will be saved (relative to workspace)
                      Note: Should be a directory path, not a file path

        Returns:
            str: The extracted text content of the paper
        """
        try:
            # 将 save_path 转换为相对于 workspace_path 的绝对路径
            # 如果 save_path 以 ./ 开头，去掉它
            if save_path.startswith('./'):
                save_path = save_path[2:]
            
            # 检查 save_path 是否错误地包含了文件名（以 .txt 或其他扩展名结尾）
            # 如果是，提取目录部分
            if save_path.endswith('.txt') or save_path.endswith('.pdf'):
                logger.warning(f"save_path appears to be a file path: {save_path}, extracting directory")
                save_path = str(Path(save_path).parent)
            
            # 构建完整路径（相对于 workspace）
            full_save_path = self.workspace_path / save_path
            
            txt_path = full_save_path / f"{paper_id}.txt"
            pdf_path = full_save_path / f"{paper_id}.pdf"

            # 计算相对于 workspace 的文件路径（用于 document_extract）
            relative_txt_path = str(Path(save_path) / f"{paper_id}.txt")
            
            # 构建 arXiv URL（用于参考文献）
            arxiv_url = f"https://arxiv.org/abs/{paper_id}"
            
            # 如果已经存在文本文件，直接读取
            if txt_path.exists():
                with open(txt_path, 'rb') as f:
                    content = f.read()
                return MCPToolResult(success=True, data={
                    "paper": content.decode('utf-8', errors='ignore'),
                    "file_path": relative_txt_path,
                    "url": arxiv_url,
                    "source_type": "arxiv",
                    "paper_id": paper_id
                })

            # 如果存在旧的PDF文件，重命名为.txt
            if pdf_path.exists():
                pdf_path.rename(txt_path)
                logger.info(f"已将arxiv论文 {paper_id} 的PDF文件重命名为.txt格式")
                with open(txt_path, 'rb') as f:
                    content = f.read()
                return MCPToolResult(success=True, data={
                    "paper": content.decode('utf-8', errors='ignore'),
                    "file_path": relative_txt_path,
                    "url": arxiv_url,
                    "source_type": "arxiv",
                    "paper_id": paper_id
                })

            # 下载文件（download_pdf现在会直接保存为.txt）
            txt_path_str = self.download_pdf(paper_id, str(full_save_path))
            with open(txt_path_str, 'rb') as f:
                content = f.read()
            return MCPToolResult(success=True, data={
                "paper": content.decode('utf-8', errors='ignore'),
                "file_path": relative_txt_path,
                "url": arxiv_url,
                "source_type": "arxiv",
                "paper_id": paper_id
            })

        except Exception as e:
            return MCPToolResult(success=False, error=f"获取arxiv论文内容失败!{e}")

    def advanced_google_scholar_search(self, query: str, author: str = None, start_year: int = None, 
                                       end_year: int = None, num_results: int = 5) -> MCPToolResult:
        """
        Search Google Scholar using advanced search filters (e.g., author, year range).

        Args:
            query: The search query (e.g., paper title or topic)
            author: The author's name to filter the results (optional)
            start_year: Start year to filter the results by publication year (optional)
            end_year: End year to filter the results by publication year (optional)
            num_results: The number of results to retrieve (default: 5)

        Returns:
            MCPToolResult: Standardized result format with search results
        """
        try:
            search_url = "https://scholar.google.com/scholar?"

            search_params = {'q': query.replace(' ', '+')}
            if author:
                search_params['as_auth'] = author
            if start_year:
                search_params['as_ylo'] = start_year
            if end_year:
                search_params['as_yhi'] = end_year

            search_url += '&'.join([f"{key}={value}" for key, value in search_params.items()])

            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            }

            response = requests.get(search_url, headers=headers, proxies=proxy, timeout=5, verify=False)

            if response.status_code != 200:
                return MCPToolResult(
                    success=False,
                    error=f"Failed to fetch data. HTTP Status code: {response.status_code}"
                )

            soup = BeautifulSoup(response.text, 'html.parser')

            results = []
            count = 0

            for item in soup.find_all('div', class_='gs_ri'):
                if count >= num_results:
                    break

                title_tag = item.find('h3', class_='gs_rt')
                title = title_tag.get_text() if title_tag else 'No title available'

                link = title_tag.find('a')['href'] if title_tag and title_tag.find('a') else 'No link available'

                authors_tag = item.find('div', class_='gs_a')
                authors = authors_tag.get_text() if authors_tag else 'No authors available'

                abstract_tag = item.find('div', class_='gs_rs')
                abstract = abstract_tag.get_text() if abstract_tag else 'No abstract available'

                result_data = {
                    'Title': title,
                    'Authors': authors,
                    'Abstract': abstract,
                    'URL': link
                }
                results.append(result_data)
                count += 1

            return MCPToolResult(
                success=True,
                data={
                    "results": results,
                    "total_results": len(results),
                    "query": query
                },
                metadata={"tool_name": "advanced_google_scholar_search"}
            )

        except Exception as e:
            logger.error(f"Advanced Google Scholar search failed: {e}")
            return MCPToolResult(
                success=False,
                error=f"Search failed: {str(e)}"
            )

    def google_scholar_search(self, query: str, num_results: int = 5) -> MCPToolResult:
        """
        Search Google Scholar using a simple keyword query.

        Args:
            query: The search query (e.g., paper title or author)
            num_results: The number of results to retrieve (default: 5)

        Returns:
            MCPToolResult: Standardized result format with search results
        """
        try:
            search_url = f"https://scholar.google.com/scholar?q={query.replace(' ', '+')}"

            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            }

            response = requests.get(search_url, proxies=proxy, timeout=5, headers=headers, verify=False)

            if response.status_code != 200:
                return MCPToolResult(
                    success=False,
                    error=f"Failed to fetch data. HTTP Status code: {response.status_code}"
                )

            soup = BeautifulSoup(response.text, 'html.parser')

            results = []
            count = 0

            for item in soup.find_all('div', class_='gs_ri'):
                if count >= num_results:
                    break

                title_tag = item.find('h3', class_='gs_rt')
                title = title_tag.get_text() if title_tag else 'No title available'

                link = title_tag.find('a')['href'] if title_tag and title_tag.find('a') else 'No link available'

                authors_tag = item.find('div', class_='gs_a')
                authors = authors_tag.get_text() if authors_tag else 'No authors available'

                abstract_tag = item.find('div', class_='gs_rs')
                abstract = abstract_tag.get_text() if abstract_tag else 'No abstract available'

                result_data = {
                    'Title': title,
                    'Authors': authors,
                    'Abstract': abstract,
                    'URL': link
                }
                results.append(result_data)
                count += 1

            return MCPToolResult(
                success=True,
                data={
                    "results": results,
                    "total_results": len(results),
                    "query": query
                },
                metadata={"tool_name": "google_scholar_search"}
            )

        except Exception as e:
            logger.error(f"Google Scholar search failed: {e}")
            return MCPToolResult(
                success=False,
                error=f"Search failed: {str(e)}"
            )

    def google_scholar_get_paper(self, paper_url: str) -> MCPToolResult:
        """
        Download and analyze a paper from Google Scholar search results.
        Similar to arxiv_read_paper and get_pubmed_article.
        
        Args:
            paper_url: The paper URL from google_scholar_search results
            
        Returns:
            MCPToolResult: Paper analysis result with file path
        """
        try:
            # Generate file path for saving
            import hashlib
            url_hash = hashlib.md5(paper_url.encode()).hexdigest()[:8]
            
            # Check if URL points to a PDF file
            is_pdf_url = paper_url.lower().endswith('.pdf') or '/pdf/' in paper_url.lower()
            
            if is_pdf_url:
                # Direct PDF download
                filename = f"google_scholar_{url_hash}.pdf"
                file_path = self.workspace_path / "url_crawler_save_files" / filename
                file_path.parent.mkdir(parents=True, exist_ok=True)
                
                try:
                    # Download PDF file directly
                    import requests
                    response = requests.get(paper_url, stream=True, timeout=30, verify=False)
                    response.raise_for_status()
                    
                    # Save PDF file
                    with open(file_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)
                    
                    saved_file_path = str(file_path.relative_to(self.workspace_path))
                    paper_title = "Google Scholar Paper"
                    
                except Exception as e:
                    logger.error(f"Failed to download PDF: {e}")
                    return MCPToolResult(
                        success=False,
                        error=f"Failed to download PDF: {str(e)}"
                    )
                
            else:
                # Use url_crawler for web pages
                filename = f"google_scholar_{url_hash}.txt"
                file_path = f"url_crawler_save_files/{filename}"
                
                crawler_result = self.url_crawler(documents=[{
                    "url": paper_url, 
                    "title": "Google Scholar Paper",
                    "file_path": file_path
                }])
                
                if not crawler_result.success:
                    return MCPToolResult(
                        success=False,
                        error=f"Failed to fetch paper: {crawler_result.error}"
                    )
                
                # Extract content from crawler result
                # url_crawler returns data as a list of result dicts directly
                crawled_data = crawler_result.data if isinstance(crawler_result.data, list) else []
                if not crawled_data:
                    return MCPToolResult(
                        success=False,
                        error="No content retrieved from URL"
                    )
                
                result_data = crawled_data[0]
                
                # Check if crawling was successful
                if not result_data.get('success', False):
                    return MCPToolResult(
                        success=False,
                        error=f"Failed to crawl URL: {result_data.get('error', 'Unknown error')}"
                    )
                
                paper_title = result_data.get('title', 'Google Scholar Paper')
                saved_file_path = result_data.get('file_path')
                
                if not saved_file_path:
                    return MCPToolResult(
                        success=False,
                        error="File path not found in crawler result"
                    )
            
            # Use document_extract to analyze the paper
            extract_result = self.document_extract(
                file_path=saved_file_path,
                task_description="Analyze this academic paper from Google Scholar"
            )
            
            if extract_result.success:
                return MCPToolResult(
                    success=True,
                    data={
                        "file_path": saved_file_path,
                        "title": paper_title,
                        "url": paper_url,
                        "analysis": extract_result.data,
                        "message": f"Successfully downloaded and analyzed paper from Google Scholar"
                    },
                    metadata={"tool_name": "google_scholar_get_paper"}
                )
            else:
                # Even if analysis fails, the file is still saved
                return MCPToolResult(
                    success=True,
                    data={
                        "file_path": saved_file_path,
                        "title": paper_title,
                        "url": paper_url,
                        "message": f"Paper downloaded but analysis failed: {extract_result.error}"
                    },
                    metadata={"tool_name": "google_scholar_get_paper"}
                )
                
        except Exception as e:
            logger.error(f"Google Scholar get paper failed: {e}")
            return MCPToolResult(
                success=False,
                error=f"Failed to get paper: {str(e)}"
            )

    def springer_search(self, query: str, max_results: int = 10, subject: str = None, 
                       start_year: int = None, end_year: int = None) -> MCPToolResult:
        """
        Search for papers on Springer Nature using their Open Access API.
        
        Args:
            query: Search query string (keywords, title, etc.)
            max_results: Maximum number of papers to return (default: 10)
            subject: Filter by subject area (optional)
            start_year: Filter by start year (optional)
            end_year: Filter by end year (optional)
            
        Returns:
            MCPToolResult with list of Paper objects
        """
        try:
            BASE_URL = "https://api.springernature.com/openaccess/json"
            
            params = {
                'q': query,
                'p': max_results,
                's': 1
            }
            
            if subject:
                params['q'] = f"{params['q']} subject:{subject}"
            
            if start_year and end_year:
                params['q'] = f"{params['q']} year:{start_year}-{end_year}"
            elif start_year:
                params['q'] = f"{params['q']} year:{start_year}-{datetime.now().year}"
            
            response = requests.get(BASE_URL, params=params, verify=False, proxies=proxy, timeout=30)
            
            if response.status_code != 200:
                return MCPToolResult(success=False, error=f"Springer API请求失败: HTTP {response.status_code}")
            
            data = response.json()
            records = data.get('records', [])
            
            papers = []
            for record in records:
                try:
                    pub_date_str = record.get('publicationDate', '')
                    try:
                        pub_date = datetime.strptime(pub_date_str, '%Y-%m-%d')
                    except:
                        try:
                            pub_date = datetime.strptime(pub_date_str, '%Y-%m')
                        except:
                            pub_date = datetime.now()
                    
                    creators = record.get('creators', [])
                    authors = [creator.get('creator', '') for creator in creators if isinstance(creator, dict)]
                    
                    doi = record.get('doi', '')
                    paper_id = doi if doi else record.get('identifier', '')
                    
                    url = record.get('url', [])
                    paper_url = url[0].get('value', '') if url and isinstance(url, list) and len(url) > 0 else f"https://doi.org/{doi}"
                    
                    pdf_url = ''
                    for url_item in url:
                        if isinstance(url_item, dict) and url_item.get('format', '') == 'pdf':
                            pdf_url = url_item.get('value', '')
                            break
                    
                    abstract = record.get('abstract', '')
                    
                    subjects = record.get('subjects', [])
                    categories = [subj.get('subject', '') for subj in subjects if isinstance(subj, dict)]
                    
                    paper = Paper(
                        paper_id=paper_id,
                        title=record.get('title', ''),
                        authors=authors,
                        abstract=abstract,
                        doi=doi,
                        published_date=pub_date,
                        pdf_url=pdf_url,
                        url=paper_url,
                        source='springer',
                        categories=categories,
                        keywords=[],
                        extra={
                            'publisher': record.get('publisher', ''),
                            'publicationType': record.get('publicationType', ''),
                            'issn': record.get('issn', ''),
                            'isbn': record.get('isbn', ''),
                            'volume': record.get('volume', ''),
                            'number': record.get('number', ''),
                            'startingPage': record.get('startingPage', ''),
                            'endingPage': record.get('endingPage', '')
                        }
                    )
                    papers.append(paper.to_dict())
                    
                except Exception as e:
                    logger.warning(f"解析Springer论文记录失败: {e}")
                    continue
            
            return MCPToolResult(success=True, data={"papers": papers, "total": len(papers)})
            
        except Exception as e:
            logger.error(f"Springer搜索失败: {e}")
            return MCPToolResult(success=False, error=f"Springer搜索失败: {e}")

    # DISABLED: Springer API currently unavailable
    # def springer_get_article(self, doi: str) -> MCPToolResult:
    #     """
    #     Get full article details from Springer Nature by DOI.
    #     
    #     Args:
    #         doi: Digital Object Identifier of the paper
    #         
    #     Returns:
    #         MCPToolResult with article content and metadata
    #     """
    #     pass

    def medrxiv_search(self, query: str, max_results: int = 10, days: int = 30) -> List[Paper]:
        """
        Search for papers on medRxiv within the last N days.
        Supports keyword matching in title/abstract and category filtering.

        Args:
            query: Search query - can be keywords (e.g., "COVID-19 vaccine") or 
                   category name (e.g., "infectious diseases").
            max_results: Maximum number of papers to return.
            days: Number of days to look back for papers.

        Returns:
            List of Paper objects matching the query within the specified date range.
        """
        # Calculate date range: last N days
        try:
            # 正确的API地址: api.biorxiv.org (不是 api.medrxiv.org)
            BASE_URL = "https://api.biorxiv.org/pubs/medrxiv"
            end_date = datetime.now().strftime('%Y-%m-%d')
            start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')

            # 准备搜索关键词（用于客户端过滤）
            query_lower = query.lower().strip()
            query_keywords = [kw.strip() for kw in query_lower.split() if len(kw.strip()) > 2]
            # 也准备类别格式用于匹配
            category_format = query_lower.replace(' ', '_')

            papers = []
            cursor = 0
            max_api_pages = 10  # 最多请求10页（1000篇）避免无限循环
            pages_fetched = 0
            
            while len(papers) < max_results and pages_fetched < max_api_pages:
                # API格式: /pubs/medrxiv/{start}/{end}/{cursor}
                url = f"{BASE_URL}/{start_date}/{end_date}/{cursor}"
                logger.info(f"medRxiv API request: {url}")

                tries = 0
                fetched_this_page = False
                while tries < self.max_retries:
                    try:
                        response = self.session.get(url, timeout=self.timeout, verify=False, proxies=proxy)
                        response.raise_for_status()
                        data = response.json()
                        collection = data.get('collection', [])
                        
                        for item in collection:
                            # 客户端过滤：匹配类别或关键词
                            item_category = item.get('category', '').lower().replace(' ', '_')
                            item_title = item.get('title', '').lower()
                            item_abstract = item.get('abstract', '').lower()
                            
                            # 匹配条件：类别匹配 OR 所有关键词出现在标题/摘要中
                            category_match = category_format and category_format in item_category
                            keyword_match = query_keywords and all(
                                kw in item_title or kw in item_abstract 
                                for kw in query_keywords
                            )
                            
                            if category_match or keyword_match:
                                date = datetime.strptime(item['date'], '%Y-%m-%d')
                                papers.append(Paper(
                                    paper_id=item['doi'],
                                    title=item['title'],
                                    authors=item['authors'].split('; '),
                                    abstract=item['abstract'],
                                    url=f"https://www.medrxiv.org/content/{item['doi']}v{item.get('version', '1')}",
                                    pdf_url=f"https://www.medrxiv.org/content/{item['doi']}v{item.get('version', '1')}.full.pdf",
                                    published_date=date,
                                    updated_date=date,
                                    source="medrxiv",
                                    categories=[item.get('category', '')],
                                    keywords=[],
                                    doi=item['doi']
                                ).to_dict())
                                
                                if len(papers) >= max_results:
                                    break
                        
                        fetched_this_page = True
                        if len(collection) < 100:
                            pages_fetched = max_api_pages  # 没有更多结果了
                        else:
                            cursor += 100
                        break  # Exit retry loop on success
                    except requests.exceptions.RequestException as e:
                        tries += 1
                        if tries == self.max_retries:
                            logger.error(f"Failed to connect to medRxiv API after {self.max_retries} attempts: {e}")
                            break
                        logger.error(f"Attempt {tries} failed, retrying...")
                
                pages_fetched += 1
                if not fetched_this_page:
                    break  # API请求失败，停止分页
            
            logger.info(f"medRxiv search found {len(papers)} papers for query '{query}' in {pages_fetched} pages")
            return MCPToolResult(success=True, data={"paper": papers})
        except Exception as e:
            return MCPToolResult(success=False, error=f"获取medrxiv论文内容失败!{e}")

    def medrxiv_download_pdf(self, paper_id: str, save_path: str) -> str:
        """
        Download a PDF for a given paper ID from medRxiv and extract text.

        Args:
            paper_id: The DOI of the paper.
            save_path: Directory to save the text file (relative to workspace).

        Returns:
            Path to the text file (relative to workspace).
        """
        if not paper_id:
            raise ValueError("Invalid paper_id: paper_id is empty")

        pdf_url = f"https://www.medrxiv.org/content/{paper_id}v1.full.pdf"
        tries = 0
        while tries < self.max_retries:
            try:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
                }
                response = self.session.get(pdf_url, timeout=self.timeout, headers=headers, verify=False, proxies=proxy)
                response.raise_for_status()
                
                # 创建目录
                os.makedirs(self.workspace_path / save_path, exist_ok=True)
                
                # 先保存PDF到临时文件
                temp_pdf_path = self.workspace_path / save_path / f"{paper_id.replace('/', '_')}_temp.pdf"
                with open(temp_pdf_path, 'wb') as f:
                    f.write(response.content)
                
                # 提取PDF文本
                extracted_text = self._read_pdf_text(temp_pdf_path)
                
                # 删除临时PDF
                if temp_pdf_path.exists():
                    temp_pdf_path.unlink()
                
                if not extracted_text or len(extracted_text.strip()) < 100:
                    raise Exception("PDF text extraction failed or content too short")
                
                # 保存为文本文件
                output_file = f"{save_path}/{paper_id.replace('/', '_')}.txt"
                with open(self.workspace_path / output_file, 'w', encoding='utf-8') as f:
                    f.write(extracted_text)
                    
                logger.info(f"Successfully extracted text from medRxiv PDF: {paper_id}")
                return output_file
            except requests.exceptions.RequestException as e:
                tries += 1
                if tries == self.max_retries:
                    raise Exception(f"Failed to download PDF after {self.max_retries} attempts: {e}")
                logger.error(f"Attempt {tries} failed, retrying...")

    def medrxiv_read_paper(self, paper_id: str, save_path: str = "./medrxiv") -> MCPToolResult:
        """
        Read a paper and convert it to text format.

        Args:
            paper_id: medRxiv DOI
            save_path: Directory where the PDF is/will be saved (relative to workspace)

        Returns:
            MCPToolResult: The extracted text content of the paper
        """
        try:
            txt_path = f"{save_path}/{paper_id.replace('/', '_')}.txt"
            pdf_path = f"{save_path}/{paper_id.replace('/', '_')}.pdf"
            
            # 使用workspace_path检查文件
            full_txt_path = self.workspace_path / txt_path
            full_pdf_path = self.workspace_path / pdf_path

            # 如果已经存在文本文件，直接读取
            if full_txt_path.exists():
                with open(full_txt_path, 'rb') as f:
                    content = f.read()
                return MCPToolResult(success=True, data={
                    "paper": content.decode('utf-8', errors='ignore'),
                    "file_path": txt_path,
                    "source_type": "medrxiv"
                })

            # 如果存在旧的PDF文件，重命名为.txt
            if full_pdf_path.exists():
                full_pdf_path.rename(full_txt_path)
                logger.info(f"已将medrxiv论文 {paper_id} 的PDF文件重命名为.txt格式")
                with open(full_txt_path, 'rb') as f:
                    content = f.read()
                return MCPToolResult(success=True, data={
                    "paper": content.decode('utf-8', errors='ignore'),
                    "file_path": txt_path,
                    "source_type": "medrxiv"
                })

            # 下载文件
            txt_path = self.medrxiv_download_pdf(paper_id, save_path)
            with open(self.workspace_path / txt_path, 'rb') as f:
                content = f.read()
            return MCPToolResult(success=True, data={
                "paper": content.decode('utf-8', errors='ignore'),
                "file_path": txt_path,
                "source_type": "medrxiv"
            })

        except Exception as e:
            return MCPToolResult(success=False, error=f"获取medrxiv论文内容失败!{e}")



def generate_pubmed_search_url(term=None, title=None, author=None, journal=None,
                               start_date=None, end_date=None, num_results=10):
    """根据用户输入的字段生成 PubMed 搜索 URL"""
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
    query_parts = []

    if term:
        query_parts.append(quote(term))
    if title:
        query_parts.append(f"{quote(title)}[Title]")
    if author:
        query_parts.append(f"{quote(author)}[Author]")
    if journal:
        query_parts.append(f"{quote(journal)}[Journal]")
    if start_date and end_date:
        query_parts.append(f"{start_date}:{end_date}[Date - Publication]")

    query = " AND ".join(query_parts)
    params = {
        "db": "pubmed",
        "term": query,
        "retmax": num_results,
        "retmode": "xml"
    }

    return f"{base_url}?{'&'.join([f'{k}={v}' for k, v in params.items()])}"


def search_pubmed(search_url):
    """从 PubMed 搜索结果中解析文章 ID"""

    response = requests.get(search_url, verify=False, proxies=proxy)

    if response.status_code == 200:
        root = ET.fromstring(response.content)
        id_list = root.find("IdList")
        if id_list is not None:
            return [id.text for id in id_list.findall("Id")]
        else:
            logger.info("No results found.")
            return []
    else:
        logger.error(f"Error: Unable to fetch data (status code: {response.status_code})")
        return []


def get_pubmed_metadata(pmid):
    """使用 PubMed API 通过 PMID 获取文章的详细元数据"""
    url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?db=pubmed&id={pmid}&retmode=xml"
    response = requests.get(url, verify=False, proxies=proxy)

    if response.status_code == 200:
        root = ET.fromstring(response.content)
        article = root.find(".//Article")
        if article is not None:
            title = article.find(".//ArticleTitle")
            title = title.text if title is not None else "No title available"

            abstract = article.find(".//Abstract/AbstractText")
            abstract = abstract.text if abstract is not None else "No abstract available"

            authors = []
            for author in article.findall(".//Author"):
                last_name = author.find(".//LastName")
                if last_name is not None and last_name.text:
                    authors.append(last_name.text)
            authors = ", ".join(authors) if authors else "No authors available"

            journal = article.find(".//Journal/Title")
            journal = journal.text if journal is not None else "No journal available"

            pub_date = article.find(".//PubDate/Year")
            pub_date = pub_date.text if pub_date is not None else "No publication date available"

            return {
                "PMID": pmid,
                "Title": title,
                "Authors": authors,
                "Journal": journal,
                "Publication Date": pub_date,
                "Abstract": abstract
            }
        else:
            logger.info(f"No article data found for PMID: {pmid}")
            return None
    else:
        logger.error(f"Error: Unable to fetch metadata (status code: {response.status_code})")
        return None


# ================ MCP TOOL SCHEMAS ================

MCP_TOOL_SCHEMAS = {
    "think": {
        "name": "think",
        "description": "Use the tool to think about something. It will not obtain new information or make any changes to the repository, but just log the thought. Use it when complex reasoning or brainstorming is needed.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "thought": {
                    "type": "string",
                    "description": "Your thoughts."
                }
            },
            "required": ["thought"]
        }
    },

    "reflect": {
        "name": "reflect",
        "description": "When multiple attempts yield no progress, use this tool to reflect on previous reasoning and planning, considering possible overlooked clues and exploring more possibilities. It will not obtain new information or make any changes to the repository.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "reflect": {
                    "type": "string",
                    "description": "The specific content of your reflection"
                }
            },
            "required": ["reflect"]
        }
    },

    "batch_web_search": {
        "name": "batch_web_search",
        "description": "Search multiple queries using configurable search API with concurrent processing (no more than 8 search queries)",
        "inputSchema": {
            "type": "object",
            "properties": {
                "queries": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of search queries"
                },
                "max_results_per_query": {
                    "type": "integer",
                    "default": 4,
                    "description": "Maximum search results per query (limited to 10)"
                },
                "max_workers": {
                    "type": "integer",
                    "default": 5,
                    "description": "Maximum number of concurrent search requests"
                }
            },
            "required": ["queries"]
        }
    },

    "url_crawler": {
        "name": "url_crawler",
        "description": "Extract content from web pages using configurable URL crawler API. Input is a list of documents with metadata including URL and local file path for saving extracted content.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "documents": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "url": {
                                "type": "string",
                                "description": "Web page URL to extract content from"
                            },
                            "file_path": {
                                "type": "string",
                                "description": "Local path to save extracted full text content"
                            },
                            "title": {
                                "type": "string",
                                "description": "Title of the web page"
                            },
                            "time": {
                                "type": "string",
                                "description": "Publication time of the web page"
                            }
                        },
                        "required": ["url", "file_path"]
                    },
                    "description": "List of documents with metadata including URL and save path"
                },
                "max_tokens_per_url": {
                    "type": "integer",
                    "default": 4000,
                    "description": "Maximum tokens per URL result"
                },
                "include_metadata": {
                    "type": "boolean",
                    "default": True,
                    "description": "Whether to include extraction metadata"
                },
                "max_workers": {
                    "type": "integer",
                    "default": 10,
                    "description": "Maximum number of concurrent extraction requests"
                }
            },
            "required": ["documents"]
        }
    },

    "concat_section_files": {
        "name": "concat_section_files",
        "description": "Concatenate the content of the saved section files into a single file",
        "inputSchema": {
            "type": "object",
            "properties": {
                "final_file_path": {
                    "type": "string",
                    "description": "The final file path to save the concatenated content, save the file in the workspace **under the relative path `./report/`**, and specify the final_file_path as `./report/final_report.md`"
                },
                "section_files": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the saved section file"
                            }
                        },
                        "required": ["file_path"]
                    },
                    "description": "List of section files to concatenate"
                }
            },
            "required": ["section_files", "final_file_path"]
        }
    },

    # TODO 需要修改schame的格式，还是存在错误
    "search_result_classifier": {
        "name": "search_result_classifier",
        "description": "Intelligently classify and organize search result files according to a structured outline for comprehensive long-form content generation. Analyzes files across fouer key dimensions (document time, source authority, core content, and task relevance) and assigns relevant files to appropriate outline sections. Files may be assigned to multiple sections when their content spans different topics.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "outline": {
                    "type": "string",
                    "description": "The outline here must be consistent with the content and structure of the outline generated above"
                },
                "key_files": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the file containing research content"
                            }
                        },
                        "required": ["file_path"]
                    },
                    "description": "List of research files to be classified according to the outline"
                },
                "model": {
                    "type": "string",
                    "default": "pangu_auto",
                    "description": "AI model to use for classification and organization"
                },
                "temperature": {
                    "type": "number",
                    "default": 0.3,
                    "description": "Creativity level for the AI classification (0-1)"
                },
                "max_tokens": {
                    "type": "integer",
                    "default": 2000,
                    "description": "Maximum tokens for the AI response"
                }
            },
            "required": ["key_files", "outline"]
        }
    },

    "document_qa": {
        "name": "document_qa",
        "description": "Answer questions based on content stored in local files. Each file has a corresponding question. Reads files and uses an AI model to answer each question using the respective file content as context.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "tasks": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the file (relative to workspace root)"
                            },
                            "question": {
                                "type": "string",
                                "description": "Question to ask about this file"
                            }
                        },
                        "required": ["file_path", "question"]
                    },
                    "description": "List of tasks, each containing a file path and a question"
                },
                "model": {
                    "type": "string",
                    "default": "gpt-4o-mini",
                    "description": "AI model to use for generating answers"
                },
                "temperature": {
                    "type": "number",
                    "default": 0.3,
                    "description": "Creativity level for the AI response (0-1)"
                },
                "max_tokens": {
                    "type": "integer",
                    "default": 2000,
                    "description": "Maximum tokens for the AI response"
                },
                "max_workers": {
                    "type": "integer",
                    "default": 5,
                    "description": "Maximum number of concurrent model API requests"
                }
            },
            "required": ["tasks"]
        }
    },

    "document_extract": {
        "name": "document_extract",
        "description": "Multi-dimensional analysis of locally stored files using AI models. Evaluates each file across four key dimensions: web page time extraction, source authority assessment, task relevance evaluation, and core content summarization (~300 words). Provides structured document analysis for research and content evaluation purposes.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "tasks": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the file (relative to workspace root)"
                            },
                            "task": {
                                "type": "string",
                                "description": "The content of the currently executed subtask"
                            }
                        },
                        "required": ["file_path", "task"]
                    },
                    "description": "List of tasks, each containing a file path and the current task"
                },
                "model": {
                    "type": "string",
                    "default": "pangu_auto",
                    "description": "AI model to use for generating answers"
                },
                "temperature": {
                    "type": "number",
                    "default": 0.3,
                    "description": "Creativity level for the AI response (0-1)"
                },
                "max_tokens": {
                    "type": "integer",
                    "default": 2000,
                    "description": "Maximum tokens for the AI response"
                },
                "max_workers": {
                    "type": "integer",
                    "default": 5,
                    "description": "Maximum number of concurrent model API requests"
                }
            },
            "required": ["tasks"]
        }
    },

    "section_writer": {
        "name": "section_writer",
        "description": "Write the current chapter content based on given web information and chapter structure; also consider user questions, completed chapters, and overall outline to ensure content relevance while avoiding duplication or contradictions.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "written_chapters_summary": {
                    "type": "string",
                    "description": "The summary of the written chapters, including the content of the chapters and the reflections on the chapters. Note that this field should be concatenated with the summaries of all previously written chapters with '\\n', and do not modify the original summary. For example, if the current chapter is the third chapter, the value of this field is 'chapter 1 summary \\n chapter 2 summary'. If not, the value is set to 'No previous chapters written yet.'"
                },
                "task_content": {
                    "type": "string",
                    "description": "Detailed description of some requirements for writing the current chapter and avoidance prompts. If there are reflections from the `think` tool on previously written chapters, they can be added to this field."
                },
                "user_query": {
                    "type": "string",
                    "description": "The user query, ensure the drafted content is highly relevant to the user's inquiry."
                },
                "current_chapter_outline": {
                    "type": "string",
                    "description": "This field represents the current chapter structure to be drafted. When composing the chapter content, do not modify content and bold formatting symbols of the existing structure's titles!!!"
                },
                "overall_outline": {
                    "type": "string",
                    "description": "This field represents the overall outline of the article. When drafting the chapter content, you should consider the overall outline to ensure the chapter content is consistent with the overall outline."
                },
                "target_file_path": {
                    "type": "string",
                    "description": "The path to save the chapter content"
                },
                "key_files": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the file containing research content"
                            }
                        },
                        "required": ["file_path"]
                    },
                    "description": "These files are the source materials required for drafting the current chapter."
                },
                "model": {
                    "type": "string",
                    "default": "pangu_auto",
                    "description": "AI model to use for classification and organization"
                },
                "temperature": {
                    "type": "number",
                    "default": 0.3,
                    "description": "Creativity level for the AI classification (0-1)"
                },
                "max_tokens": {
                    "type": "integer",
                    "default": 5000,
                    "description": "Maximum tokens for the AI response"
                },
            },
            "required": ["user_query", "current_chapter_outline", "overall_outline", "target_file_path", "key_files"]
        }
    },

    "download_files": {
        "name": "download_files",
        "description": "Download files from URLs to the workspace",
        "inputSchema": {
            "type": "object",
            "properties": {
                "urls": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of URLs to download"
                },
                "target_directory": {
                    "type": "string",
                    "description": "Directory to save files"
                },
                "overwrite": {
                    "type": "boolean",
                    "default": False,
                    "description": "Whether to overwrite existing files"
                },
                "max_file_size_mb": {
                    "type": "integer",
                    "default": 100,
                    "description": "Maximum file size in MB"
                }
            },
            "required": ["urls"]
        }
    },

    "process_user_uploaded_files": {
        "name": "process_user_uploaded_files",
        "description": "Process and download user-uploaded files from the Flask backend. This tool fetches files uploaded by users (e.g., PDFs, documents) and saves them to the workspace with high priority markers. Use this tool FIRST when user files are available to ensure they are analyzed before web search results.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_ids": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of file IDs from user uploads"
                },
                "backend_url": {
                    "type": "string",
                    "default": "http://localhost:5000",
                    "description": "Flask backend URL"
                }
            },
            "required": ["file_ids"]
        }
    },

    "process_library_files": {
        "name": "process_library_files",
        "description": "Process and download user-selected files from the document library. This tool fetches files that users have selected from their document library and saves them to the workspace. These files are treated equally with web search results - the LLM will judge their relevance and decide whether to cite them based on task_relevance, source_authority, and information_richness dimensions. Use this tool when users have selected specific files from their document library.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_ids": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of file IDs from document library"
                },
                "backend_url": {
                    "type": "string",
                    "default": "http://localhost:5000",
                    "description": "Flask backend URL"
                }
            },
            "required": ["file_ids"]
        }
    },

    "list_workspace": {
        "name": "list_workspace",
        "description": "List files and directories in the workspace",
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Specify the directory path to list, using a relative path"
                },
                "recursive": {
                    "type": "boolean",
                    "default": False,
                    "description": "Whether to list recursively"
                },
                "include_hidden": {
                    "type": "boolean",
                    "default": False,
                    "description": "Whether to include hidden files"
                },
                "max_depth": {
                    "type": "integer",
                    "default": 3,
                    "description": "Maximum recursion depth"
                }
            },
            "required": []
        }
    },

    "str_replace_based_edit_tool": {
        "name": "str_replace_based_edit_tool",
        "description": "Create, view, and edit files with various operations",
        "inputSchema": {
            "type": "object",
            "properties": {
                "action": {
                    "type": "string",
                    "enum": ["create", "view", "str_replace", "insert", "append", "delete"],
                    "description": "Action to perform"
                },
                "file_path": {
                    "type": "string",
                    "description": "Path to the file"
                },
                "content": {
                    "type": "string",
                    "description": "Content for create/insert/append actions"
                },
                "old_str": {
                    "type": "string",
                    "description": "String to replace (for str_replace)"
                },
                "new_str": {
                    "type": "string",
                    "description": "Replacement string (for str_replace)"
                },
                "line_number": {
                    "type": "integer",
                    "description": "Line number for insert action"
                }
            },
            "required": ["action", "file_path"]
        }
    },

    "file_read": {
        "name": "file_read",
        "description": "Read file content",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Relative path to the file (relative to workspace root)"
                },
                "encoding": {
                    "type": "string",
                    "default": "utf-8",
                    "description": "File encoding"
                }
            },
            "required": ["file_path"]
        }
    },

    "load_json": {
        "name": "load_json",
        "description": "Read json format file",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Relative path to the file (relative to workspace root)"
                },
                "encoding": {
                    "type": "string",
                    "default": "utf-8",
                    "description": "File encoding"
                }
            },
            "required": ["file_path"]
        }
    },

    "file_write": {
        "name": "file_write",
        "description": "Write content to file",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Relative path to the file (relative to workspace root)"
                },
                "content": {
                    "type": "string",
                    "description": "Content to write"
                },
                "encoding": {
                    "type": "string",
                    "default": "utf-8",
                    "description": "File encoding"
                },
                "create_dirs": {
                    "type": "boolean",
                    "default": True,
                    "description": "Create parent directories"
                }
            },
            "required": ["file_path", "content"]
        }
    },

    "file_grep_search": {
        "name": "file_grep_search",
        "description": "Search for pattern in files",
        "inputSchema": {
            "type": "object",
            "properties": {
                "pattern": {
                    "type": "string",
                    "description": "Regex pattern to search for"
                },
                "file_pattern": {
                    "type": "string",
                    "default": "*",
                    "description": "File pattern to search in"
                },
                "recursive": {
                    "type": "boolean",
                    "default": True,
                    "description": "Search recursively"
                },
                "ignore_case": {
                    "type": "boolean",
                    "default": False,
                    "description": "Ignore case in search"
                },
                "max_matches": {
                    "type": "integer",
                    "default": 100,
                    "description": "Maximum number of matches"
                }
            },
            "required": ["pattern"]
        }
    },

    "file_find_by_name": {
        "name": "file_find_by_name",
        "description": "Find files by name pattern",
        "inputSchema": {
            "type": "object",
            "properties": {
                "name_pattern": {
                    "type": "string",
                    "description": "Name pattern to search for"
                },
                "recursive": {
                    "type": "boolean",
                    "default": True,
                    "description": "Search recursively"
                },
                "case_sensitive": {
                    "type": "boolean",
                    "default": False,
                    "description": "Case sensitive search"
                },
                "max_results": {
                    "type": "integer",
                    "default": 100,
                    "description": "Maximum number of results"
                }
            },
            "required": ["name_pattern"]
        }
    },

    "bash": {
        "name": "bash",
        "description": "Execute bash command in the workspace",
        "inputSchema": {
            "type": "object",
            "properties": {
                "command": {
                    "type": "string",
                    "description": "Bash command to execute"
                },
                "timeout": {
                    "type": "integer",
                    "default": 30,
                    "description": "Command timeout in seconds"
                },
                "capture_output": {
                    "type": "boolean",
                    "default": True,
                    "description": "Whether to capture stdout/stderr"
                },
                "working_directory": {
                    "type": "string",
                    "description": "Working directory for command"
                }
            },
            "required": ["command"]
        }
    },

    "info_seeker_task_done": {
        "name": "info_seeker_task_done",
        "description": "Information Seeker Agent task completion reporting with information collection summary and related files.",
        "inputSchema": {
            "type": "object",
            "properties": {
                # "save_analysis_file_path": {
                #     "type": "string",
                #     "description": "The path to save the analysis file, save the analysis file in the workspace **under the relative path `./doc_analysis/`**, and specify the file path as `/doc_analysis/file_analysis.jsonl`"
                # },
                "task_summary": {
                    "type": "string",
                    "description": "Simple summary of what information has been collected for the current task and what new discoveries have been made.",
                    "format": "markdown"
                },
                "key_files": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "file_path": {
                                "type": "string",
                                "description": "Relative path to the file with collected content"
                            },
                        },
                        "required": ["file_path"]
                    },
                    "description": "Collect files highly relevant to this task. "
                },
                "completion_status": {
                    "type": "string",
                    "enum": ["completed", "partial", "failed"],
                    "description": "Final status of the information gathering task"
                },
                "completion_analysis": {
                    "type": "string",
                    "description": "Brief analysis of task completion quality, information thoroughness, and any limitations or gaps."
                }
            },
            "required": ["task_summary", "key_files", "completion_status", "completion_analysis"]
        }
    },

    "section_writer_task_done": {
        "name": "section_writer_task_done",
        "description": "Section Writer Agent task completion reporting for chapter/section writing. Called when a chapter, section, or paragraph is completed to provide a brief overview of the written content and completion status.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "chapter_summary": {
                    "type": "string",
                    "description": "Brief summary of the content written in the current chapter/section, including main topics covered and key points addressed.",
                    "format": "markdown"
                },
                "key_topics_covered": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of main topics or themes addressed in the written chapter/section"
                },
                "completion_status": {
                    "type": "string",
                    "enum": ["completed", "partial", "failed"],
                    "description": "Final status of the chapter/section writing task"
                },
                "completion_analysis": {
                    "type": "string",
                    "description": "Brief analysis of the writing task completion including: assessment of content quality, evaluation of outline adherence, identification of any challenges encountered, and overall evaluation of the writing process success."
                }
            },
            "required": ["chapter_summary", "key_topics_covered", "completion_status", "completion_analysis"]
        }
    },

    "writer_task_done": {
        "name": "writer_task_done",
        "description": "Writer Agent task completion reporting for complete long-form content. Called after all chapters/sections are written to provide a summary of the complete long article, final completion status and analysis, and the storage path of the final consolidated article.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "final_article_path": {
                    "type": "string",
                    "description": "The file path where the final article is saved."
                },
                "article_summary": {
                    "type": "string",
                    "description": "Comprehensive summary of the complete long-form article, including main themes, key points covered, and overall narrative structure.",
                    "format": "markdown"
                },
                "completion_status": {
                    "type": "string",
                    "enum": ["completed", "partial", "failed"],
                    "description": "Final status of the complete long-form writing task"
                },
                "completion_analysis": {
                    "type": "string",
                    "description": "Analysis of the overall writing project completion including: assessment of article coherence and quality, evaluation of content organization and flow, identification of any challenges in the writing process, and overall evaluation of the long-form content creation success."
                }
            },
            "required": ["final_article_path", "article_summary", "completion_status", "completion_analysis"]
        }
    },

    "semantic_search": {
        "name": "semantic_search",
        "description": "Search semantically through system-maintained knowledge index using OpenAI embeddings",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search query - can be natural language question or keywords"
                },
                "max_tokens": {
                    "type": "integer",
                    "default": 2000,
                    "description": "Maximum tokens to return in results (controls result size)"
                },
                "max_results": {
                    "type": "integer",
                    "default": 5,
                    "description": "Maximum number of results to return"
                },
                "similarity_threshold": {
                    "type": "number",
                    "default": 0.7,
                    "description": "Minimum similarity score (0-1) for results"
                },
                "filters": {
                    "type": "object",
                    "properties": {
                        "task_name": {
                            "type": "string",
                            "description": "Filter by specific task name"
                        },
                        "file_path": {
                            "type": "string",
                            "description": "Filter by files containing this path"
                        },
                        "is_final_output": {
                            "type": "boolean",
                            "description": "Filter by final output files only"
                        }
                    },
                    "description": "Optional filters to narrow search results"
                }
            },
            "required": ["query"]
        }
    },

    "knowledge_status": {
        "name": "knowledge_status",
        "description": "Get status and statistics about the system-managed knowledge index",
        "inputSchema": {
            "type": "object",
            "properties": {},
            "required": []
        }
    },

    "search_pubmed_key_words": {
        "name": "search_pubmed_key_words",
        "description": "Search for biological articles by keywords",
        "inputSchema": {
            "type": "object",
            "properties": {
                "keywords": {
                    "type": "string",
                    "description": "Search query string, only supports english"
                },
                "max_results": {
                    "type": "integer",
                    "description": "Number of results to return (default: 10)"
                }

            },
            "required": ["keywords"]
        }
    },

    "search_pubmed_advanced": {
        "name": "search_pubmed_advanced",
        "description": "Perform an advanced search for biological articles on PubMed.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "term": {
                    "type": "string",
                    "description": "General search term, only supports english"
                },
                "title": {
                    "type": "string",
                    "description": "Search in title, only supports english"
                },
                "author": {
                    "type": "string",
                    "description": "Author name, only supports english"
                },
                "journal": {
                    "type": "string",
                    "description": "Journal name, only supports english"
                },
                "start_date": {
                    "type": "string",
                    "description": "Start date for search range (format: YYYY/MM/DD)"
                },
                "end_date": {
                    "type": "string",
                    "description": "End date for search range (format: YYYY/MM/DD)"
                },
                "num_results": {
                    "type": "integer",
                    "description": "Number of results to return (default: 10)"
                }
            },
            "required": []
        }
    },
    "get_pubmed_article": {
        "name": "get_pubmed_article",
        "description": "Obtain articles of biology on PubMed via PMID. Before calling this function, first use search_key_words or search_advanced to obtain the article's PMID.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "pmid": {
                    "type": "string",
                    "description": "PMID"
                }
            },
            "required": ["pmid"]
        }
    },
    "arxiv_search": {
        "name": "arxiv_search",
        "description": "Searcher for arXiv papers, return the metadata of papers. You can get paper_id with this function and then use it for reading paper.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search query string, only supports english"
                },
                "max_results": {
                    "type": "integer",
                    "description": "Max number of searched papers"
                },
            },
            "required": ["query"]
        }
    },
    "arxiv_read_paper": {
        "name": "arxiv_read_paper",
        "description": "Obtain Arxiv article content via paper_id. Before calling this function, first use arxiv_search to obtain the article's paper_id.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "paper_id": {
                    "type": "string",
                    "description": "arXiv paper ID"
                },
                "save_path": {
                    "type": "string",
                    "description": "Directory where the PDF is/will be saved"
                }

            },
            "required": ["paper_id"]
        }
    },
    "medrxiv_search": {
        "name": "medrxiv_search",
        "description": "Searcher for biologically relevant papers, return the metadata of papers. You can get paper_id with this function and then use it for medrxiv_read_paper.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Category name to search for (e.g., \"cardiovascular medicine\"), only supports english"
                },
                "max_results": {
                    "type": "integer",
                    "description": "Max number of searched papers"
                },
                "days": {
                    "type": "integer",
                    "description": "Number of days to look back for papers."
                }
            },
            "required": ["query"]
        }
    },
    "medrxiv_read_paper": {
        "name": "medrxiv_read_paper",
        "description": "Obtain medrxiv article content via paper_id. Before calling this function, first use medrxiv_search to obtain the article's paper_id.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "paper_id": {
                    "type": "string",
                    "description": "medrxiv paper ID"
                },
                "save_path": {
                    "type": "string",
                    "description": "Directory where the PDF is/will be saved"
                }

            },
            "required": ["paper_id"]
        }
    },
    "google_scholar_search": {
        "name": "google_scholar_search",
        "description": "Search Google Scholar for academic papers using a simple keyword query. Returns paper titles, authors, abstracts and URLs. Covers all academic disciplines. Use this for broad academic search across all fields.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search query string (keywords, paper title, etc.), only supports english"
                },
                "num_results": {
                    "type": "integer",
                    "description": "Number of results to retrieve (default: 5)"
                }
            },
            "required": ["query"]
        }
    },
    "advanced_google_scholar_search": {
        "name": "advanced_google_scholar_search",
        "description": "Search Google Scholar with advanced filters including author name and year range. Returns paper titles, authors, abstracts and URLs. Use this when you need to filter by specific author or publication year range.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search query string (keywords, paper title, etc.), only supports english"
                },
                "author": {
                    "type": "string",
                    "description": "Author name to filter results (optional)"
                },
                "start_year": {
                    "type": "integer",
                    "description": "Start year to filter by publication year (optional)"
                },
                "end_year": {
                    "type": "integer",
                    "description": "End year to filter by publication year (optional)"
                },
                "num_results": {
                    "type": "integer",
                    "description": "Number of results to retrieve (default: 5)"
                }
            },
            "required": ["query"]
        }
    },
    "google_scholar_get_paper": {
        "name": "google_scholar_get_paper",
        "description": "Download and analyze a paper from Google Scholar search results. Similar to arxiv_read_paper and get_pubmed_article. Use this after google_scholar_search to get full paper content for analysis. The paper will be downloaded, saved to workspace, and analyzed automatically.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "paper_url": {
                    "type": "string",
                    "description": "The paper URL from google_scholar_search results"
                }
            },
            "required": ["paper_url"]
        }
    },
    # DISABLED: Springer API currently unavailable
    # "springer_search": {
    #     "name": "springer_search",
    #     "description": "Search for open access papers on Springer Nature across multiple disciplines. Returns metadata of papers including DOI, PDF links, and abstracts. Only searches open access content. Supports English queries only.",
    #     "inputSchema": {
    #         "type": "object",
    #         "properties": {
    #             "query": {
    #                 "type": "string",
    #                 "description": "Search query string (keywords, title, etc.), only supports English"
    #             },
    #             "max_results": {
    #                 "type": "integer",
    #                 "description": "Maximum number of papers to return (default: 10)"
    #             },
    #             "subject": {
    #                 "type": "string",
    #                 "description": "Filter by subject area (e.g., 'Computer Science', 'Earth Sciences', 'Life Sciences', 'Medicine', 'Physics', 'Chemistry', 'Mathematics', 'Engineering')"
    #             },
    #             "start_year": {
    #                 "type": "integer",
    #                 "description": "Filter by start year (e.g., 2020)"
    #             },
    #             "end_year": {
    #                 "type": "integer",
    #                 "description": "Filter by end year (e.g., 2024)"
    #             }
    #         },
    #         "required": ["query"]
    #     }
    # },
    # "springer_get_article": {
    #     "name": "springer_get_article",
    #     "description": "Get full article details from Springer Nature by DOI. Returns article content, metadata, and PDF link if available. Before calling this function, first use springer_search to obtain the article's DOI.",
    #     "inputSchema": {
    #         "type": "object",
    #         "properties": {
    #             "doi": {
    #                 "type": "string",
    #                 "description": "Digital Object Identifier (DOI) of the paper"
    #             }
    #         },
    #         "required": ["doi"]
    #     }
    # },
    "file_stats": {
        "name": "file_stats",
        "description": "Get comprehensive file statistics without reading full content - perfect for deciding reading strategy",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the file (relative to workspace)"
                }
            },
            "required": ["file_path"]
        }
    },

    "file_read_lines": {
        "name": "file_read_lines",
        "description": "Read specific line ranges from a file without loading entire file - perfect for large files",
        "inputSchema": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the file"
                },
                "start_line": {
                    "type": "integer",
                    "default": 1,
                    "description": "Starting line number (1-based)"
                },
                "end_line": {
                    "type": "integer",
                    "description": "Ending line number (1-based, None for end of file)"
                },
                "max_lines": {
                    "type": "integer",
                    "default": 1000,
                    "description": "Maximum number of lines to read (safety limit)"
                }
            },
            "required": ["file_path"]
        }
    }

    # NOTE: Task assignment tool schemas removed - these are now built-in methods of PlannerAgent
    # to avoid circular dependency issues with sub-agents trying to create MCP client connections
}


# ================ MAIN INTERFACE ================

def create_mcp_tools(workspace_path: str = None) -> MCPTools:
    """Create and return MCP tools instance"""
    return MCPTools(workspace_path)


def get_tool_schemas() -> Dict[str, Any]:
    """Get all tool schemas for MCP registration"""
    return MCP_TOOL_SCHEMAS
